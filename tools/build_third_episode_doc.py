from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Iterable, Sequence

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
except ModuleNotFoundError:  # Optional until the document command is actually used.
    Document = None
    WD_SECTION = WD_CELL_VERTICAL_ALIGNMENT = WD_TABLE_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = WD_BREAK = WD_LINE_SPACING = None
    OxmlElement = qn = RT = None
    Inches = Pt = RGBColor = None
    DOCX_AVAILABLE = False

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from lib.workspace_paths import WorkspacePaths

DEFAULT_PROJECT_ID = "003-tech-chat"
OUT_DIR = REPO_ROOT / "content" / "episodes" / DEFAULT_PROJECT_ID / "docs"
DOCX_PATH = OUT_DIR / "第三期_科技闲聊_抖音发布与数字人台词文档.docx"
MD_PATH = OUT_DIR / "第三期_科技闲聊_抖音发布与数字人台词文档.md"


def configure_output(project_id: str = DEFAULT_PROJECT_ID, output_dir: Path | None = None) -> None:
    """Configure output paths without relying on a machine-specific folder."""

    global OUT_DIR, DOCX_PATH, MD_PATH
    workspace = WorkspacePaths.from_repo_root(REPO_ROOT)
    OUT_DIR = (output_dir or workspace.project_docs(project_id)).resolve()
    DOCX_PATH = OUT_DIR / "第三期_科技闲聊_抖音发布与数字人台词文档.docx"
    MD_PATH = OUT_DIR / "第三期_科技闲聊_抖音发布与数字人台词文档.md"

SOURCE_URL = "https://help.openai.com/en/articles/20001371-evolving-atlas-into-chatgpt-for-browser-based-agentic-work"
SOURCE_TITLE = "OpenAI Help Center：Evolving Atlas into ChatGPT for browser-based agentic work"
CHECK_DATE = "2026-08-10"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "555555"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "1F6B52"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"

TITLE_MAIN = "第三期｜科技闲聊"
TITLE_SUB = "AI 浏览器刚开始流行，为什么 Atlas 却先退场？"

TITLE_OPTIONS = [
    ("新闻快讯型", "AI 浏览器刚开始流行，为什么 Atlas 却先退场？", "首发推荐；信息清楚，反差足，适合科技资讯受众。"),
    ("新闻快讯型", "Atlas 将停止运行：AI 浏览器的第一轮试验结束了吗？", "适合强调事件本身，标题更像一条科技快讯。"),
    ("反差悬念型", "一个 AI 浏览器为什么突然退场？答案可能和你想的不一样", "适合抖音信息流，先制造疑问，再在正文里给判断。"),
    ("反差悬念型", "OpenAI 放弃 Atlas，可能不是放弃 AI 浏览", "冲突感强，但发布时要避免把‘退场’解释成失败定论。"),
    ("观点判断型", "浏览器不是重点，AI 替你办事才是", "适合打造账号观点，强调节目核心判断。"),
    ("观点判断型", "AI 浏览器的未来，不在‘浏览器’三个字", "适合偏深度、偏播客的内容定位。"),
    ("生活化提问型", "以后买东西，真的不用自己一个个网页点了吗？", "把抽象的智能体趋势落到普通人的购物场景。"),
    ("生活化提问型", "AI 会替你搜、比、选，浏览器会变成什么？", "适合评论区讨论，问题导向明显。"),
    ("趋势观察型", "从 Atlas 退场，看智能体如何接管互联网", "适合做系列化科技观察，突出趋势价值。"),
    ("趋势观察型", "Atlas 的终点，可能是 AI 浏览器的起点", "适合做更有态度的封面标题，但需要正文支撑。"),
]

YAYA_LINES = [
    "最近有一个消息挺有意思，OpenAI 的 AI 浏览器 Atlas，原定在 8 月 9 日停止运行。",
    "这也是我觉得它有意思的地方。Atlas 可能不是因为“AI 浏览”这件事失败了，而是因为 OpenAI 发现，浏览器本身可能不是最重要的产品形态。",
    "对。用户真正需要的，可能是一个能够帮他搜索、阅读、整理、比较，甚至执行操作的 AI 助手。至于它是在一个新浏览器里完成，还是在 ChatGPT、Codex，甚至 Chrome 的侧边栏里完成，反而没那么重要。",
    "没错。浏览器正在从“一个你主动操作的工具”，变成“AI 代替你操作互联网的工作台”。",
    "这就是智能体时代最值得关注的地方。AI 不只是回答问题了，它可能开始读取网页、调用工具、下载文件，甚至代表你完成一部分任务。能力越强，权限管理和安全边界就越重要。",
    "我的判断是，AI 浏览器不会彻底消失，但它可能会被“吸收”进更大的产品里。以后我们未必会记住自己用的是哪个浏览器，但会越来越习惯让 AI 帮我们完成一整段流程。",
    "对。今天这个消息真正值得聊的，不是 Atlas 这款产品停不停，而是一个趋势：AI 正在从“帮你找答案”，走向“帮你完成事情”。",
    "欢迎把你的答案留在评论区。下一期，我们继续聊那些正在改变普通人生活方式的科技动态。",
]

MENGMENG_LINES = [
    "等一下，AI 浏览器不是刚刚开始热起来吗？怎么这么快就退场了？",
    "你的意思是，用户不一定需要单独安装一个 AI 浏览器？",
    "这就像以前大家会专门打开一个搜索引擎，但以后可能直接问 AI：“帮我找三款适合我的电脑，然后比较价格和配置。”",
    "但这也带来一个问题。以前我们自己点网页、看商品、做决定。以后如果 AI 替我们完成这些动作，它是不是会知道更多隐私？",
    "所以 Atlas 的退场，并不能说明 AI 浏览器没有未来。",
    "那未来浏览器的竞争，可能就不是谁的界面更漂亮，而是谁的 AI 更懂用户，谁的智能体更可靠。",
    "那你会愿意把浏览、搜索和购买交给 AI 吗？",
]

DIALOGUE = [
    ("雅雅", YAYA_LINES[0]),
    ("檬檬", MENGMENG_LINES[0]),
    ("雅雅", YAYA_LINES[1]),
    ("檬檬", MENGMENG_LINES[1]),
    ("雅雅", YAYA_LINES[2]),
    ("檬檬", MENGMENG_LINES[2]),
    ("雅雅", YAYA_LINES[3]),
    ("檬檬", MENGMENG_LINES[3]),
    ("雅雅", YAYA_LINES[4]),
    ("檬檬", MENGMENG_LINES[4]),
    ("雅雅", YAYA_LINES[5]),
    ("檬檬", MENGMENG_LINES[5]),
    ("雅雅", YAYA_LINES[6]),
    ("檬檬", MENGMENG_LINES[6]),
    ("雅雅", YAYA_LINES[7]),
]


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int]):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="D9E1EA", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_borders(paragraph, color=BLUE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "6")
        element.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    relationship_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Microsoft YaHei")
    fonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(fonts)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_callout(doc, title, body, fill=CALLOUT, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.keep_together = True
    set_para_spacing(p, before=3, after=7, line=1.25)
    shade_paragraph(p, fill)
    set_paragraph_borders(p, color=accent, size="8")
    r = p.add_run(title + "\n")
    set_run_font(r, size=10.5, color=accent, bold=True)
    r2 = p.add_run(body)
    set_run_font(r2, size=10.5, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_body(doc, text, bold_prefix=None, color=None):
    p = doc.add_paragraph()
    set_para_spacing(p, after=6, line=1.25)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=color or INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color or INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, after=4, line=1.25)
    for run in p.runs:
        set_run_font(run, size=11, color=INK)
    if not p.runs:
        r = p.add_run(text)
    else:
        p.runs[0].text = text
        r = p.runs[0]
    set_run_font(r, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para_spacing(p, after=4, line=1.25)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_speaker_line(doc, speaker, text):
    p = doc.add_paragraph()
    set_para_spacing(p, after=6, line=1.25)
    r1 = p.add_run(f"{speaker}：")
    set_run_font(r1, size=11, color=GREEN if speaker == "雅雅" else DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=INK)
    return p


def add_clean_script(doc, speaker, lines, note):
    add_heading(doc, f"{speaker}（台词纯净版）", 2)
    add_callout(doc, "生成注意", note, fill="F8FAFC", accent=GREEN if speaker == "雅雅" else DARK_BLUE)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        set_para_spacing(p, after=8, line=1.35)
        r = p.add_run(line)
        set_run_font(r, size=12, color=INK)


def add_meta_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1900, 7460])
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "内容"
    for cell in hdr:
        shade_cell(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            set_para_spacing(p, after=0, line=1.15)
            for r in p.runs:
                set_run_font(r, size=10.5, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for i, cell in enumerate(cells):
            if i == 0:
                shade_cell(cell, "F8FAFC")
            for p in cell.paragraphs:
                set_para_spacing(p, after=0, line=1.20)
                for r in p.runs:
                    set_run_font(r, size=10.5, color=INK, bold=(i == 0))
    return table


def add_title_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1550, 5400, 2410])
    set_table_borders(table, color="D9E1EA")
    headers = ["风格", "标题", "适用场景 / 注意"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = text
        shade_cell(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            set_para_spacing(p, after=0, line=1.15)
            for r in p.runs:
                set_run_font(r, size=10, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for style, title, note in TITLE_OPTIONS:
        cells = table.add_row().cells
        cells[0].text = style
        cells[1].text = title
        cells[2].text = note
        for idx, cell in enumerate(cells):
            if idx == 0:
                shade_cell(cell, "F8FAFC")
            for p in cell.paragraphs:
                set_para_spacing(p, after=0, line=1.16)
                for r in p.runs:
                    set_run_font(r, size=9.5, color=INK, bold=(idx == 0))
    return table


def build_docx():
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is required for DOCX output. Install project requirements first."
        )
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.20
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    # Quiet running furniture for a multi-page operator document.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(header, after=0, line=1.0)
    hr = header.add_run("第三期｜科技闲聊 · 抖音发布与数字人制作稿")
    set_run_font(hr, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(footer, after=0, line=1.0)
    fr = footer.add_run("OpenMontage 内容制作资料  |  第 ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(footer)
    fr2 = footer.add_run(" 页")
    set_run_font(fr2, size=9, color=MUTED)

    # Cover / title block.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(TITLE_MAIN)
    set_run_font(r, size=28, color=INK, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p2, after=12, line=1.20)
    r = p2.add_run(TITLE_SUB)
    set_run_font(r, size=16, color=BLUE, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p3, after=22, line=1.20)
    r = p3.add_run("抖音发布方案｜双人数字人播客｜可直接拆分生成雅雅与檬檬视频")
    set_run_font(r, size=11, color=MUTED)
    add_callout(
        doc,
        "本文件怎么用",
        "第 5 节用于生成完整对话稿；第 6、7 节是分别喂给两位数字人的纯净台词。生成数字人时，只复制纯净台词正文，不要把‘雅雅（台词纯净版）’、说明文字或发布文案一起复制。",
        fill=CALLOUT,
        accent=BLUE,
    )
    add_body(doc, f"资料核对时间：{CHECK_DATE}。新闻内容有时效性，正式发布前请按第 10 节重新核对产品状态。")
    doc.add_page_break()

    add_heading(doc, "1. 事实核对与节目口径", 1)
    add_body(doc, "本期不是把 Atlas 的退场简单解释成‘产品失败’，而是围绕一个更适合播客讨论的判断展开：浏览器可能只是 AI 代理执行任务的载体，真正有价值的是搜索、阅读、整理、比较和执行这一整段能力。")
    add_meta_table(doc, [
        ("官方事实", "OpenAI 官方帮助中心称正在逐步淘汰 Atlas，并将基于浏览器的智能体能力迁移到 ChatGPT 和 Codex；页面明确写明 Atlas 计划于 2026 年 8 月 9 日停止运行。"),
        ("节目观点", "这更像是产品形态迁移：能力进入更大的产品入口，而不一定意味着 AI 浏览这一方向没有未来。该判断是节目分析，不是 OpenAI 对 Atlas 退场原因的官方解释。"),
        ("发布风险", "不要使用‘Atlas 因为失败被砍’、‘OpenAI 彻底放弃 AI 浏览’等绝对化表述；应使用‘官方计划停止运行’、‘能力迁移’、‘我们判断’等口径。"),
        ("日期提醒", "当前资料核对日为 2026-08-10，已超过官方计划日期。若发布时客户端状态、官方页面或迁移安排有变化，应更新开头台词和画面角标。"),
    ])
    add_body(doc, "官方来源：")
    p = doc.add_paragraph()
    set_para_spacing(p, after=8, line=1.20)
    add_hyperlink(p, SOURCE_TITLE, SOURCE_URL)

    add_heading(doc, "2. 推荐发布方案", 1)
    add_meta_table(doc, [
        ("内容类型", "数字人播客 / 录音间双人聊天 / 科技趋势闲聊"),
        ("平台", "抖音，建议 9:16 竖屏；主体时长约 3 分 30 秒至 4 分 30 秒，最终以两位数字人的实际音频时长为准。"),
        ("人物关系", "雅雅负责提出事实、解释和判断；檬檬负责追问、生活化转译和安全隐私反问。"),
        ("核心受众", "对 AI、浏览器、智能体和效率工具感兴趣，但不一定熟悉行业术语的普通用户。"),
        ("一句话卖点", "Atlas 可能不是把 AI 浏览器带到了终点，而是提醒我们：未来重要的不是浏览器叫什么，而是 AI 能不能替你完成一整段事情。"),
        ("推荐主标题", "AI 浏览器刚开始流行，为什么 Atlas 却先退场？"),
        ("封面大字", "Atlas 要退场了？"),
        ("封面副标题", "AI 浏览器输了，还是换了形态？"),
    ])

    add_heading(doc, "3. 多风格视频标题库", 1)
    add_body(doc, "首发建议：优先使用标题 1 或标题 3。标题 1 信息准确、反差自然；标题 3 更适合抖音信息流，但需要在正文中及时解释‘退场’不是官方认定的失败。")
    add_title_table(doc)

    add_heading(doc, "4. 视频详情与发布文案", 1)
    add_heading(doc, "4.1 推荐视频详情", 2)
    add_body(doc, "OpenAI 官方页面宣布，Atlas 将进入停止运行节点，浏览器智能体能力会继续向 ChatGPT 和 Codex 等产品入口迁移。可这件事真正有意思的地方，不是某个浏览器产品退场，而是浏览器的角色正在变化：从我们主动点击的工具，变成 AI 代替我们操作互联网的工作台。")
    add_body(doc, "这期雅雅和檬檬会聊三个问题：AI 浏览器到底需要不需要独立存在？AI 代替我们搜索、比较、购买时，隐私边界怎么办？未来的浏览器竞争，究竟是界面竞争，还是智能体可靠性的竞争？")
    add_heading(doc, "4.2 抖音发布区文案（长版）", 2)
    add_callout(doc, "发布文案", "AI 浏览器刚开始热起来，Atlas 却迎来了停止运行节点。\n\n这不一定代表 AI 浏览失败了，也可能意味着：浏览器只是载体，真正重要的是 AI 能不能替你搜索、阅读、比较，甚至完成一整段任务。\n\n如果以后 AI 可以替你上网、选商品、处理网页，你愿意把多少权限交给它？", fill="F8FAFC", accent=BLUE)
    add_heading(doc, "4.3 抖音发布区文案（短版）", 2)
    add_body(doc, "Atlas 的退场，可能不是 AI 浏览器的终点，而是产品形态的一次迁移：从‘我来操作浏览器’，到‘AI 替我完成事情’。")
    add_heading(doc, "4.4 置顶评论与话题标签", 2)
    add_body(doc, "置顶评论：如果 AI 可以替你搜索、比较、下单，你会愿意把浏览和购买交给它吗？你最担心的是效率、隐私，还是出错？")
    add_body(doc, "推荐标签：#科技闲聊  #人工智能  #AI浏览器  #智能体  #OpenAI  #科技新闻  #效率工具")

    add_heading(doc, "5. 完整双人对话文案", 1)
    add_callout(doc, "录制规则", "以下为完整对话顺序。两位数字人的口播视频生成后，必须以音频内容为主时钟，再用字幕和说话人身份做三重核对。不要在后期另外叠加同内容的 TTS，避免嘴型、音频和字幕出现双轨。", fill="FFF9E6", accent=GOLD)
    for speaker, line in DIALOGUE:
        add_speaker_line(doc, speaker, line)

    add_heading(doc, "6. 雅雅（台词纯净版）", 1)
    add_clean_script(doc, "雅雅", YAYA_LINES, "复制生成时只保留下面 8 段正文；不要复制本标题、‘生成注意’或任何舞台说明。建议保持自然聊天速度，句号和问号处保留真实停顿。")

    add_heading(doc, "7. 檬檬（台词纯净版）", 1)
    add_clean_script(doc, "檬檬", MENGMENG_LINES, "复制生成时只保留下面 7 段正文；不要复制本标题、‘生成注意’或任何舞台说明。檬檬的语气应更像即时追问和生活化转译，不要读成新闻播报。")

    add_heading(doc, "8. 数字人视频生成与合成执行单", 1)
    add_heading(doc, "8.1 生成顺序", 2)
    for item in [
        "分别使用雅雅纯净台词和檬檬纯净台词生成两条数字人素材；如果工具支持按段落生成，优先一段一段生成并保留段落边界。",
        "每一位数字人都以自身视频自带音频为主轨，不额外覆盖同内容音频；需要补环境声时，只能在不影响主口播的情况下低音量加入。",
        "按完整对话稿的顺序交替拼接：雅雅 1 → 檬檬 1 → 雅雅 2 → 檬檬 2，以此类推。不要按‘人物全部说完再拼接’的方式排列。",
        "如果需要字幕，先从最终合成音频生成字幕，再把字幕与实际音频逐句比对；不要直接用脚本段落假定时间轴。",
        "保留原始素材、分段素材、合成中间版和最终版，文件名中写入版本号，便于发现顺序问题时返工。",
    ]:
        add_number(doc, item)
    add_heading(doc, "8.2 建议生成参数", 2)
    add_meta_table(doc, [
        ("语言", "普通话，偏自然聊天；专有名词 Atlas、ChatGPT、Codex 保持清晰，不要过度英文腔。"),
        ("语速", "0.95 至 1.00 倍作为起点；檬檬追问可以略快，雅雅解释段落略稳。最终以嘴型自然和信息可听清为准。"),
        ("情绪", "雅雅：有判断但不武断；檬檬：有好奇心、会追问、偶尔带一点惊讶。"),
        ("停顿", "逗号短停，句号中停，问号有明显但自然的抬头感；不要通过强行插入过长静音制造播客感。"),
        ("输出检查", "先看音频波形和字幕文本，再看嘴型同步、说话人身份、台词顺序和画面是否出现重复人物。"),
    ])

    add_heading(doc, "9. 录音间画面与节奏建议", 1)
    add_body(doc, "画面基调：真实、克制、有信息密度的科技播客。两位数字人不需要持续同框，建议通过‘双人对话 - 单人近景 - 信息卡 - 网页操作感 B-roll - 双人回到问题’形成节奏。")
    add_heading(doc, "9.1 建议的画面节点", 2)
    scene_rows = [
        ("0:00-0:06", "开头钩子", "双人录音间快速切入，屏幕出现‘Atlas 要退场了？’；雅雅第一句出现时，右侧保留标题卡，不要留大面积空白。"),
        ("0:06-0:35", "反差建立", "檬檬追问时切檬檬近景，随后短暂切回双人；突出‘刚热起来 / 却要退场’的反差。"),
        ("0:35-1:35", "解释产品形态", "用抽象浏览器窗口、标签页、搜索结果、整理清单等素材，解释 AI 从工具变成工作台。"),
        ("1:35-2:20", "生活化例子", "出现‘帮我找三款电脑并比较价格和配置’的聊天界面或信息卡，画面文字要短，避免把整段台词塞进画面。"),
        ("2:20-2:55", "隐私与权限", "画面转深色，加入权限、Cookie、登录、下载等抽象符号；用字幕强调‘能力越强，边界越重要’。"),
        ("2:55-3:45", "趋势判断", "雅雅近景与浏览器窗口、ChatGPT/Codex 能力迁移示意交替出现，避免右侧再次叠加同一位雅雅的静态图片。"),
        ("3:45-结尾", "评论区互动", "回到双人或雅雅主镜头，屏幕只保留一个问题：‘你愿意把浏览、搜索和购买交给 AI 吗？’。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1500, 2200, 5660])
    set_table_borders(table)
    for idx, text in enumerate(["时间段", "功能", "画面执行"]):
        cell = table.cell(0, idx)
        cell.text = text
        shade_cell(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            set_para_spacing(p, after=0, line=1.15)
            for r in p.runs:
                set_run_font(r, size=10, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in scene_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if idx == 0:
                shade_cell(cells[idx], "F8FAFC")
            for p in cells[idx].paragraphs:
                set_para_spacing(p, after=0, line=1.18)
                for r in p.runs:
                    set_run_font(r, size=9.5, color=INK, bold=(idx < 2))
    add_body(doc, "素材原则：可以使用抽象化浏览器界面、搜索页、权限弹窗、标签页和信息卡；若使用官方产品截图或 Logo，应保留来源语境，不要让画面暗示官方背书。")

    add_heading(doc, "10. 发布前审查清单", 1)
    add_body(doc, "本清单是这期视频的最低通过标准。任一项不通过，都应返工对应片段后再导出。")
    for item in [
        "事实审查：开头关于 Atlas 的日期、状态和‘能力迁移’表述与官方页面一致；没有把节目推论写成官方结论。",
        "音频审查：最终视频只有一条主口播音轨；没有数字人自带音频与外置 TTS 重复播放、相位叠音或音量忽大忽小。",
        "台词审查：完整对话顺序与第 5 节逐句一致；雅雅、檬檬没有互换；没有漏句、重复句或把下一段提前。",
        "字幕审查：字幕以实际最终音频重新生成或校准；每句入出点与说话内容一致，字幕不压脸、不贴右下角、不被平台控件遮挡。",
        "画面审查：说话人画面与声音一致；雅雅说话时不叠加重复的雅雅静态图片；信息卡只展示关键词，不堆满全文。",
        "节奏审查：0:00-0:06 有明确钩子；解释段有 B-roll 或信息卡变化；隐私段有视觉转折；结尾留出评论问题。",
        "平台审查：竖屏安全区、封面文字可读、背景音乐不盖过人声；简介和标签没有未经核实的绝对化宣传。",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "最终通过条件", "把视频静音看一遍，仍能通过画面看出‘Atlas 退场 → AI 浏览能力迁移 → 智能体权限问题 → 未来趋势’；再只听音频一遍，仍能听清两位角色、顺序和观点。两项都通过后再发布。", fill="EAF6F1", accent=GREEN)

    add_heading(doc, "附录：版本与来源记录", 1)
    add_meta_table(doc, [
        ("文档版本", "v1.0"),
        ("制作日期", CHECK_DATE),
        ("主题来源", SOURCE_TITLE),
        ("来源地址", SOURCE_URL),
        ("使用提醒", "如果节目不是在 2026-08-10 附近发布，发布前必须重新核对 Atlas 当前状态和 ChatGPT/Codex 相关能力的可用性。"),
    ])

    doc.core_properties.title = f"{TITLE_MAIN}｜{TITLE_SUB}"
    doc.core_properties.subject = "抖音发布文案、完整对话稿与双人数字人纯净台词"
    doc.core_properties.author = "OpenMontage 内容制作"
    doc.core_properties.comments = "资料核对时间：2026-08-10；发布前复核事实时效。"
    doc.save(DOCX_PATH)


def md_escape(text: str) -> str:
    return text.replace("\\", "\\\\")


def build_markdown():
    lines: list[str] = []
    lines += [f"# {TITLE_MAIN}", f"## {TITLE_SUB}", "", "抖音发布方案｜双人数字人播客｜可直接拆分生成雅雅与檬檬视频", "", f"资料核对时间：{CHECK_DATE}", ""]
    lines += ["> 使用说明：第 5 节是完整对话稿；第 6、7 节是分别喂给数字人的纯净台词。生成时只复制台词正文，不要复制标题、说明或发布文案。", ""]
    lines += ["## 1. 事实核对与节目口径", "", "- 官方事实：OpenAI 官方帮助中心称正在逐步淘汰 Atlas，并将基于浏览器的智能体能力迁移到 ChatGPT 和 Codex；页面明确写明 Atlas 计划于 2026 年 8 月 9 日停止运行。", "- 节目观点：这更像是产品形态迁移，能力进入更大的产品入口；这是节目分析，不是 OpenAI 对 Atlas 退场原因的官方解释。", "- 发布风险：不要使用‘Atlas 因为失败被砍’、‘OpenAI 彻底放弃 AI 浏览’等绝对化表述。", "- 日期提醒：当前资料核对日为 2026-08-10，已超过官方计划日期；发布前重新核对产品状态。", "", f"官方来源：[{SOURCE_TITLE}]({SOURCE_URL})", ""]
    lines += ["## 2. 推荐发布方案", "", "| 项目 | 内容 |", "|---|---|", "| 内容类型 | 数字人播客 / 录音间双人聊天 / 科技趋势闲聊 |", "| 平台 | 抖音 9:16 竖屏；建议 3 分 30 秒至 4 分 30 秒，以实际音频为准 |", "| 人物关系 | 雅雅负责事实、解释和判断；檬檬负责追问、生活化转译和隐私反问 |", "| 一句话卖点 | 浏览器可能只是 AI 代理执行任务的载体，真正有价值的是 AI 能否替你完成一整段事情 |", "| 推荐主标题 | AI 浏览器刚开始流行，为什么 Atlas 却先退场？ |", "| 封面大字 | Atlas 要退场了？ |", "| 封面副标题 | AI 浏览器输了，还是换了形态？ |", ""]
    lines += ["## 3. 多风格视频标题库", "", "首发建议：优先使用标题 1 或标题 3。", "", "| 风格 | 标题 | 适用场景 / 注意 |", "|---|---|---|"]
    lines += [f"| {style} | {title} | {note} |" for style, title, note in TITLE_OPTIONS]
    lines += ["", "## 4. 视频详情与发布文案", "", "### 4.1 推荐视频详情", "", "OpenAI 官方页面宣布，Atlas 将进入停止运行节点，浏览器智能体能力会继续向 ChatGPT 和 Codex 等产品入口迁移。可这件事真正有意思的地方，不是某个浏览器产品退场，而是浏览器的角色正在变化：从我们主动点击的工具，变成 AI 代替我们操作互联网的工作台。", "", "这期雅雅和檬檬会聊三个问题：AI 浏览器到底需不需要独立存在？AI 代替我们搜索、比较、购买时，隐私边界怎么办？未来的浏览器竞争，究竟是界面竞争，还是智能体可靠性的竞争？", "", "### 4.2 抖音发布区文案（长版）", "", "AI 浏览器刚开始热起来，Atlas 却迎来了停止运行节点。\n\n这不一定代表 AI 浏览失败了，也可能意味着：浏览器只是载体，真正重要的是 AI 能不能替你搜索、阅读、比较，甚至完成一整段任务。\n\n如果以后 AI 可以替你上网、选商品、处理网页，你愿意把多少权限交给它？", "", "### 4.3 抖音发布区文案（短版）", "", "Atlas 的退场，可能不是 AI 浏览器的终点，而是产品形态的一次迁移：从‘我来操作浏览器’，到‘AI 替我完成事情’。", "", "### 4.4 置顶评论与话题标签", "", "置顶评论：如果 AI 可以替你搜索、比较、下单，你会愿意把浏览和购买交给它吗？你最担心的是效率、隐私，还是出错？", "", "推荐标签：#科技闲聊  #人工智能  #AI浏览器  #智能体  #OpenAI  #科技新闻  #效率工具", ""]
    lines += ["## 5. 完整双人对话文案", "", "录制规则：以下为完整对话顺序。生成后以音频内容为主时钟，再用字幕和说话人身份做三重核对。不要另外叠加同内容的 TTS。", ""]
    lines += [f"**{speaker}：** {line}" for speaker, line in DIALOGUE]
    lines += ["", "## 6. 雅雅（台词纯净版）", "", "复制生成时只保留下面 8 段正文，不要复制本标题、说明或舞台文字。", ""]
    lines += [md_escape(line) for line in YAYA_LINES]
    lines += ["", "## 7. 檬檬（台词纯净版）", "", "复制生成时只保留下面 7 段正文，不要复制本标题、说明或舞台文字。", ""]
    lines += [md_escape(line) for line in MENGMENG_LINES]
    lines += ["", "## 8. 数字人视频生成与合成执行单", "", "1. 分别使用雅雅纯净台词和檬檬纯净台词生成两条数字人素材。", "2. 以每位数字人视频自带音频为主轨，不额外覆盖同内容音频。", "3. 按完整对话稿顺序交替拼接：雅雅 1 → 檬檬 1 → 雅雅 2 → 檬檬 2，以此类推。", "4. 字幕从最终合成音频生成或校准，不要用脚本段落假定时间轴。", "5. 普通话、自然聊天语气；语速建议 0.95-1.00 倍。", ""]
    lines += ["## 9. 录音间画面与节奏建议", "", "- 0:00-0:06：双人录音间快速切入，出现‘Atlas 要退场了？’。", "- 0:06-0:35：用双人和单人近景建立‘刚热起来 / 却要退场’的反差。", "- 0:35-1:35：浏览器窗口、标签页、搜索结果、整理清单等 B-roll，解释‘工作台’。", "- 1:35-2:20：用‘帮我找三款电脑并比较价格和配置’的信息卡承接生活化例子。", "- 2:20-2:55：权限、Cookie、登录、下载等画面，强调安全边界。", "- 2:55-结尾：趋势判断后回到双人或雅雅主镜头，留下评论问题。", "", "素材原则：可使用抽象化浏览器界面和信息卡；官方截图或 Logo 要保留来源语境，不要暗示官方背书。", ""]
    lines += ["## 10. 发布前审查清单", "", "- 事实与日期表述和官方页面一致。", "- 只有一条主口播音轨，没有 TTS 重复、相位叠音或音量跳变。", "- 台词顺序与完整对话稿逐句一致，没有漏句、重复句或人物互换。", "- 字幕按最终音频重新生成或校准，不压脸、不贴右下角、不被平台控件遮挡。", "- 说话人画面与声音一致，不叠加重复的雅雅静态图。", "- 0:00-0:06 有钩子，解释段有视觉变化，结尾有评论问题。", "- 封面和简介没有未经核实的绝对化宣传。", "", "最终通过条件：静音看画面能看懂主线；只听音频能听清角色、顺序和观点。", ""]
    lines += ["## 附录：版本与来源记录", "", f"- 文档版本：v1.0", f"- 制作日期：{CHECK_DATE}", f"- 来源：[{SOURCE_TITLE}]({SOURCE_URL})", "- 发布前：重新核对 Atlas 当前状态和 ChatGPT/Codex 相关能力的可用性。", ""]
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the third-episode publishing and digital-human document")
    parser.add_argument("--project", default=DEFAULT_PROJECT_ID, help="Project id under content/episodes")
    parser.add_argument("--output-dir", type=Path, help="Optional output directory; defaults to the project docs directory")
    args = parser.parse_args()
    configure_output(args.project, args.output_dir)
    build_docx()
    build_markdown()
    print(DOCX_PATH)
    print(MD_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
