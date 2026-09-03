from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TOOLS_DIR = Path(__file__).resolve().parent
REPO_ROOT = TOOLS_DIR.parent
if str(TOOLS_DIR) not in sys.path:
    sys.path.insert(0, str(TOOLS_DIR))

import build_third_episode_doc as ui


PROJECT_DIR = REPO_ROOT / "content" / "episodes" / "004-tech-brief"
OUT_DIR = PROJECT_DIR / "docs"
DOCX_PATH = OUT_DIR / "第四期_科技快报_抖音发布与数字人台词文档_v1.1.docx"
MD_PATH = OUT_DIR / "第四期_科技快报_抖音发布与数字人台词文档_v1.1.md"
CHECK_DATE = "2026-08-10"

SOURCES = [
    (
        "AI 漫剧数字角色商业化讨论",
        "新浪娱乐：短剧角色段宴、容寄侨拥有角色账号并参与商业化",
        "https://ent.sina.cn/2026-07-30/detail-inikphnv1416832.d.html",
    ),
    (
        "宇树科技申购热度",
        "抖音热榜聚合页：宇树科技今日申购",
        "https://hotflashnews.com/platform/douyin",
    ),
    (
        "荣耀 Robot Phone 发布信息",
        "IT之家：荣耀 Robot Phone 将于 8 月 12 日发布",
        "https://www.ithome.com/0/981/843.htm",
    ),
    (
        "苹果设备与千问相关讨论",
        "8 月 9 日热点归档：苹果设备接入千问相关话题",
        "https://hotflashnews.com/archive/2026-08-09",
    ),
]

TITLE_OPTIONS = [
    ("首发推荐", "AI角色开始带货了，真人演员还需要出镜吗？", "事件清楚、人物冲突直接，最贴合本期主讨论。"),
    ("就业讨论", "数字人开始抢真人饭碗了吗？", "冲突更强，封面吸睛；正文必须及时解释并非整个人被替代。"),
    ("四条快报", "AI从演员走进手机：今天四条科技快报", "适合突出栏目属性，弱化单一新闻。"),
    ("数字人视角", "让数字人聊数字人就业，这事有点意思", "突出雅雅、檬檬身份，适合账号老观众。"),
    ("趋势判断", "AI不是只会聊天了，它开始进入工作现场", "适合把四条新闻串成同一条趋势。"),
]

DIALOGUE = [
    ("雅雅", "欢迎收听今天的《科技快报》。今天我们关注四条消息：AI角色开始带货、宇树科技申购、荣耀发布机器人手机，以及苹果设备接入千问。"),
    ("檬檬", "先看第一条。最近，AI漫剧里的虚拟角色段宴和容寄侨，不仅拥有固定形象、角色账号和粉丝，还带动了相关商品热销。"),
    ("雅雅", "重点不是他们又演了什么，而是一个不存在于现实中的数字人，已经可以像真人演员一样积累粉丝、参与宣传并产生商业价值。"),
    ("檬檬", "那我作为数字人就要问了：如果数字人可以出镜、介绍商品，还能长时间工作，真人演员和主播会不会少掉很多机会？"),
    ("雅雅", "最先受到影响的，可能是固定讲解和重复宣传。但数字人不会自己写内容、判断风险，背后依然需要真人创作、运营和承担责任。"),
    ("檬檬", "所以数字人不一定替代整个人，而是先拿走一部分比较重复的工作。"),
    ("雅雅", "对。真人的价值会更多转向创意、判断、真实体验和建立信任。未来的问题不是“真人还要不要出镜”，而是哪些内容必须由真人完成，哪些可以交给数字人。"),
    ("檬檬", "第二条，宇树科技开始申购。过去大家看机器人会不会翻跟头，现在更关心它能不能量产、进工厂、真正产生价值。"),
    ("雅雅", "机器人正在从舞台表演走向产业应用。"),
    ("檬檬", "第三条，荣耀 Robot Phone 将于8月12日发布。它不是把机器人塞进手机，而是让手机里的 AI 理解任务、调用应用，替你完成一些操作。"),
    ("雅雅", "手机正在从“你操作它”，变成“它帮你做事”。"),
    ("檬檬", "第四条，苹果设备接入千问的消息引发讨论。看起来是模型合作，背后其实是在争夺手机里的 AI 入口。"),
    ("雅雅", "以后你使用 AI，可能不需要单独打开一个软件，它会直接藏在手机系统、相册和语音助手里。"),
    ("檬檬", "四条消息放在一起看，AI正在进入内容、机器人、手机和操作系统。"),
    ("雅雅", "你最愿意把哪件事交给 AI？欢迎在评论区告诉我们。"),
]

YAYA_LINES = [text for speaker, text in DIALOGUE if speaker == "雅雅"]
MENGMENG_LINES = [text for speaker, text in DIALOGUE if speaker == "檬檬"]


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(ui.INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ui.BLUE, 18, 10),
        ("Heading 2", 13, ui.BLUE, 14, 7),
        ("Heading 3", 12, ui.DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(ui.INK)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ui.set_para_spacing(header, after=0, line=1.0)
    ui.set_run_font(header.add_run("第四期｜科技快报 · 抖音发布与数字人制作稿"), size=8.5, color=ui.MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ui.set_para_spacing(footer, after=0, line=1.0)
    ui.set_run_font(footer.add_run("OpenMontage 内容制作资料  |  第 "), size=9, color=ui.MUTED)
    ui.add_page_field(footer)
    ui.set_run_font(footer.add_run(" 页"), size=9, color=ui.MUTED)
    return doc


def add_title_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    ui.set_table_geometry(table, [1550, 5400, 2410])
    ui.set_table_borders(table)
    for idx, text in enumerate(["风格", "标题", "适用场景 / 注意"]):
        cell = table.cell(0, idx)
        cell.text = text
        ui.shade_cell(cell, ui.LIGHT_BLUE)
        for paragraph in cell.paragraphs:
            ui.set_para_spacing(paragraph, after=0, line=1.15)
            for run in paragraph.runs:
                ui.set_run_font(run, size=10, color=ui.INK, bold=True)
    ui.set_repeat_table_header(table.rows[0])
    for style, title, note in TITLE_OPTIONS:
        cells = table.add_row().cells
        for idx, value in enumerate([style, title, note]):
            cells[idx].text = value
            if idx == 0:
                ui.shade_cell(cells[idx], "F8FAFC")
            for paragraph in cells[idx].paragraphs:
                ui.set_para_spacing(paragraph, after=0, line=1.16)
                for run in paragraph.runs:
                    ui.set_run_font(run, size=9.5, color=ui.INK, bold=(idx == 0))


def add_sources(doc: Document) -> None:
    for topic, title, url in SOURCES:
        p = doc.add_paragraph()
        ui.set_para_spacing(p, after=5, line=1.2)
        ui.set_run_font(p.add_run(f"{topic}："), size=10.5, color=ui.INK, bold=True)
        ui.add_hyperlink(p, title, url)


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(8)
    ui.set_run_font(p.add_run("第四期｜科技快报"), size=28, color=ui.INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ui.set_para_spacing(p, after=12, line=1.2)
    ui.set_run_font(p.add_run("AI角色开始带货了，真人演员还需要出镜吗？"), size=16, color=ui.BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ui.set_para_spacing(p, after=22, line=1.2)
    ui.set_run_font(p.add_run("两分钟科技快讯｜四条热点｜雅雅与檬檬双数字人播客"), size=11, color=ui.MUTED)
    ui.add_callout(
        doc,
        "本文件怎么用",
        "第 5 节用 T001—T015 锁定最终对话顺序；推荐按轮次分别生成数字人视频。第 6、7 节按人物合并的纯净台词只用于必须整段生成的平台兼容。编号、标题、姓名和说明文字都不属于口播内容；成片必须控制在 2 分钟以内。",
        fill=ui.CALLOUT,
        accent=ui.BLUE,
    )
    ui.add_body(doc, f"资料核对时间：{CHECK_DATE}。新闻时效性强，发布前必须按第 10 节重新核验。")
    doc.add_page_break()

    ui.add_heading(doc, "1. 事实核对与节目口径", 1)
    ui.add_body(doc, "本期用四条新闻说明同一个趋势：AI 正从单独的软件能力，进入内容生产、机器人产业、手机操作和系统入口。主新闻聚焦数字人带货与就业变化，不展开漫剧剧情。")
    ui.add_meta_table(doc, [
        ("AI角色带货", "虚拟角色段宴、容寄侨近期因角色账号、粉丝运营与商业化受到关注。节目只说‘带动相关商品热销’，不写具体销量、成交额或未经核验的转化参数。"),
        ("就业讨论", "‘数字人先影响固定讲解和重复宣传’属于节目分析，不是统计结论。必须同时说明真人仍承担创意、判断、体验、运营和责任。"),
        ("宇树科技", "节目采用‘开始申购/进入申购阶段’口径。发布前复核交易所或公司最新公告，不把热榜标题当作唯一事实来源。"),
        ("荣耀 Robot Phone", "节目口径为‘将于 8 月 12 日发布’。若视频在发布会后上线，应改为已发布事实，并更新产品能力描述。"),
        ("苹果与千问", "当前只说‘接入千问的消息引发讨论’，不表述为苹果官方已经确认。若无官方确认，画面角标应注明‘媒体消息/市场讨论’。"),
        ("时长规则", "最终成片不超过 2:00。主新闻约占四成，其余三条各约两成；超时优先压缩空白停顿和转场，不擅自删改已确认观点。"),
    ])
    ui.add_heading(doc, "1.1 资料来源", 2)
    add_sources(doc)

    ui.add_heading(doc, "2. 推荐发布方案", 1)
    ui.add_meta_table(doc, [
        ("内容类型", "两分钟科技快讯 / 双数字人播客 / 热点解释"),
        ("平台与画幅", "抖音，9:16 竖屏；目标时长 1:45—2:00。"),
        ("人物分工", "雅雅负责事实、解释和判断；檬檬负责提问、生活化转译，并以数字人身份追问就业影响。"),
        ("核心受众", "关心 AI、机器人和手机科技，但不熟悉行业术语的普通观众。"),
        ("一句话卖点", "让两个数字人讨论数字人会不会抢真人饭碗，再用三条快讯把 AI 进入现实世界的路径串起来。"),
        ("推荐主标题", "AI角色开始带货了，真人演员还需要出镜吗？"),
        ("封面大字", "数字人要抢饭碗？"),
        ("封面副标题", "两分钟看懂四条科技热点"),
    ])

    ui.add_heading(doc, "3. 多风格视频标题库", 1)
    ui.add_body(doc, "首发建议使用标题 1。它把事件和讨论问题同时说清楚，观众不需要猜视频内容；标题 2 更有冲突感，但封面和开头必须避免制造‘真人马上失业’的误解。")
    add_title_table(doc)

    ui.add_heading(doc, "4. 视频详情与发布文案", 1)
    ui.add_heading(doc, "4.1 推荐视频详情", 2)
    ui.add_body(doc, "AI漫剧里的虚拟角色开始积累粉丝、参与宣传并带动商品热销；与此同时，机器人走向资本与产业，手机开始替用户执行任务，国产模型也在争夺系统入口。四条消息看似分散，其实都在说明：AI 正从‘能聊天’走向‘能参与工作’。")
    ui.add_heading(doc, "4.2 抖音发布区文案", 2)
    ui.add_callout(doc, "发布文案", "AI角色开始带货，宇树科技进入申购热度，Robot Phone 准备发布，苹果设备与千问的消息也引发讨论。\n\n这期雅雅和檬檬用两分钟聊清楚四条科技新闻：数字人会不会抢真人饭碗？机器人什么时候真正进工厂？手机里的 AI 又会替你做什么？", fill="F8FAFC", accent=ui.BLUE)
    ui.add_heading(doc, "4.3 置顶评论与话题标签", 2)
    ui.add_body(doc, "置顶评论：如果数字人可以替你完成一部分工作，你最愿意把哪件事交给它？又有哪些事情，你只信真人？")
    ui.add_body(doc, "推荐标签：#科技快报  #数字人  #AI演员  #人形机器人  #RobotPhone  #千问  #人工智能")

    ui.add_heading(doc, "5. 完整双人对话文案", 1)
    ui.add_callout(doc, "录制规则", "以下顺序即最终剪辑顺序。成片以数字人自带口播音频为主时钟；不要重复叠加同内容 TTS。目标 1:45—2:00，口播自然偏快，但不能牺牲专有名词和核心判断的清晰度。", fill="FFF9E6", accent=ui.GOLD)
    for index, (speaker, line) in enumerate(DIALOGUE, 1):
        ui.add_speaker_line(doc, f"T{index:03d} · {speaker}", line)

    ui.add_heading(doc, "6. 雅雅（台词纯净版）", 1)
    ui.add_clean_script(doc, "雅雅", YAYA_LINES, f"长视频兼容输入：只保留下面 {len(YAYA_LINES)} 段正文，不复制标题、姓名、编号、‘生成注意’或舞台说明。推荐生产仍使用第 5 节的逐轮次导出。语气清楚、自然、有判断但不武断。")

    ui.add_heading(doc, "7. 檬檬（台词纯净版）", 1)
    ui.add_clean_script(doc, "檬檬", MENGMENG_LINES, f"长视频兼容输入：只保留下面 {len(MENGMENG_LINES)} 段正文，不复制标题、姓名、编号、‘生成注意’或舞台说明。推荐生产仍使用第 5 节的逐轮次导出。语气像普通人即时追问，数字人自我调侃要自然，不要演得过满。")

    ui.add_heading(doc, "8. 数字人视频生成与合成执行单", 1)
    ui.add_heading(doc, "8.1 生成顺序", 2)
    for item in [
        "推荐按第 5 节逐轮次生成：一条 Txxx 台词对应一个视频，文件名使用 T001_YAYA.mp4、T002_MENGMENG.mp4 等稳定格式。",
        "把全部视频上传到导演工作台的‘数字人导入’页面；批量上传会从文件名中的 Txxx 自动匹配人物和顺序。",
        "先执行媒体检查，再用本地 ASR 核对每轮台词；缺句、错人、重复文件或覆盖率不足都必须返工对应轮次。",
        "每段数字人视频的原生音频作为唯一主轨，不再生成或叠加相同内容的外置 TTS。",
        "母版合成后，以实际原声重新生成时间线和字幕。实测超过 2:00 时，先缩短人物间空白、信息卡停留和转场；仍超时再回到脚本确认，不自行删新闻事实。",
    ]:
        ui.add_number(doc, item)
    ui.add_heading(doc, "8.2 建议生成参数", 2)
    ui.add_meta_table(doc, [
        ("语言", "普通话，自然聊天；段宴、容寄侨、宇树科技、Robot Phone、千问等专有名词要清楚。"),
        ("语速", "自然偏快，建议从 1.05 倍测试；以实际成片不超过 2:00 为准，不使用机械式加速。"),
        ("情绪", "雅雅：清楚、稳、像在解释；檬檬：好奇、反应快、带一点自然的数字人自我意识。"),
        ("停顿", "句内短停，新闻切换处留轻微节奏点；避免每句之间都留长空白。"),
        ("输出检查", "先听完整音频计时，再检查嘴型、说话人、字幕、日期角标和画面新闻来源。"),
        ("母版规格", "H.264 / AAC，25fps，48kHz；以数字人原声时间线为准，脚本时间只作生成前估算。"),
    ])

    ui.add_heading(doc, "8.3 轮次文件交付清单", 2)
    ui.add_body(doc, "下列文件名用于自动归位。每个视频必须同时包含画面与数字人原声；编号和人物后缀不属于口播内容。")
    table = doc.add_table(rows=1, cols=3)
    ui.set_table_geometry(table, [1500, 2200, 5860])
    ui.set_table_borders(table)
    for idx, text in enumerate(["轮次", "人物", "建议文件名"]):
        cell = table.cell(0, idx)
        cell.text = text
        ui.shade_cell(cell, ui.LIGHT_BLUE)
        for paragraph in cell.paragraphs:
            ui.set_para_spacing(paragraph, after=0, line=1.15)
            for run in paragraph.runs:
                ui.set_run_font(run, size=10, color=ui.INK, bold=True)
    ui.set_repeat_table_header(table.rows[0])
    speaker_ids = {"雅雅": "YAYA", "檬檬": "MENGMENG"}
    for index, (speaker, _line) in enumerate(DIALOGUE, 1):
        turn_id = f"T{index:03d}"
        cells = table.add_row().cells
        for col, value in enumerate([turn_id, speaker, f"{turn_id}_{speaker_ids[speaker]}.mp4"]):
            cells[col].text = value
            for paragraph in cells[col].paragraphs:
                ui.set_para_spacing(paragraph, after=0, line=1.1)
                for run in paragraph.runs:
                    ui.set_run_font(run, size=9.5, color=ui.INK, bold=(col == 0))

    ui.add_heading(doc, "9. 画面与节奏建议", 1)
    ui.add_body(doc, "画面不追求复杂，重点是让观众两分钟内听懂。主新闻多给人物反应和就业关键词，后三条用简洁信息卡快速推进。")
    table = doc.add_table(rows=1, cols=3)
    ui.set_table_geometry(table, [1500, 2200, 5660])
    ui.set_table_borders(table)
    for idx, text in enumerate(["时间段", "内容", "画面执行"]):
        cell = table.cell(0, idx)
        cell.text = text
        ui.shade_cell(cell, ui.LIGHT_BLUE)
        for paragraph in cell.paragraphs:
            ui.set_para_spacing(paragraph, after=0, line=1.15)
            for run in paragraph.runs:
                ui.set_run_font(run, size=10, color=ui.INK, bold=True)
    ui.set_repeat_table_header(table.rows[0])
    rows = [
        ("0:00—0:10", "栏目开场", "双人快速入镜；四条消息以四格关键词闪现，主标题立刻可读。"),
        ("0:10—0:52", "数字人带货与就业", "段宴、容寄侨角色卡 + ‘带货/粉丝/重复工作/真人价值’关键词；重点给檬檬的数字人身份反应。"),
        ("0:52—1:12", "宇树科技", "机器人从舞台到工厂的对照图；字幕只保留‘量产、工厂、价值’。"),
        ("1:12—1:32", "荣耀 Robot Phone", "手机调用应用的简化流程图；不要真的画一个机器人钻进手机。"),
        ("1:32—1:50", "苹果与千问", "系统、相册、语音助手三个入口图标；若未获官方确认，加‘媒体消息’角标。"),
        ("1:50—2:00", "总结与互动", "回到双人，屏幕只留问题：‘你最愿意把哪件事交给 AI？’。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if idx == 0:
                ui.shade_cell(cells[idx], "F8FAFC")
            for paragraph in cells[idx].paragraphs:
                ui.set_para_spacing(paragraph, after=0, line=1.18)
                for run in paragraph.runs:
                    ui.set_run_font(run, size=9.5, color=ui.INK, bold=(idx < 2))

    ui.add_heading(doc, "10. 发布前审查清单", 1)
    ui.add_body(doc, "以下任一项未通过，都应返工对应片段后再导出。")
    for item in [
        "时效审查：四条新闻的日期、状态和官方/媒体口径均在发布当天重新核验。",
        "销量审查：没有出现具体销量、成交额、转化率等未经确认的参数，只使用‘热销、热卖、受到关注’等描述。",
        "主次审查：第一条重点是数字人带货和就业影响，没有把篇幅浪费在漫剧剧情；后三条均保留发生了什么和为什么重要。",
        "就业审查：没有简单制造‘数字人让真人失业’恐慌，同时清楚解释重复工作、创意判断、真实体验和责任边界。",
        "时长审查：最终导出文件不超过 2:00；开头 10 秒内出现栏目、四条消息和主话题。",
        "音频审查：只有一条主口播轨，没有数字人原声与 TTS 重复；音量稳定、专有名词清楚。",
        "台词审查：顺序与第 5 节逐句一致，雅雅、檬檬未互换，没有漏句、重句和提前抢话。",
        "字幕审查：字幕依据最终音频生成，入出点准确，不压脸，不被抖音界面遮挡。",
        "互动审查：结尾明确出现‘你最愿意把哪件事交给 AI？’，并给观众留出短暂阅读时间。",
    ]:
        ui.add_bullet(doc, item)
    ui.add_callout(doc, "最终通过条件", "只听音频能听清四件事分别发生了什么、为什么重要；静音看画面仍能看出‘数字人就业—机器人产业—手机执行—系统入口’这条主线。两项都通过后再发布。", fill="EAF6F1", accent=ui.GREEN)

    ui.add_heading(doc, "附录：版本与来源记录", 1)
    ui.add_meta_table(doc, [
        ("文档版本", "v1.1（数字人逐轮次导入版）"),
        ("制作日期", CHECK_DATE),
        ("栏目", "今日科技快报"),
        ("主话题", "AI角色开始带货了，真人演员还需要出镜吗？"),
        ("使用提醒", "新闻内容时效性强；发布日与制作日不同时，必须更新产品状态、日期角标和事实口径。"),
    ])
    ui.add_heading(doc, "来源链接", 2)
    add_sources(doc)

    doc.core_properties.title = "第四期｜科技快报｜AI角色开始带货了，真人演员还需要出镜吗？"
    doc.core_properties.subject = "抖音发布文案、完整对话稿、双数字人纯净台词与制作审查规则"
    doc.core_properties.author = "OpenMontage 内容制作"
    doc.core_properties.comments = "资料核对时间：2026-08-10；发布前复核全部新闻时效。"
    doc.save(DOCX_PATH)


def build_markdown() -> None:
    lines = [
        "# 第四期｜科技快报",
        "",
        "## AI角色开始带货了，真人演员还需要出镜吗？",
        "",
        f"> 资料核对时间：{CHECK_DATE}。目标成片不超过 2 分钟。",
        "",
        "## 完整双人对话文案",
        "",
    ]
    for index, (speaker, text) in enumerate(DIALOGUE, 1):
        lines.extend([f"**T{index:03d} · {speaker}：**{text}", ""])
    lines.extend([
        "## 数字人逐轮次导入",
        "",
        "推荐一条 Txxx 台词对应一个数字人视频，按 `T001_YAYA.mp4`、`T002_MENGMENG.mp4` 的格式命名后上传到导演工作台。依次执行媒体检查、ASR 核词和原声母版合成；数字人视频自带声音是唯一主音轨，不再叠加同内容 TTS。",
        "",
        "按人物合并的纯净台词仅用于必须整段生成的平台兼容。最终字幕、轮次切点和两分钟检查都以实际数字人原声为准。",
        "",
    ])
    lines.extend(["## 雅雅（台词纯净版）", ""])
    for text in YAYA_LINES:
        lines.extend([text, ""])
    lines.extend(["## 檬檬（台词纯净版）", ""])
    for text in MENGMENG_LINES:
        lines.extend([text, ""])
    lines.extend([
        "## 发布文案",
        "",
        "AI角色开始带货，宇树科技进入申购热度，Robot Phone 准备发布，苹果设备与千问的消息也引发讨论。这期雅雅和檬檬用两分钟聊清楚四条科技新闻。",
        "",
        "置顶评论：如果数字人可以替你完成一部分工作，你最愿意把哪件事交给它？又有哪些事情，你只信真人？",
        "",
        "标签：#科技快报 #数字人 #AI演员 #人形机器人 #RobotPhone #千问 #人工智能",
    ])
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    build_docx()
    build_markdown()
    print(DOCX_PATH)
    print(MD_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
