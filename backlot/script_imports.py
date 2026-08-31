"""Deterministic user-script intake for avatar projects.

This module deliberately does not call an LLM.  It turns DOCX paragraphs or
pasted dialogue into a reviewable staging record, then builds the exact script
contract the avatar pipeline expects after explicit human confirmation.
"""

from __future__ import annotations

import hashlib
import io
import json
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document


MAX_SCRIPT_IMPORT_BYTES = 10 * 1024 * 1024
STAGING_DIRECTORY = Path("artifacts/script_import_staging")
TOKEN_RE = re.compile(r"^SIP-[A-F0-9]{24}$")
TURN_RE = re.compile(r"^T(?P<number>\d{1,6})$", re.IGNORECASE)
SPEAKER_ID_RE = re.compile(r"^[a-z][a-z0-9_-]{1,31}$")
EXPLICIT_DIALOGUE_RE = re.compile(
    r"^\s*(?P<turn>T\d{1,6})\s*(?:[·•.\-—–]\s*)?"
    r"(?P<speaker>[^：:\r\n]{1,48})[：:]\s*(?P<text>.+?)\s*$",
    re.IGNORECASE,
)
IMPLICIT_DIALOGUE_RE = re.compile(
    r"^\s*(?P<speaker>[^：:\r\n]{1,48})[：:]\s*(?P<text>.+?)\s*$"
)
KNOWN_SPEAKER_IDS = {
    "雅雅": "yaya",
    "檬檬": "mengmeng",
    "萌萌": "mengmeng",
    "旁白": "narrator",
    "主持人": "host",
}
NON_DIALOGUE_LABELS = {
    "录制规则", "使用说明", "导入说明", "注意事项", "提示", "规则",
    "项目名称", "标题", "摘要", "内容目标", "纯净台词版",
}


class ScriptImportError(ValueError):
    """The uploaded document cannot safely become a production script."""


def _now() -> str:
    return datetime.now(UTC).isoformat()


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _atomic_json(path: Path, value: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(path)


def _normalise_turn_id(value: str | None, index: int) -> str:
    if not value:
        return f"T{index:03d}"
    match = TURN_RE.fullmatch(value.strip())
    if not match:
        raise ScriptImportError(f"轮次编号格式不正确：{value}")
    return f"T{int(match.group('number')):03d}"


def _speaker_id(name: str, order: int) -> str:
    cleaned = re.sub(r"\s+", "", name).strip("*_·")
    known = KNOWN_SPEAKER_IDS.get(cleaned)
    if known:
        return known
    ascii_value = re.sub(r"[^a-z0-9_-]+", "-", cleaned.lower()).strip("-")
    if SPEAKER_ID_RE.fullmatch(ascii_value or ""):
        return ascii_value
    return f"speaker-{order}"


def _estimate_seconds(text: str) -> float:
    meaningful = len(re.sub(r"\s+", "", text))
    return round(max(1.5, meaningful / 4.2), 2)


def _clean_rows(rows: list[tuple[int, str]]) -> list[tuple[int, str]]:
    return [(line, re.sub(r"\s+", " ", text).strip()) for line, text in rows if text and text.strip()]


def _parse_rows(rows: list[tuple[int, str]], *, title: str, source_kind: str) -> dict[str, Any]:
    rows = _clean_rows(rows)
    explicit: list[dict[str, Any]] = []
    warnings: list[str] = []
    previous_number = 0
    for line_number, line in rows:
        match = EXPLICIT_DIALOGUE_RE.match(line)
        if not match:
            continue
        current_number = int(TURN_RE.fullmatch(match.group("turn")).group("number"))
        if explicit and current_number <= previous_number:
            warnings.append(
                f"在第 {line_number} 段发现轮次重新从 {match.group('turn').upper()} 开始；"
                "系统只导入前一组完整编号台词，避免把纯净台词版重复导入。"
            )
            break
        explicit.append({
            "raw_turn_id": match.group("turn"),
            "speaker_name": match.group("speaker").strip(" *_·"),
            "text": match.group("text").strip(),
            "source_location": f"第 {line_number} 段",
        })
        previous_number = current_number

    candidates = explicit
    if not candidates:
        implicit: list[dict[str, Any]] = []
        for line_number, line in rows:
            match = IMPLICIT_DIALOGUE_RE.match(line)
            if not match:
                continue
            speaker = match.group("speaker").strip(" *_·")
            if speaker in NON_DIALOGUE_LABELS or len(speaker) > 24:
                continue
            implicit.append({
                "raw_turn_id": None,
                "speaker_name": speaker,
                "text": match.group("text").strip(),
                "source_location": f"第 {line_number} 段",
            })
        candidates = implicit
        if candidates:
            warnings.append("原稿没有显式轮次编号，系统已按台词出现顺序补齐 T001…；请在提交前核对。")

    if not candidates:
        raise ScriptImportError("没有识别到“角色：台词”格式。请保留角色名和冒号后重新上传或粘贴。")

    turns: list[dict[str, Any]] = []
    seen_turns: set[str] = set()
    speaker_order: dict[str, int] = {}
    speaker_ids: dict[str, str] = {}
    for index, raw in enumerate(candidates, 1):
        turn_id = _normalise_turn_id(raw.get("raw_turn_id"), index)
        if turn_id in seen_turns:
            raise ScriptImportError(f"轮次编号重复：{turn_id}")
        seen_turns.add(turn_id)
        name = str(raw["speaker_name"]).strip()
        text = str(raw["text"]).strip()
        if not name or name in NON_DIALOGUE_LABELS:
            raise ScriptImportError(f"{turn_id} 的说话人名称无效")
        if not text:
            raise ScriptImportError(f"{turn_id} 的台词为空")
        if name not in speaker_order:
            speaker_order[name] = len(speaker_order) + 1
            speaker_ids[name] = _speaker_id(name, speaker_order[name])
        seconds = _estimate_seconds(text)
        if seconds > 20:
            warnings.append(f"{turn_id} 预计约 {seconds:.1f} 秒，建议后续拆分为更短的驱动音频。")
        turns.append({
            "turn_id": turn_id,
            "speaker_name": name,
            "speaker_id": speaker_ids[name],
            "text": text,
            "estimated_seconds": seconds,
            "source_location": raw["source_location"],
        })

    start = 0.0
    for index, turn in enumerate(turns, 1):
        turn["id"] = f"section-{index:03d}"
        turn["start_seconds"] = round(start, 2)
        start += float(turn["estimated_seconds"])
        turn["end_seconds"] = round(start, 2)

    return {
        "title": title.strip() or "未命名数字人口播脚本",
        "source_kind": source_kind,
        "turns": turns,
        "turn_count": len(turns),
        "speakers": [
            {"speaker_id": speaker_ids[name], "name": name}
            for name in speaker_order
        ],
        "estimated_total_duration_seconds": round(start, 2),
        "warnings": list(dict.fromkeys(warnings)),
    }


def parse_text_script(text: str, *, title: str = "") -> dict[str, Any]:
    raw = str(text or "")
    if not raw.strip():
        raise ScriptImportError("粘贴的脚本为空")
    if len(raw.encode("utf-8")) > MAX_SCRIPT_IMPORT_BYTES:
        raise ScriptImportError("脚本文本不能超过 10 MB")
    rows = list(enumerate(raw.splitlines(), 1))
    inferred_title = title.strip() or next((line.strip() for _, line in rows if line.strip() and not EXPLICIT_DIALOGUE_RE.match(line)), "")
    return _parse_rows(rows, title=inferred_title, source_kind="pasted_text")


def parse_docx_script(data: bytes, *, filename: str) -> dict[str, Any]:
    if not data:
        raise ScriptImportError("上传的 Word 文件为空")
    if len(data) > MAX_SCRIPT_IMPORT_BYTES:
        raise ScriptImportError("Word 脚本不能超过 10 MB")
    if Path(filename).suffix.lower() != ".docx":
        raise ScriptImportError("当前只支持 .docx 格式的 Word 脚本")
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ScriptImportError("无法读取这个 Word 文件；请确认文件未损坏且格式为 .docx") from exc
    if document.tables:
        raise ScriptImportError("这个 Word 脚本包含表格；V1 暂只支持按段落排列的“角色：台词”脚本")
    rows = [(index, paragraph.text) for index, paragraph in enumerate(document.paragraphs, 1)]
    title = next((text.strip() for _, text in rows if text and text.strip() and not EXPLICIT_DIALOGUE_RE.match(text)), Path(filename).stem)
    return _parse_rows(rows, title=title or Path(filename).stem, source_kind="docx")


def _stage(project_dir: Path, preview: dict[str, Any], *, source_bytes: bytes, filename: str) -> dict[str, Any]:
    token = f"SIP-{uuid4().hex[:24].upper()}"
    digest = _sha256_bytes(source_bytes)
    suffix = ".docx" if preview["source_kind"] == "docx" else ".txt"
    directory = project_dir / STAGING_DIRECTORY
    directory.mkdir(parents=True, exist_ok=True)
    source_path = directory / f"{token}{suffix}"
    source_path.write_bytes(source_bytes)
    record = {
        "version": "1.0",
        "import_token": token,
        "project_id": project_dir.name,
        "created_at": _now(),
        "filename": Path(filename).name,
        "source_sha256": digest,
        "source_path": source_path.relative_to(project_dir).as_posix(),
        "preview": {**preview, "source_sha256": digest, "filename": Path(filename).name},
    }
    _atomic_json(directory / f"{token}.json", record)
    return {"import_token": token, **record["preview"]}


def stage_docx_preview(project_dir: Path, data: bytes, *, filename: str) -> dict[str, Any]:
    return _stage(project_dir, parse_docx_script(data, filename=filename), source_bytes=data, filename=filename)


def stage_text_preview(project_dir: Path, text: str, *, title: str = "") -> dict[str, Any]:
    source = str(text or "").encode("utf-8")
    return _stage(project_dir, parse_text_script(text, title=title), source_bytes=source, filename="pasted-script.txt")


def load_staged_import(project_dir: Path, token: str) -> dict[str, Any]:
    value = str(token or "").strip().upper()
    if not TOKEN_RE.fullmatch(value):
        raise ScriptImportError("脚本预览令牌无效，请重新解析脚本")
    record_path = project_dir / STAGING_DIRECTORY / f"{value}.json"
    if not record_path.is_file():
        raise ScriptImportError("脚本预览已失效，请重新上传或粘贴脚本")
    try:
        record = json.loads(record_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ScriptImportError("脚本预览记录损坏，请重新解析脚本") from exc
    if record.get("project_id") != project_dir.name:
        raise ScriptImportError("脚本预览不属于当前项目")
    source_path = project_dir / str(record.get("source_path") or "")
    if not source_path.is_file() or _sha256_bytes(source_path.read_bytes()) != record.get("source_sha256"):
        raise ScriptImportError("脚本来源校验失败，请重新上传或粘贴脚本")
    return record


def build_script_from_staged_import(record: dict, *, speaker_overrides: dict[str, str] | None = None) -> tuple[dict, dict, bytes]:
    preview = record.get("preview") if isinstance(record.get("preview"), dict) else {}
    turns = preview.get("turns") if isinstance(preview.get("turns"), list) else []
    if not turns:
        raise ScriptImportError("脚本预览中没有可提交的轮次")
    overrides = speaker_overrides or {}
    if not isinstance(overrides, dict):
        raise ScriptImportError("说话人编号映射格式无效")
    names_to_ids: dict[str, str] = {}
    ids_to_names: dict[str, str] = {}
    for turn in turns:
        name = str(turn.get("speaker_name") or "").strip()
        speaker_id = str(overrides.get(name) or turn.get("speaker_id") or "").strip().lower()
        if not SPEAKER_ID_RE.fullmatch(speaker_id):
            raise ScriptImportError(f"“{name}”的说话人编号不合法：{speaker_id or '空值'}")
        previous = names_to_ids.get(name)
        if previous and previous != speaker_id:
            raise ScriptImportError(f"“{name}”不能同时映射到多个说话人编号")
        other = ids_to_names.get(speaker_id)
        if other and other != name:
            raise ScriptImportError(f"说话人编号“{speaker_id}”不能同时绑定“{other}”和“{name}”")
        names_to_ids[name] = speaker_id
        ids_to_names[speaker_id] = name

    sections: list[dict[str, Any]] = []
    for turn in turns:
        name = str(turn["speaker_name"])
        speaker_id = names_to_ids[name]
        sections.append({
            "id": str(turn["id"]),
            "turn_id": str(turn["turn_id"]),
            "speaker_id": speaker_id,
            "speaker_name": name,
            "expected_asset_filename": f"{turn['turn_id']}_{speaker_id.upper()}.mp4",
            "label": f"{turn['turn_id']} · {name}",
            "text": str(turn["text"]),
            "start_seconds": float(turn["start_seconds"]),
            "end_seconds": float(turn["end_seconds"]),
            "visual_contract": {
                "visual_intent": f"数字人口播：{name} 讲述本段内容",
                "required_assets": [],
                "forbidden_states": ["拉伸或变速驱动音频"],
                "min_visual_coverage": 1,
            },
        })
    script = {
        "version": "1.0",
        "title": str(preview.get("title") or "未命名数字人口播脚本"),
        "total_duration_seconds": float(preview.get("estimated_total_duration_seconds") or 1),
        "sections": sections,
        "metadata": {
            "source": str(preview.get("source_kind") or "user_script"),
            "source_filename": str(record.get("filename") or ""),
            "source_sha256": str(record.get("source_sha256") or ""),
            "timing_basis": "script_estimate_pending_native_avatar_audio",
            "text_policy": "verbatim_no_ai_rewrite",
        },
    }
    provenance = {
        "source_kind": str(preview.get("source_kind") or "user_script"),
        "filename": str(record.get("filename") or ""),
        "title": script["title"],
        "source_sha256": str(record.get("source_sha256") or ""),
        "warnings": list(preview.get("warnings") or []),
        "import_token": str(record.get("import_token") or ""),
    }
    source_path = Path(str(record.get("source_path") or ""))
    return script, provenance, source_path


def consume_staged_import(project_dir: Path, record: dict) -> None:
    token = str(record.get("import_token") or "")
    source_path = project_dir / str(record.get("source_path") or "")
    record_path = project_dir / STAGING_DIRECTORY / f"{token}.json"
    for path in (record_path, source_path):
        try:
            path.unlink()
        except OSError:
            pass
