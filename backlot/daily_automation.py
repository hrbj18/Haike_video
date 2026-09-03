"""Durable daily technology-news production contracts.

The module deliberately keeps the paid media stages behind explicit, persisted
gates.  Research and script validation can be exercised without Voicebox or
RunningHub, while a resumed process can continue from the last successful
stage instead of submitting duplicate paid jobs.
"""

from __future__ import annotations

import hashlib
import html
import json
import os
import re
import subprocess
import shutil
import tempfile
import threading
import time
import uuid
import xml.etree.ElementTree as ET
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import UTC, date, datetime, time as datetime_time, timedelta, timezone
from email.utils import parsedate_to_datetime
from pathlib import Path
from typing import Any, Callable, Iterable
from urllib.parse import urlsplit, urljoin, urlunsplit

import requests
from docx import Document

from backlot.ai_text import (
    TextAIError,
    _chat_json,
    _effective_value,
    _read_env_file,
    _secrets_path,
    read_text_ai_config,
    read_text_provider_status,
)
from backlot.copy_skill_hotspot_feed import (
    DEFAULT_COPY_SKILL_HOTSPOT_ROOT,
    feed_to_discovery_candidates,
    feed_to_heat_signals,
    try_load_copy_skill_hotspot_feed,
)
from backlot.state import PROJECTS_DIR, REPO_ROOT
from lib.checkpoint import init_project
from lib.env_loader import load_env


load_env(REPO_ROOT)


# China Standard Time has had no daylight-saving transitions since 1991.
# A fixed UTC+8 timezone keeps the scheduler portable on Windows machines
# where the optional IANA ``tzdata`` package is not installed.
LOCAL_TIMEZONE = timezone(timedelta(hours=8), "Asia/Shanghai")
CONFIG_PATH = REPO_ROOT / "config" / "daily_tech_brief.json"
RUNS_ROOT = REPO_ROOT / ".backlot" / "daily-runs"
RUN_LOCK_PATH = RUNS_ROOT / ".daily-production.lock"
DEFAULT_MAX_BUDGET_CNY = 5.0
MAX_TOPIC_SELECTION_CANDIDATES = 28
INTRO_TEXT = "欢迎收听今天的科技快报。一起看看过去二十四小时，科技圈发生了什么。"
STAGE_ORDER = (
    "research",
    "script",
    "project",
    "voice",
    "avatar",
    "align",
    "visuals",
    "compose",
    "review_ready",
)
TERMINAL_RUN_STATES = {"review_ready", "failed", "cancelled"}
ACTIVE_RUN_STATES = {"queued", "running"}
MEDIA_RELEASE_RULE_VERSION = "daily-media-release-v2-review-candidate"
OOM_PATTERNS = (
    r"\bout of memory\b",
    r"outofmemory(?:error)?",
    r"\bcuda oom\b",
    r"insufficient.*(?:gpu|vram|video memory)",
    r"failed to allocate.*(?:cuda|gpu|vram)",
    r"显存不足",
    r"爆显存",
    r"内存分配失败.*(?:gpu|显存)",
)
NON_OOM_PATTERNS = (
    r"timeout",
    r"timed out",
    r"queue",
    r"rate limit",
    r"too many requests",
    r"network",
    r"connection",
    r"排队",
    r"超时",
    r"限流",
    r"网络",
)
DEFAULT_NEWS_SOURCES = (
    {
        "id": "ithome-direct",
        "name": "IT之家",
        "url": "https://www.ithome.com/rss/",
        "authority": "media",
        "max_candidates": 40,
    },
    {
        "id": "qbitai-direct",
        "name": "量子位",
        "url": "https://www.qbitai.com/feed",
        "authority": "media",
        "max_candidates": 30,
    },
    {
        "id": "ifanr-direct",
        "name": "爱范儿",
        "url": "https://www.ifanr.com/feed",
        "authority": "media",
        "max_candidates": 30,
    },
    {
        "id": "google-deepmind",
        "name": "Google DeepMind 官方博客",
        "url": "https://deepmind.google/blog/rss.xml",
        "authority": "official",
    },
    {
        "id": "openai-news",
        "name": "OpenAI 官方新闻",
        "url": "https://openai.com/news/rss.xml",
        "authority": "official",
    },
    {
        "id": "nvidia-blog",
        "name": "NVIDIA 官方博客",
        "url": "https://blogs.nvidia.com/feed/",
        "authority": "official",
    },
    {
        "id": "nvidia-developer",
        "name": "NVIDIA 开发者博客",
        "url": "https://developer.nvidia.com/blog/feed/",
        "authority": "official",
    },
    {
        "id": "microsoft-blog",
        "name": "Microsoft 官方博客",
        "url": "https://blogs.microsoft.com/feed/",
        "authority": "official",
    },
    {
        "id": "techcrunch-ai",
        "name": "TechCrunch AI",
        "url": "https://techcrunch.com/category/artificial-intelligence/feed/",
        "authority": "media",
    },
    {
        "id": "the-verge",
        "name": "The Verge",
        "url": "https://www.theverge.com/rss/index.xml",
        "authority": "media",
    },
    {
        "id": "ars-technica-technology",
        "name": "Ars Technica Technology Lab",
        "url": "https://feeds.arstechnica.com/arstechnica/technology-lab",
        "authority": "media",
    },
    {
        "id": "venturebeat-ai",
        "name": "VentureBeat AI",
        "url": "https://venturebeat.com/category/ai/feed/",
        "authority": "media",
    },
    {
        "id": "google-news-cn-big-tech",
        "name": "中国大厂科技新闻聚合",
        "url": "https://news.google.com/rss/search?q=%28%E8%85%BE%E8%AE%AF+OR+%E5%AD%97%E8%8A%82%E8%B7%B3%E5%8A%A8+OR+%E9%98%BF%E9%87%8C%E5%B7%B4%E5%B7%B4+OR+%E7%99%BE%E5%BA%A6+OR+%E5%8D%8E%E4%B8%BA+OR+DeepSeek%29+%28AI+OR+%E5%A4%A7%E6%A8%A1%E5%9E%8B+OR+%E8%8A%AF%E7%89%87+OR+%E6%9C%BA%E5%99%A8%E4%BA%BA%29+after:{previous_date}+before:{next_date}&hl=zh-CN&gl=CN&ceid=CN:zh-Hans",
        "authority": "media",
        "max_candidates": 40,
    },
    {
        "id": "google-news-cn-tech-media",
        "name": "中文科技媒体新闻聚合",
        "url": "https://news.google.com/rss/search?q=%28site%3A36kr.com+OR+site%3Aithome.com+OR+site%3Ageekpark.net+OR+site%3Aleiphone.com%29+%28AI+OR+%E5%A4%A7%E6%A8%A1%E5%9E%8B+OR+%E8%8A%AF%E7%89%87+OR+%E6%9C%BA%E5%99%A8%E4%BA%BA%29+after:{previous_date}+before:{next_date}&hl=zh-CN&gl=CN&ceid=CN:zh-Hans",
        "authority": "media",
        "max_candidates": 40,
    },
    {
        "id": "google-news-cn-consumer-tech",
        "name": "中文消费科技新闻聚合",
        "url": "https://news.google.com/rss/search?q=%28site%3Aithome.com%20OR%20site%3A36kr.com%20OR%20site%3Ageekpark.net%20OR%20site%3Aleiphone.com%29%20%28%E5%BE%AE%E4%BF%A1%20OR%20%E6%89%8B%E6%9C%BA%20OR%20%E6%95%B0%E7%A0%81%20OR%20%E6%B8%B8%E6%88%8F%20OR%20%E6%B5%8F%E8%A7%88%E5%99%A8%20OR%20%E8%80%B3%E6%9C%BA%20OR%20%E9%BC%A0%E6%A0%87%20OR%20%E9%99%8D%E4%BB%B7%20OR%20%E6%B6%A8%E4%BB%B7%29%20after%3A{previous_date}%20before%3A{next_date}&hl=zh-CN&gl=CN&ceid=CN:zh-Hans",
        "authority": "media",
        "max_candidates": 40,
    },
)

DEFAULT_HEAT_SOURCES = (
    {
        "id": "baidu-realtime",
        "name": "百度实时热榜",
        "kind": "baidu_board",
        "url": "https://top.baidu.com/board?tab=realtime",
        "max_items": 50,
    },
)

# 抖音热榜与对标账号是"选题信号"，不是事实证据：标题只用于热度排序补盲，
# 事实仍由原站正文证据门冻结。配置 API 时优先请求，失败回落快照；两者
# 都不可用时明确记录 skipped，不改变既有新闻生产行为。
DEFAULT_DOUYIN_SOURCES = (
    {
        "id": "douyin-hotboard",
        "name": "抖音热榜",
        "kind": "douyin_board",
        "url": "https://www.douyin.com/hot",
        "api_url": "",
        "max_items": 50,
        "snapshot_path": "config/douyin_snapshots/hotboard.json",
    },
    {
        "id": "douyin-benchmark-accounts",
        "name": "对标账号最新视频",
        "kind": "douyin_creator",
        "url": "",
        "api_url": "",
        "max_items": 30,
        "snapshot_path": "config/douyin_snapshots/benchmark_accounts.json",
    },
)

# These names are not a substitute for evidence.  They are an audience
# relevance hint used only to help the writing model choose between already
# verified candidates.  The daily script may still select a lesser-known item
# when it has a concrete Chinese user impact and explains it plainly.
CHINA_SHORT_VIDEO_PRIORITY_BRANDS = (
    "字节", "豆包", "火山", "阿里", "通义", "夸克", "腾讯", "混元", "元宝", "百度", "文心",
    "华为", "小米", "荣耀", "deepseek", "openai", "chatgpt", "google", "gemini", "苹果",
    "apple", "微软", "microsoft", "英伟达", "nvidia", "亚马逊", "amazon", "特斯拉", "tesla",
)


class DailyAutomationError(RuntimeError):
    """A safe, user-facing automation error."""


class DailyScriptValidationError(DailyAutomationError):
    def __init__(self, issues: list[str], *, candidate: dict[str, Any] | None = None):
        self.issues = issues
        self.candidate = candidate
        super().__init__("；".join(issues))


class DailyTopicValidationError(DailyAutomationError):
    def __init__(self, issues: list[str]):
        self.issues = issues
        super().__init__("；".join(issues))


def _now() -> str:
    return datetime.now(UTC).isoformat(timespec="seconds")


def _atomic_json(path: Path, value: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    handle, temporary = tempfile.mkstemp(prefix=f".{path.name}.", suffix=".tmp", dir=path.parent)
    try:
        with os.fdopen(handle, "w", encoding="utf-8", newline="\n") as stream:
            json.dump(value, stream, ensure_ascii=False, indent=2)
            stream.write("\n")
        os.replace(temporary, path)
    except Exception:
        try:
            Path(temporary).unlink()
        except OSError:
            pass
        raise


def _read_json(path: Path) -> dict[str, Any] | None:
    if not path.is_file():
        return None
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise DailyAutomationError(f"自动化状态文件损坏：{path.name}") from exc
    return value if isinstance(value, dict) else None


def default_config() -> dict[str, Any]:
    return {
        "version": "1.0",
        "enabled": False,
        "schedule_time": "03:00",
        "timezone": "Asia/Shanghai",
        "window_mode": "previous_calendar_day",
        "aspect": "portrait",
        "max_budget_cny": DEFAULT_MAX_BUDGET_CNY,
        "story_count": {"minimum": 3, "maximum": 5},
        "text_resilience": {
            "max_episode_combinations": 3,
            "max_editorial_reviews": 2,
            "max_text_attempts": 3,
            "max_rescue_research_rounds": 1,
            "max_wall_clock_seconds": 720,
        },
        "media_release": {
            "rule_version": MEDIA_RELEASE_RULE_VERSION,
            "premium_total": 85,
            "premium_hook": 16,
            "premium_dialogue": 16,
            "premium_information_density": 20,
        },
        "runninghub": {
            "primary_instance": "default",
            "standard_instance_type": "default",
            "allow_plus": False,
            "max_concurrency": 1,
        },
        "avatar": {
            "default_treatment": "pip_top_left",
            "shape": "circle",
            "source_aspect": "4:5",
            "preserve_source_aspect": True,
            "source_directory": "",
        },
        "visuals": {"static_images_enabled": False, "mode": "ai_balanced_motion"},
        "background_music": {"enabled": True, "use_first_news_track": True},
        "news_sources": list(DEFAULT_NEWS_SOURCES),
        "heat_sources": list(DEFAULT_HEAT_SOURCES),
        "douyin_sources": list(DEFAULT_DOUYIN_SOURCES),
        "copy_skill_hotspot_feed": {
            "enabled": False,
            "root": str(DEFAULT_COPY_SKILL_HOTSPOT_ROOT),
        },
        "last_updated_at": None,
    }


def _merge_dict(base: dict[str, Any], changes: dict[str, Any]) -> dict[str, Any]:
    result = dict(base)
    for key, value in changes.items():
        if isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = _merge_dict(result[key], value)
        else:
            result[key] = value
    return result


def _validate_config(config: dict[str, Any]) -> None:
    if not re.fullmatch(r"(?:[01]\d|2[0-3]):[0-5]\d", str(config.get("schedule_time") or "")):
        raise DailyAutomationError("每日运行时间必须使用 24 小时制 HH:MM")
    if str(config.get("timezone") or "") != "Asia/Shanghai":
        raise DailyAutomationError("V1 只支持北京时间 Asia/Shanghai")
    if str(config.get("aspect") or "") != "portrait":
        raise DailyAutomationError("V1 只支持竖版 9:16")
    try:
        budget = float(config.get("max_budget_cny"))
    except (TypeError, ValueError) as exc:
        raise DailyAutomationError("每日预算必须是有效数字") from exc
    if budget <= 0 or budget > DEFAULT_MAX_BUDGET_CNY:
        raise DailyAutomationError("无人值守 RunningHub 每日预算必须在 0—5 元之间")
    runninghub = config.get("runninghub") if isinstance(config.get("runninghub"), dict) else {}
    primary_instance = str(runninghub.get("primary_instance") or "").strip().lower()
    if primary_instance not in {"lite", "default"}:
        raise DailyAutomationError("RunningHub 默认实例只能是企业 Lite 或 Standard 24GB（default）")
    if primary_instance == "lite" and runninghub.get("lite_request_mode") != "verified_lite_only":
        raise DailyAutomationError("企业 Lite 必须先通过实际账单验证，不能把自动调度当作 Lite 保证")
    if primary_instance == "default":
        if str(runninghub.get("standard_instance_type") or "") != "default":
            raise DailyAutomationError("Standard 24GB 的 RunningHub instanceType 必须为 default")
        if runninghub.get("allow_plus") is not False:
            raise DailyAutomationError("每日自动生产禁止使用 RunningHub Plus 48GB")
    if int(runninghub.get("max_concurrency") or 0) not in {1, 2}:
        raise DailyAutomationError("RunningHub 并发数只能是 1 或 2")
    resilience = config.get("text_resilience") if isinstance(config.get("text_resilience"), dict) else {}
    resilience_bounds = {
        "max_episode_combinations": (1, 3),
        "max_editorial_reviews": (1, 2),
        "max_text_attempts": (1, 4),
        "max_rescue_research_rounds": (0, 1),
        "max_wall_clock_seconds": (60, 1800),
    }
    for key, (minimum, maximum) in resilience_bounds.items():
        try:
            value = int(resilience.get(key))
        except (TypeError, ValueError) as exc:
            raise DailyAutomationError(f"文本韧性配置 {key} 必须是整数") from exc
        if not minimum <= value <= maximum:
            raise DailyAutomationError(f"文本韧性配置 {key} 必须在 {minimum}—{maximum} 之间")
    avatar = config.get("avatar") if isinstance(config.get("avatar"), dict) else {}
    if str(avatar.get("shape") or "") not in {"rectangle", "rounded", "circle"}:
        raise DailyAutomationError("数字人外框只能是圆角矩形、圆形或直角矩形")
    for source in config.get("douyin_sources") or []:
        if not isinstance(source, dict) or str(source.get("kind") or "") not in {"douyin_board", "douyin_creator"}:
            raise DailyAutomationError("抖音信号来源类型必须是 douyin_board 或 douyin_creator")
        api_url = str(source.get("api_url") or "").strip()
        if api_url:
            parsed = urlsplit(api_url)
            if parsed.scheme.lower() != "https" or not parsed.netloc:
                raise DailyAutomationError("抖音数据接口必须使用有效的 HTTPS 地址")
            if re.search(r"(?:key|token|secret)=", parsed.query, re.IGNORECASE):
                raise DailyAutomationError("抖音数据接口密钥不得写入 URL 查询参数")
    copy_skill_feed = config.get("copy_skill_hotspot_feed")
    if copy_skill_feed is not None:
        if not isinstance(copy_skill_feed, dict):
            raise DailyAutomationError("copy_skill 热点候选池配置必须是对象")
        if copy_skill_feed.get("enabled") is not False and not str(copy_skill_feed.get("root") or "").strip():
            raise DailyAutomationError("copy_skill 热点候选池启用时必须配置输出根目录")


def read_config() -> dict[str, Any]:
    persisted = _read_json(CONFIG_PATH) or {}
    value = _merge_dict(default_config(), persisted)
    _validate_config(value)
    return value


def save_config(payload: dict[str, Any]) -> dict[str, Any]:
    current = read_config()
    value = _build_config(current, payload)
    _atomic_json(CONFIG_PATH, value)
    return value


def _build_config(current: dict[str, Any], payload: dict[str, Any]) -> dict[str, Any]:
    allowed = {
        "enabled",
        "schedule_time",
        "max_budget_cny",
        "story_count",
        "avatar",
        "visuals",
        "background_music",
        "news_sources",
        "heat_sources",
        "douyin_sources",
        "copy_skill_hotspot_feed",
    }
    changes = {key: value for key, value in payload.items() if key in allowed}
    value = _merge_dict(current, changes)
    value["last_updated_at"] = _now()
    _validate_config(value)
    return value


def apply_config_with_scheduler(payload: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    """Apply user settings and the Windows task as one observable transaction.

    The old implementation persisted ``enabled=true`` before asking Windows to
    create the task.  A Task Scheduler error therefore left the project card
    green even though nothing would run at 03:00.  Build the candidate first,
    change the fixed scheduler task, and only then publish the configuration.
    If persisting fails, make a best-effort rollback to the previous task.
    """

    previous = read_config()
    candidate = _build_config(previous, payload)
    scheduler = sync_windows_scheduler(candidate)
    try:
        _atomic_json(CONFIG_PATH, candidate)
    except Exception as exc:
        try:
            sync_windows_scheduler(previous)
        except Exception:
            pass
        raise DailyAutomationError("自动化配置写入失败，计划任务已回滚") from exc
    return candidate, scheduler


def previous_target_date(now: datetime | None = None) -> date:
    current = now or datetime.now(LOCAL_TIMEZONE)
    if current.tzinfo is None:
        current = current.replace(tzinfo=LOCAL_TIMEZONE)
    return current.astimezone(LOCAL_TIMEZONE).date() - timedelta(days=1)


def target_window(target: date) -> tuple[datetime, datetime]:
    start = datetime.combine(target, datetime_time.min, LOCAL_TIMEZONE)
    return start, start + timedelta(days=1)


def next_scheduled_run(config: dict[str, Any] | None = None, now: datetime | None = None) -> dict[str, str]:
    settings = config or read_config()
    local_now = (now or datetime.now(LOCAL_TIMEZONE)).astimezone(LOCAL_TIMEZONE)
    hour, minute = (int(item) for item in str(settings["schedule_time"]).split(":", 1))
    scheduled = datetime.combine(local_now.date(), datetime_time(hour, minute), LOCAL_TIMEZONE)
    if scheduled <= local_now:
        scheduled += timedelta(days=1)
    return {
        "starts_at": scheduled.isoformat(timespec="minutes"),
        "target_date": (scheduled.date() - timedelta(days=1)).isoformat(),
    }


def _run_path(target: date | str) -> Path:
    value = date.fromisoformat(target) if isinstance(target, str) else target
    return RUNS_ROOT / value.isoformat() / "daily_run.json"


def _new_stage(name: str) -> dict[str, Any]:
    return {
        "name": name,
        "status": "pending",
        "attempts": 0,
        "started_at": None,
        "finished_at": None,
        "message": None,
        "error": None,
        "output": {},
    }


def create_or_resume_run(target: date | str, *, trigger: str = "manual") -> dict[str, Any]:
    target_value = date.fromisoformat(target) if isinstance(target, str) else target
    path = _run_path(target_value)
    existing = _read_json(path)
    if existing:
        return existing
    start, end = target_window(target_value)
    config = read_config()
    runninghub_config = config.get("runninghub") if isinstance(config.get("runninghub"), dict) else {}
    standard_default = str(runninghub_config.get("primary_instance") or "").strip().lower() == "default"
    provider_policy = {
        "runninghub_primary": "standard_24gb" if standard_default else "lite",
        "ordinary_timeout_never_upgrades": True,
        "plus_48gb_allowed": False,
    }
    approval_policy = {
        "scope": "unattended_review_candidate",
        "user_preapproved": True,
        "fallback_script_approved": False,
        "budget_limit_cny": float(config["max_budget_cny"]),
        "formal_publish_requires_human": True,
    }
    if standard_default:
        provider_policy.update({
            "authorized_instance": "default",
            "authorization_scope": "global_default",
            "authorization_target_date": target_value.isoformat(),
            "authorization_reason": "用户已将 RunningHub Standard 24GB 设为每日默认机型",
            "authorization_recorded_at": _now(),
        })
        approval_policy.update({
            "authorized_instance": "default",
            "runninghub_standard_24gb_preapproved": True,
        })
    else:
        provider_policy.update({
            "lite_request_mode": "verified_lite_only",
            "lite_verified": daily_billing_safety().get("state") == "verified_lite",
            "standard_24gb_only_on_oom": True,
        })
    run = {
        "version": "1.0",
        "run_id": f"daily-tech-{target_value.isoformat()}",
        "target_date": target_value.isoformat(),
        "window": {"start": start.isoformat(), "end": end.isoformat(), "timezone": "Asia/Shanghai"},
        "trigger": trigger,
        "status": "queued",
        "current_stage": "research",
        "created_at": _now(),
        "updated_at": _now(),
        "project_id": None,
        "stages": {name: _new_stage(name) for name in STAGE_ORDER},
        "budget": {
            "currency": "CNY",
            "limit": float(config["max_budget_cny"]),
            "reserved": 0.0,
            "spent": 0.0,
            "entries": [],
        },
        "provider_policy": provider_policy,
        "approval_policy": approval_policy,
    }
    _atomic_json(path, run)
    return run


def read_run(target: date | str) -> dict[str, Any] | None:
    return _read_json(_run_path(target))


def _stable_fingerprint(value: object) -> str:
    payload = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def evaluate_media_release(script: dict[str, Any] | None) -> dict[str, Any]:
    """Return the versioned content decision independently of provider/cost gates."""
    document = script if isinstance(script, dict) else {}
    review = document.get("editorial_review") if isinstance(document.get("editorial_review"), dict) else {}
    validation = document.get("validation") if isinstance(document.get("validation"), dict) else {}
    selection = document.get("topic_selection") if isinstance(document.get("topic_selection"), dict) else {}
    scores = review.get("scores") if isinstance(review.get("scores"), dict) else {}
    stories = selection.get("selected_stories") if isinstance(selection.get("selected_stories"), list) else []
    if not stories:
        stories = document.get("stories") if isinstance(document.get("stories"), list) else []
    levels = [str(item.get("heat_level") or "") for item in stories if isinstance(item, dict)]
    input_view = {
        "script": document,
        "rule_version": MEDIA_RELEASE_RULE_VERSION,
    }
    fingerprint = _stable_fingerprint(input_view)
    reasons: list[str] = []
    structured_issues = (
        review.get("structured_issues")
        if isinstance(review.get("structured_issues"), list)
        else []
    )
    factual_or_duration_redline = any(
        isinstance(item, dict)
        and (
            item.get("hard_fact_boundary") is True
            or str(item.get("code") or "")
            in {"factual_boundary", "episode_duration_over_target"}
        )
        for item in structured_issues
    )
    hard_fail = (
        validation.get("passed") is False
        or validation.get("valid") is False
        or bool(review.get("hard_gate_failures"))
        or bool(validation.get("errors"))
        or factual_or_duration_redline
    )
    if hard_fail:
        decision = "blocked"
        reasons.append("事实、结构、来源或独立复验硬门未通过")
    else:
        total = int(review.get("total") or 0)
        hook = int(scores.get("hook") or 0)
        dialogue = int(scores.get("dialogue") or 0)
        density = int(scores.get("information_density") or scores.get("info_density") or 0)
        premium = total >= 85 and hook >= 16 and dialogue >= 16 and density >= 20
        reliable = total >= 78 and hook >= 14 and dialogue >= 15 and density >= 19
        if premium:
            decision = "auto_release"
            reasons.append("传播总分及钩子、对话、信息密度高级单项全部达到自动放行线")
        elif reliable:
            decision = "fallback_review_candidate"
            reasons.append("脚本达到可靠可用线；允许生成带醒目标记的待审视频，但禁止自动发布")
        else:
            decision = "blocked"
            reasons.append("脚本未达到78分可靠线或最低钩子、对话、信息密度单项门")
        if "H3" in levels:
            reasons.append("本期含H3热点，沿用85分优质线")
        elif levels:
            reasons.append("本期最高为H1/H2；达到可靠线后只自动生成待审视频")
    return {
        "decision": decision,
        "rule_version": MEDIA_RELEASE_RULE_VERSION,
        "input_fingerprint": fingerprint,
        "reasons": reasons,
        "media_generation_allowed": decision in {"auto_release", "fallback_review_candidate"},
        "publish_requires_human": True,
        "evaluated_at": _now(),
    }


def script_requires_fallback_approval(script: dict[str, Any] | None) -> bool:
    """Keep the legacy manual-approval API valid for review candidates.

    Approval is no longer required to render a review-only video, but an
    editor may still record an explicit acknowledgement before a manual
    resume or audit export.
    """
    return evaluate_media_release(script).get("decision") == "fallback_review_candidate"


def fallback_approval_matches(run: dict[str, Any], script: dict[str, Any]) -> bool:
    policy = run.get("approval_policy") if isinstance(run.get("approval_policy"), dict) else {}
    return bool(
        policy.get("fallback_script_approved") is True
        and str(policy.get("fallback_script_fingerprint") or "")
        == str(evaluate_media_release(script).get("input_fingerprint") or "")
    )


def approve_fallback_script(target: date | str) -> dict[str, Any]:
    """Persist an optional human acknowledgement for a reliable review draft."""
    run = read_run(target)
    if not run:
        raise DailyAutomationError("未找到该日期的每日任务")
    script = _read_json(_run_path(target).parent / "daily_script.json") or {}
    if not script_requires_fallback_approval(script):
        raise DailyAutomationError("该脚本不是达到可靠线的待审候选稿")
    policy = run.setdefault("approval_policy", {})
    policy.update({
        "fallback_script_approved": True,
        "fallback_script_approved_at": _now(),
        "fallback_script_approved_by": "human",
        "fallback_script_fingerprint": evaluate_media_release(script)["input_fingerprint"],
    })
    run["status"] = "queued"
    run["current_stage"] = "voice"
    return _save_run(run)


def request_text_story_replacement(target: date | str) -> dict[str, Any]:
    """Keep the current lead, queue a different bounded episode combination, and preserve history."""
    run = read_run(target)
    if not run:
        raise DailyAutomationError("未找到该日期的每日任务")
    target_dir = _run_path(target).parent
    selection_path = target_dir / "topic_selection_v2.json"
    selection = _read_json(selection_path) or {}
    combinations = [
        item for item in selection.get("episode_combinations") or []
        if isinstance(item, dict) and item.get("selected_stories")
    ]
    if len(combinations) < 2:
        raise DailyAutomationError("当前没有可替换的差异化候选组合")
    ledger_path = target_dir / "daily_text_attempts.json"
    ledger = _read_json(ledger_path) or {}
    current_ids = set((ledger.get("best_candidate") or {}).get("event_ids") or [])
    if not current_ids:
        current_ids = {
            str(item.get("event_id") or "")
            for item in selection.get("selected_stories") or []
            if isinstance(item, dict)
        }
    current_stories = selection.get("selected_stories") or []
    lead_id = str((current_stories[0] if current_stories else {}).get("event_id") or "")
    if not lead_id:
        raise DailyAutomationError("当前选题缺少可锁定的头条")
    alternate = next(
        (
            item for item in combinations
            if lead_id in set(item.get("event_ids") or [])
            and set(item.get("event_ids") or []) != current_ids
            and not item.get("blocking_issues")
        ),
        None,
    )
    if not alternate:
        raise DailyAutomationError("保留当前头条后，没有其他通过组合门的候选")
    if ledger:
        stamp = re.sub(r"[^0-9]", "", _now())[:14]
        _atomic_json(target_dir / f"daily_text_attempts.before-manual-{stamp}.json", ledger)
    ordered = [alternate, *(item for item in combinations if item is not alternate)]
    for index, item in enumerate(ordered, 1):
        item["rank"] = index
    selection["episode_combinations"] = ordered
    selection["selected_stories"] = alternate["selected_stories"]
    selection["manual_recovery_revision"] = int(selection.get("manual_recovery_revision") or 0) + 1
    selection["manual_preferences"] = {
        "locked_event_ids": [lead_id],
        "requested_at": _now(),
        "action": "keep_lead_replace_weak_story",
    }
    summary = selection.setdefault("selection_summary", {})
    summary.update(
        {
            "selected_combination_id": alternate.get("combination_id"),
            "selected_count": len(alternate.get("selected_stories") or []),
            "episode_score": alternate.get("episode_score"),
            "duration_profile": alternate.get("duration_profile"),
        }
    )
    _atomic_json(selection_path, selection)
    stage = run["stages"]["script"]
    stage.update(
        {
            "status": "pending",
            "finished_at": None,
            "message": "已锁定头条并切换候选组合；等待重新执行文本双门",
            "error": None,
        }
    )
    run["status"] = "queued"
    run["current_stage"] = "script"
    run.setdefault("approval_policy", {}).pop("editorial_recovery_reason", None)
    return _save_run(run)


def authorize_runninghub_standard_for_run(
    target: date | str,
    *,
    reason: str,
    max_budget_cny: float | None = None,
) -> dict[str, Any]:
    """Persist a one-run Standard 24GB authorization without changing global defaults."""
    run = create_or_resume_run(target, trigger="scheduled")
    configured_limit = float((run.get("budget") or {}).get("limit") or read_config()["max_budget_cny"])
    authorized_limit = configured_limit if max_budget_cny is None else min(configured_limit, float(max_budget_cny))
    if authorized_limit <= 0:
        raise DailyAutomationError("RunningHub Standard 单次授权预算必须大于0")
    provider_policy = run.setdefault("provider_policy", {})
    provider_policy.update(
        {
            "runninghub_primary": "standard_24gb",
            "authorized_instance": "default",
            "authorization_scope": "single_daily_run",
            "authorization_target_date": str(run["target_date"]),
            "authorization_reason": re.sub(r"\s+", " ", str(reason or "")).strip()[:300],
            "authorization_recorded_at": _now(),
            "ordinary_timeout_never_upgrades": True,
        }
    )
    approval_policy = run.setdefault("approval_policy", {})
    approval_policy.update(
        {
            "authorized_instance": "default",
            "runninghub_standard_24gb_preapproved": True,
            "budget_limit_cny": authorized_limit,
            "formal_publish_requires_human": True,
        }
    )
    run.setdefault("budget", {})["limit"] = authorized_limit
    return _save_run(run)


def _save_run(run: dict[str, Any]) -> dict[str, Any]:
    run["updated_at"] = _now()
    _atomic_json(_run_path(str(run["target_date"])), run)
    operations = run.get("paid_operations")
    if isinstance(operations, dict):
        _atomic_json(_run_path(str(run["target_date"])).parent / "paid_operations.json", operations)
    project_id = str(run.get("project_id") or "")
    if project_id:
        project_dir = PROJECTS_DIR / project_id
        if project_dir.is_dir():
            _atomic_json(project_dir / "artifacts" / "daily_run.json", run)
    return run


def update_stage(
    run: dict[str, Any],
    stage_name: str,
    status: str,
    *,
    message: str | None = None,
    error: str | None = None,
    output: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if stage_name not in STAGE_ORDER:
        raise DailyAutomationError(f"未知自动化阶段：{stage_name}")
    stage = run["stages"][stage_name]
    if status == "running":
        stage["attempts"] = int(stage.get("attempts") or 0) + 1
        stage["started_at"] = _now()
        stage["finished_at"] = None
        stage["error"] = None
        run["status"] = "running"
        run["current_stage"] = stage_name
    elif status in {"succeeded", "failed", "skipped"}:
        stage["finished_at"] = _now()
        if status == "failed":
            run["status"] = "failed"
    stage["status"] = status
    if message is not None:
        stage["message"] = message
    if error is not None:
        stage["error"] = str(error)[:1000]
    if output:
        stage["output"] = _merge_dict(stage.get("output") or {}, output)
    if status == "succeeded":
        current_index = STAGE_ORDER.index(stage_name)
        if current_index + 1 < len(STAGE_ORDER):
            run["current_stage"] = STAGE_ORDER[current_index + 1]
        else:
            run["status"] = "review_ready"
    return _save_run(run)


def heartbeat_stage(
    run: dict[str, Any],
    stage_name: str,
    *,
    message: str | None = None,
    output: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Persist progress without starting a new attempt or resetting timers."""
    if stage_name not in STAGE_ORDER:
        raise DailyAutomationError(f"未知自动化阶段：{stage_name}")
    stage = run["stages"][stage_name]
    if stage.get("status") != "running":
        raise DailyAutomationError(f"{stage_name} 当前不是运行中状态")
    if message is not None:
        stage["message"] = message
    if output:
        stage["output"] = _merge_dict(stage.get("output") or {}, output)
    return _save_run(run)


def list_runs(limit: int = 20) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if RUNS_ROOT.is_dir():
        for path in sorted(RUNS_ROOT.glob("*/daily_run.json"), reverse=True):
            value = _read_json(path)
            if value:
                rows.append(value)
            if len(rows) >= limit:
                break
    return rows


def daily_billing_safety() -> dict[str, Any]:
    """Require provider billing proof before unattended Lite submissions.

    RunningHub documents Lite as system-scheduled, not as a guaranteed public
    ``instanceType``.  Both an omitted field and an explicit ``lite`` probe
    have been observed billing at Standard rates.  A non-Lite observation is
    therefore a hard pre-submit blocker until a later task proves the actual
    0.4 CNY/hour rate.
    """
    for run in list_runs(30):
        roles = (((run.get("stages") or {}).get("avatar") or {}).get("output") or {}).get("roles") or {}
        if not isinstance(roles, dict):
            continue
        for role, record in roles.items():
            if not isinstance(record, dict):
                continue
            requested = str(record.get("requested_instance") or record.get("instance") or "")
            billing = record.get("billing") if isinstance(record.get("billing"), dict) else {}
            observed = str(record.get("observed_instance") or billing.get("observed_instance") or "unverified")
            if requested not in {"auto_lite", "lite", "explicit_lite_probe"}:
                continue
            latest_seen = {
                "target_date": run.get("target_date"),
                "role": role,
                "task_id": record.get("task_id"),
                "observed_instance": observed,
                "billing": billing,
            }
            if observed == "lite":
                return {
                    "auto_schedule_eligible": True,
                    "state": "verified_lite",
                    "message": "已核验企业 Lite 实际费率；允许开启凌晨自动生产。",
                    "latest_evidence": latest_seen,
                }
            if observed in {"standard_24gb", "plus_48gb", "unverified"}:
                return {
                    "auto_schedule_eligible": False,
                    "state": "provider_did_not_honor_lite",
                    "message": "RunningHub 的 Lite 请求已被实际账单按 Standard 结算；在供应商提供可保证 0.4 元/小时的调用方式前，禁止自动提交付费数字人。",
                    "latest_evidence": latest_seen,
                }
    return {
        "auto_schedule_eligible": False,
        "state": "lite_verification_required",
        "message": "尚无 0.4 元/小时 Lite 账单证据；请先运行单任务最小探测，自动生产不会拿正式角色试错。",
        "latest_evidence": None,
    }


def provider_media_eligibility(run: dict[str, Any]) -> dict[str, Any]:
    """Evaluate provider authorization after content release, never as a script error."""
    approval = run.get("approval_policy") if isinstance(run.get("approval_policy"), dict) else {}
    if approval.get("runninghub_standard_24gb_preapproved") is True and str(approval.get("authorized_instance")) == "default":
        return {
            "eligible": True, "state": "authorized_standard_24gb",
            "reason": "本日期运行已获得持久化的单次 Standard 24GB 授权",
        }
    safety = daily_billing_safety()
    return {
        "eligible": safety.get("auto_schedule_eligible") is True,
        "state": safety.get("state") or "lite_verification_required",
        "reason": safety.get("message") or "RunningHub Lite 尚未取得可核验账单证据",
    }


def douyin_signal_status(config: dict[str, Any], latest_run: dict[str, Any] | None = None) -> dict[str, Any]:
    sources = config.get("douyin_sources") if isinstance(config.get("douyin_sources"), list) else list(DEFAULT_DOUYIN_SOURCES)
    api_configured = bool(_douyin_api_key() and any(str(item.get("api_url") or "").strip() for item in sources if isinstance(item, dict)))
    snapshot_count = sum(1 for item in sources if isinstance(item, dict) and _read_douyin_snapshot(item) is not None)
    latest_results: list[dict[str, Any]] = []
    target_date = str((latest_run or {}).get("target_date") or "")
    if target_date:
        research = _read_json(RUNS_ROOT / target_date / "news_research.json") or {}
        latest_results = [item for item in research.get("douyin_sources") or [] if isinstance(item, dict)]
    ok_results = [item for item in latest_results if item.get("status") == "ok"]
    modes = list(dict.fromkeys(str(item.get("mode") or "") for item in ok_results if item.get("mode")))
    return {
        "api_configured": api_configured,
        "snapshot_count": snapshot_count,
        "latest_target_date": target_date,
        "latest_ok_count": sum(int(item.get("count") or 0) for item in ok_results),
        "latest_modes": modes,
        "latest_sources": latest_results,
        "state": "ok" if ok_results else ("ready" if api_configured or snapshot_count else "unconfigured"),
    }


def copy_skill_hotspot_status(config: dict[str, Any], latest_run: dict[str, Any] | None = None) -> dict[str, Any]:
    feed_config = config.get("copy_skill_hotspot_feed") if isinstance(config.get("copy_skill_hotspot_feed"), dict) else {}
    target_date = str((latest_run or {}).get("target_date") or "")
    research = _read_json(RUNS_ROOT / target_date / "news_research.json") if target_date else None
    latest = research.get("copy_skill_feed") if isinstance(research, dict) and isinstance(research.get("copy_skill_feed"), dict) else {}
    if target_date and not latest:
        source_index = _read_json(RUNS_ROOT / target_date / "inputs" / "copy_skill_hotspot" / "source-index.json") or {}
        if source_index:
            latest = {
                "feed_status": source_index.get("feed_status"),
                "business_date": source_index.get("business_date"),
                "run_id": source_index.get("run_id"),
                "counts": {"candidates": int(source_index.get("candidate_count") or 0)},
                "coverage_warning": source_index.get("coverage_warning"),
                "manifest_validation": source_index.get("manifest_validation"),
            }
    return {
        "enabled": feed_config.get("enabled") is not False,
        "root": str(feed_config.get("root") or DEFAULT_COPY_SKILL_HOTSPOT_ROOT),
        "latest_target_date": target_date or None,
        "state": str(latest.get("feed_status") or ("ready" if feed_config.get("enabled") is not False else "disabled")),
        "business_date": latest.get("business_date"),
        "run_id": latest.get("run_id"),
        "candidate_count": int(((latest.get("counts") or {}).get("candidates") or 0)),
        "coverage_warning": latest.get("coverage_warning"),
        "manifest_validation": latest.get("manifest_validation") or None,
    }


def read_status() -> dict[str, Any]:
    config = read_config()
    runs = list_runs(10)
    active = next((item for item in runs if item.get("status") in ACTIVE_RUN_STATES), None)
    latest = runs[0] if runs else None
    lock = _read_json(RUN_LOCK_PATH)
    lock_pid = int((lock or {}).get("pid") or 0)
    try:
        from backlot.daily_script_v2 import golden_script_status

        golden_scripts = golden_script_status()
    except Exception:  # noqa: BLE001 - status page must survive one malformed optional sample.
        golden_scripts = {"available_count": 0, "loaded_count": 0, "loaded": [], "ignored": []}
    return {
        "config": config,
        "next_run": next_scheduled_run(config),
        "active_run": _public_run(active),
        "latest_run": _public_run(latest),
        "history": [_public_run(item) for item in runs[:5]],
        "run_lock": {
            "present": bool(lock),
            "pid": lock_pid or None,
            "alive": bool(lock_pid and _process_is_alive(lock_pid)),
            "target_date": (lock or {}).get("target_date"),
            "trigger": (lock or {}).get("trigger"),
            "acquired_at": (lock or {}).get("acquired_at"),
        },
        "text_ai": read_text_ai_config(),
        "text_providers": read_text_provider_status(),
        "golden_scripts": golden_scripts,
        "douyin_signals": douyin_signal_status(config, latest),
        "copy_skill_hotspot_feed": copy_skill_hotspot_status(config, latest),
        "billing_safety": daily_billing_safety(),
    }


def _public_run(run: dict[str, Any] | None) -> dict[str, Any] | None:
    if not run:
        return None
    public = {
        key: run.get(key)
        for key in (
            "run_id",
            "target_date",
            "trigger",
            "status",
            "current_stage",
            "created_at",
            "updated_at",
            "project_id",
            "window",
            "stages",
            "budget",
            "provider_policy",
            "approval_policy",
            "media_release_decision",
            "provider_eligibility",
            "copy_skill_feed",
        )
    }
    project_id = str(run.get("project_id") or "")
    target_date = str(run.get("target_date") or "")
    script = _read_json(RUNS_ROOT / target_date / "daily_script.json") if target_date else None
    selection = _read_json(RUNS_ROOT / target_date / "topic_selection_v2.json") if target_date else None
    research = _read_json(RUNS_ROOT / target_date / "news_research.json") if target_date else None
    if "copy_skill_feed" not in public and isinstance(research, dict) and isinstance(research.get("copy_skill_feed"), dict):
        public["copy_skill_feed"] = research["copy_skill_feed"]
    if "copy_skill_feed" not in public and target_date:
        source_index = _read_json(RUNS_ROOT / target_date / "inputs" / "copy_skill_hotspot" / "source-index.json") or {}
        if source_index:
            public["copy_skill_feed"] = {
                "feed_status": source_index.get("feed_status"),
                "business_date": source_index.get("business_date"),
                "run_id": source_index.get("run_id"),
                "counts": {"candidates": int(source_index.get("candidate_count") or 0)},
                "coverage_warning": source_index.get("coverage_warning"),
                "manifest_validation": source_index.get("manifest_validation"),
                "snapshot_paths": {
                    "source_index": str(RUNS_ROOT / target_date / "inputs" / "copy_skill_hotspot" / "source-index.json"),
                    "candidate_snapshot": str(RUNS_ROOT / target_date / "inputs" / "copy_skill_hotspot" / "candidate-snapshot.json"),
                    "top10_report": str(RUNS_ROOT / target_date / "inputs" / "copy_skill_hotspot" / "top10-report.json"),
                },
            }
    text_ledger = _read_json(RUNS_ROOT / target_date / "daily_text_attempts.json") if target_date else None
    if script:
        stories = script.get("stories") if isinstance(script.get("stories"), list) else []
        lines = script.get("lines") if isinstance(script.get("lines"), list) else []
        public["script_summary"] = {
            "story_count": len(stories),
            "line_count": len(lines),
            "headlines": [str(item.get("headline") or "") for item in stories[:5] if isinstance(item, dict)],
            "intro": next((str(item.get("text") or "") for item in lines if isinstance(item, dict)), ""),
            "outro": next((str(item.get("text") or "") for item in reversed(lines) if isinstance(item, dict) and item.get("kind") == "outro"), ""),
            "quality_band": str(((script.get("editorial_review") or {}).get("quality_band") or "")),
            "editorial_score": int(((script.get("editorial_review") or {}).get("total") or 0)),
        }
    if selection or text_ledger:
        combinations = []
        for item in (selection or {}).get("episode_combinations") or []:
            if not isinstance(item, dict):
                continue
            combinations.append(
                {
                    "combination_id": item.get("combination_id"),
                    "rank": item.get("rank"),
                    "episode_score": item.get("episode_score"),
                    "story_count": item.get("story_count"),
                    "duration_profile": item.get("duration_profile"),
                    "topic_families": item.get("topic_families") or [],
                    "event_forms": item.get("event_forms") or [],
                    "blocking_issues": item.get("blocking_issues") or [],
                    "stories": [
                        {
                            "event_id": story.get("event_id"),
                            "headline": story.get("canonical_title"),
                            "heat_level": story.get("heat_level"),
                            "topic_family": story.get("topic_family"),
                            "event_form": story.get("event_form"),
                            "marginal_contribution": story.get("marginal_contribution"),
                            "replacement_priority": story.get("replacement_priority"),
                        }
                        for story in item.get("selected_stories") or []
                        if isinstance(story, dict)
                    ],
                }
            )
        ledger = text_ledger or {}
        public["text_resilience"] = {
            "state": ledger.get("recovery_state") or ((((run.get("stages") or {}).get("script") or {}).get("output") or {}).get("text_resilience") or {}).get("state"),
            "safe_resume_point": ledger.get("safe_resume_point"),
            "terminal_reason": ledger.get("terminal_reason"),
            "policy": ledger.get("policy") or {},
            "editorial_reviews_used": int(ledger.get("editorial_reviews_used") or 0),
            "attempt_count": len(ledger.get("attempts") or []),
            "best_score": int(((ledger.get("best_candidate") or {}).get("editorial_score") or 0)),
            "best_combination_id": (ledger.get("best_candidate") or {}).get("combination_id"),
            "attempts": [
                {
                    "attempt_id": item.get("attempt_id"),
                    "combination_id": item.get("combination_id"),
                    "status": item.get("status"),
                    "editorial_score": item.get("editorial_score"),
                    "event_ids": item.get("event_ids") or [],
                    "recovery": item.get("recovery") or {},
                    "started_at": item.get("started_at"),
                    "finished_at": item.get("finished_at"),
                }
                for item in ledger.get("attempts") or []
                if isinstance(item, dict)
            ],
            "combinations": combinations,
        }
    if run.get("status") == "failed":
        failed_name = next((name for name in STAGE_ORDER if ((run.get("stages") or {}).get(name) or {}).get("status") == "failed"), None)
        failed_stage = ((run.get("stages") or {}).get(failed_name) or {}) if failed_name else {}
        raw_error = str(failed_stage.get("error") or "")
        classification = classify_runninghub_failure(raw_error)
        if classification.get("is_oom"):
            summary = "供应商明确报告显存不足；恢复时允许只为该角色升级 Standard 24GB。"
        elif classification.get("kind") == "transient":
            summary = "网络或供应商临时连接失败；可以从当前阶段继续，仍保持企业 Lite，不会因此升级算力。"
        else:
            summary = "本阶段未完成；已成功产物会保留，恢复前请核对配置与错误详情。"
        public["failure"] = {
            "stage": failed_name,
            "summary": summary,
            "retryable": bool(raw_error),
            "may_upgrade_to_standard": bool(classification.get("may_upgrade_to_standard")),
        }
    if project_id and run.get("current_stage") == "visuals":
        workbench = _read_json(PROJECTS_DIR / project_id / "artifacts" / "workbench.json") or {}
        batch = ((workbench.get("automation") or {}).get("visual_batch") or {})
        if batch:
            public["live_progress"] = {
                "kind": "visual_batch",
                "status": batch.get("status"),
                "completed": int(batch.get("completed_slots") or 0),
                "total": int(batch.get("total_slots") or 0),
                "failed": int(batch.get("failed_slots") or 0),
                "updated_at": workbench.get("updated_at"),
            }
    avatar_roles = (((run.get("stages") or {}).get("avatar") or {}).get("output") or {}).get("roles") or {}
    if isinstance(avatar_roles, dict):
        public["billing_audit"] = {
            role: {
                "task_id": value.get("task_id"),
                "requested_instance": value.get("requested_instance") or value.get("instance"),
                "observed_instance": value.get("observed_instance"),
                "actual_cost_cny": value.get("actual_cost_cny"),
                "billing": value.get("billing"),
                "billing_blocker": value.get("billing_blocker"),
            }
            for role, value in avatar_roles.items()
            if isinstance(value, dict)
        }
    operations = ((run.get("paid_operations") or {}).get("operations") or {}) if isinstance(run.get("paid_operations"), dict) else {}
    public["paid_operations"] = [
        {
            "operation_id": value.get("operation_id"), "stage": value.get("stage"),
            "role": value.get("role"), "state": value.get("state"),
            "task_id": value.get("task_id"), "requested_instance": value.get("requested_instance"),
            "updated_at": value.get("updated_at"),
        }
        for value in operations.values() if isinstance(value, dict)
    ]
    return public


def classify_runninghub_failure(error: object) -> dict[str, Any]:
    def safe_fragments(value: object, *, depth: int = 0) -> list[str]:
        if depth > 3:
            return []
        if isinstance(value, dict):
            allowed = (
                "error", "message", "reason", "status", "exception_type",
                "exception_message", "error_message", "node_name", "node_id",
                "failure_details",
            )
            fragments: list[str] = []
            for key in allowed:
                if key in value:
                    fragments.extend(safe_fragments(value.get(key), depth=depth + 1))
            return fragments
        if isinstance(value, (list, tuple)):
            fragments = []
            for item in value[:20]:
                fragments.extend(safe_fragments(item, depth=depth + 1))
            return fragments
        if value is None:
            return []
        return [str(value)[:500]]

    fragments = [re.sub(r"\s+", " ", item).strip() for item in safe_fragments(error) if str(item).strip()]
    message = " | ".join(fragments)
    oom_evidence = [
        pattern for pattern in OOM_PATTERNS
        if any(re.search(pattern, fragment, re.IGNORECASE) for fragment in fragments)
    ]
    non_oom_evidence = [
        pattern for pattern in NON_OOM_PATTERNS
        if any(re.search(pattern, fragment, re.IGNORECASE) for fragment in fragments)
    ]
    # A provider-supplied OutOfMemoryError/显存不足 marker is stronger than
    # generic transport words that may also appear in a diagnostic summary.
    is_oom = bool(oom_evidence)
    return {
        "kind": "oom" if is_oom else ("transient" if non_oom_evidence else "unknown"),
        "is_oom": is_oom,
        "explicit": is_oom,
        "may_upgrade_to_standard": is_oom,
        "evidence": oom_evidence if is_oom else non_oom_evidence,
        "message": message[:500],
    }


@dataclass
class BudgetLedger:
    run: dict[str, Any]

    @property
    def limit(self) -> float:
        return float(self.run["budget"]["limit"])

    @property
    def committed(self) -> float:
        return round(float(self.run["budget"].get("reserved") or 0) + float(self.run["budget"].get("spent") or 0), 4)

    def _entry(self, operation_id: str, kind: str) -> dict[str, Any] | None:
        return next((item for item in self.run["budget"].get("entries") or []
                     if item.get("operation_id") == operation_id and item.get("type") == kind), None)

    def _assert_invariants(self) -> None:
        budget = self.run["budget"]
        reserved = round(float(budget.get("reserved") or 0), 4)
        spent = round(float(budget.get("spent") or 0), 4)
        if reserved < 0 or spent < 0 or reserved + spent > self.limit + 1e-9:
            raise DailyAutomationError("付费预算账本不变量被破坏，已停止后续提交")

    def reserve_once(self, operation_id: str, amount: float, *, purpose: str, provider: str = "runninghub") -> bool:
        if self._entry(operation_id, "reserve"):
            return False
        self.reserve(amount, purpose=purpose, provider=provider, operation_id=operation_id)
        return True

    def settle_once(self, operation_id: str, reserved: float, actual: float, *, purpose: str, task_id: str | None = None) -> bool:
        if self._entry(operation_id, "settle"):
            return False
        if self._entry(operation_id, "release"):
            raise DailyAutomationError("已释放的付费操作不能再结算")
        self.settle(reserved, actual, purpose=purpose, task_id=task_id, operation_id=operation_id)
        return True

    def release_once(self, operation_id: str, amount: float, *, purpose: str, reason: str) -> bool:
        if self._entry(operation_id, "release"):
            return False
        if self._entry(operation_id, "settle"):
            raise DailyAutomationError("已结算的付费操作不能再释放")
        self.release(amount, purpose=purpose, reason=reason, operation_id=operation_id)
        return True

    def reserve(self, amount: float, *, purpose: str, provider: str = "runninghub", operation_id: str | None = None) -> None:
        value = round(float(amount), 4)
        if value <= 0:
            raise DailyAutomationError("预算预留金额必须大于 0")
        if self.committed + value > self.limit + 1e-9:
            raise DailyAutomationError(f"本次付费任务将超过每日 {self.limit:g} 元预算，已停止提交")
        budget = self.run["budget"]
        budget["reserved"] = round(float(budget.get("reserved") or 0) + value, 4)
        budget["entries"].append({
            "at": _now(), "type": "reserve", "provider": provider, "purpose": purpose, "amount": value,
            "operation_id": operation_id,
        })
        self._assert_invariants()
        _save_run(self.run)

    def settle(self, reserved: float, actual: float, *, purpose: str, task_id: str | None = None, operation_id: str | None = None) -> None:
        reserve_value = max(0.0, round(float(reserved), 4))
        actual_value = max(0.0, round(float(actual), 4))
        budget = self.run["budget"]
        budget["reserved"] = round(max(0.0, float(budget.get("reserved") or 0) - reserve_value), 4)
        if float(budget.get("spent") or 0) + actual_value > self.limit + 1e-9:
            raise DailyAutomationError("RunningHub 实际费用超过每日预算上限，已阻止后续付费任务")
        budget["spent"] = round(float(budget.get("spent") or 0) + actual_value, 4)
        budget["entries"].append({
            "at": _now(), "type": "settle", "provider": "runninghub", "purpose": purpose,
            "reserved": reserve_value, "actual": actual_value, "task_id": task_id,
            "operation_id": operation_id,
        })
        self._assert_invariants()
        _save_run(self.run)

    def release(self, amount: float, *, purpose: str, reason: str, operation_id: str | None = None) -> None:
        """Release an unused reservation without disguising it as free work."""
        value = max(0.0, round(float(amount), 4))
        budget = self.run["budget"]
        budget["reserved"] = round(max(0.0, float(budget.get("reserved") or 0) - value), 4)
        budget["entries"].append({
            "at": _now(), "type": "release", "provider": "runninghub",
            "purpose": purpose, "amount": value, "reason": str(reason)[:300],
            "operation_id": operation_id,
        })
        self._assert_invariants()
        _save_run(self.run)


def ensure_paid_operation(run: dict[str, Any], operation_id: str, **fields: Any) -> dict[str, Any]:
    ledger = run.setdefault("paid_operations", {"version": "1.0", "operations": {}})
    operations = ledger.setdefault("operations", {})
    operation = operations.setdefault(operation_id, {
        "operation_id": operation_id, "state": "planned", "history": [], "created_at": _now(),
    })
    for key, value in fields.items():
        operation.setdefault(key, value)
    return operation


def transition_paid_operation(run: dict[str, Any], operation_id: str, state: str, **fields: Any) -> dict[str, Any]:
    operation = ensure_paid_operation(run, operation_id)
    allowed = {
        "planned": {"reserved", "cancelled"},
        "reserved": {"submitting", "released"},
        "submitting": {"submitted", "ambiguous", "released"},
        "submitted": {"running", "succeeded", "failed"},
        "running": {"running", "succeeded", "failed"},
        "ambiguous": {"submitted", "released"},
        "failed": {"reserved"},
        "succeeded": set(), "released": set(), "cancelled": set(),
    }
    previous = str(operation.get("state") or "planned")
    if state != previous and state not in allowed.get(previous, set()):
        raise DailyAutomationError(f"非法付费操作状态迁移：{previous} → {state}")
    operation.update(fields)
    operation["state"] = state
    operation["updated_at"] = _now()
    operation.setdefault("history", []).append({"at": _now(), "from": previous, "to": state})
    _save_run(run)
    return operation


def _strip_markup(value: str) -> str:
    text = re.sub(r"<[^>]+>", " ", html.unescape(value or ""))
    return re.sub(r"\s+", " ", text).strip()


def _first_text(element: ET.Element, names: Iterable[str]) -> str:
    wanted = {name.lower() for name in names}
    for child in element.iter():
        name = child.tag.rsplit("}", 1)[-1].lower()
        if name in wanted and child.text and child.text.strip():
            return child.text.strip()
    return ""


def _entry_url(element: ET.Element) -> str:
    for child in element.iter():
        if child.tag.rsplit("}", 1)[-1].lower() != "link":
            continue
        href = str(child.attrib.get("href") or "").strip()
        relation = str(child.attrib.get("rel") or "alternate").lower()
        if href and relation in {"alternate", ""}:
            return href
        if child.text and child.text.strip().startswith("http"):
            return child.text.strip()
    return ""


def _parse_published(value: str) -> datetime | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except ValueError:
        try:
            parsed = parsedate_to_datetime(raw)
        except (TypeError, ValueError, OverflowError):
            return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=UTC)
    return parsed.astimezone(LOCAL_TIMEZONE)


def _canonical_url(value: str) -> str:
    try:
        parts = urlsplit(value.strip())
    except ValueError:
        return value.strip()
    path = re.sub(r"/+", "/", parts.path).rstrip("/") or "/"
    return urlunsplit((parts.scheme.lower(), parts.netloc.lower(), path, "", ""))


def _decode_google_news_url(value: str, *, session: requests.Session | None = None) -> str:
    """Resolve a Google News RSS wrapper to the publisher URL.

    Current Google News links need the article id, timestamp, and signature from
    the wrapper page.  Resolution is best-effort: research collection must stay
    usable when Google rate-limits this non-critical enrichment step.
    """
    parts = urlsplit(str(value or ""))
    if parts.netloc.lower() != "news.google.com" or "/articles/" not in parts.path:
        return value
    client = session or requests.Session()
    headers = {"User-Agent": "Mozilla/5.0 OpenMontage/1.0"}
    response = client.get(value, timeout=20, headers=headers)
    response.raise_for_status()
    page = response.text
    matches = {
        name: re.search(rf'data-n-a-{name}="([^"]+)"', page)
        for name in ("id", "ts", "sg")
    }
    if not all(matches.values()):
        raise DailyAutomationError("Google 新闻链接缺少原站解析参数")
    article_id = matches["id"].group(1)
    timestamp = int(matches["ts"].group(1))
    signature = matches["sg"].group(1)
    request_value = [
        "garturlreq",
        [["X", "X", ["X", "X"], None, None, 1, 1, "US:en", None, 1, None, None, None, None, None, 0, 1],
         "X", "X", 1, [1, 1, 1], 1, 1, None, 0, 0, None, 0],
        article_id,
        timestamp,
        signature,
    ]
    envelope = [[["Fbv4je", json.dumps(request_value, separators=(",", ":")), None, "generic"]]]
    decoded = client.post(
        "https://news.google.com/_/DotsSplashUi/data/batchexecute",
        data={"f.req": json.dumps(envelope, separators=(",", ":"))},
        timeout=20,
        headers={**headers, "Content-Type": "application/x-www-form-urlencoded;charset=UTF-8"},
    )
    decoded.raise_for_status()
    for block in decoded.text.split("\n\n"):
        try:
            rows = json.loads(block)
        except json.JSONDecodeError:
            continue
        for row in rows if isinstance(rows, list) else []:
            if not isinstance(row, list) or len(row) < 3 or row[:2] != ["wrb.fr", "Fbv4je"]:
                continue
            try:
                result = json.loads(row[2])
            except (TypeError, json.JSONDecodeError):
                continue
            if isinstance(result, list) and len(result) > 1 and str(result[1]).startswith("http"):
                return str(result[1])
    raise DailyAutomationError("Google 新闻链接未返回原站地址")


def _extract_article_evidence(page: str) -> str:
    """Extract a compact evidence excerpt without adding a crawler dependency."""
    fragments: list[str] = []
    for match in re.finditer(
        r'<meta[^>]+(?:name|property)=["\'](?:description|og:description)["\'][^>]+content=["\']([^"\']+)',
        page,
        re.IGNORECASE,
    ):
        fragments.append(_strip_markup(match.group(1)))
    for match in re.finditer(r'"articleBody"\s*:\s*"((?:\\.|[^"\\])*)"', page, re.IGNORECASE):
        try:
            fragments.append(json.loads(f'"{match.group(1)}"'))
        except json.JSONDecodeError:
            continue
    # Publisher pages commonly keep article copy in paragraphs even when the
    # surrounding container differs.  Navigation snippets are too short and
    # are discarded here.
    for match in re.finditer(r"<p\b[^>]*>(.*?)</p>", page, re.IGNORECASE | re.DOTALL):
        paragraph = _strip_markup(match.group(1))
        if len(paragraph) >= 24:
            fragments.append(paragraph)
    unique: list[str] = []
    seen: set[str] = set()
    for fragment in fragments:
        normalized = _clean_line_text(fragment)
        key = re.sub(r"\W+", "", normalized)
        if len(normalized) < 24 or key in seen:
            continue
        seen.add(key)
        unique.append(normalized)
    combined = "\n".join(unique)
    if len(combined) <= 3600:
        return combined
    # Keep both the event explanation near the top and publisher warnings or
    # user guidance commonly placed near the end.
    return f"{combined[:2500]}\n[正文中段已压缩]\n{combined[-1100:]}"


def _enrich_candidate_evidence(candidate: dict[str, Any]) -> None:
    """Attach publisher-level evidence to one shortlisted candidate in place."""
    candidate["evidence_status"] = "failed"
    try:
        client = requests.Session()
        original_url = _decode_google_news_url(str(candidate.get("url") or ""), session=client)
        response = client.get(original_url, timeout=20, headers={"User-Agent": "Mozilla/5.0 OpenMontage/1.0"})
        response.raise_for_status()
        response.encoding = response.apparent_encoding or response.encoding
        evidence = _extract_article_evidence(response.text)
        if len(evidence) < 80:
            raise DailyAutomationError("原文正文证据不足")
        candidate["evidence_url"] = _canonical_url(original_url)
        candidate["evidence_excerpt"] = evidence
        candidate["evidence_status"] = "ok"
        official = _extract_article_og_image(response.text, base_url=original_url)
        if official["image_url"]:
            host = urlsplit(original_url).netloc.lower().removeprefix("www.")
            candidate["evidence_image_url"] = _canonical_url(official["image_url"])
            candidate["evidence_image_attribution"] = official["attribution"] or host
            # 官方/媒体配图属于新闻报道用途；后续画面须保留署名来源，不得曲解或商用篡改。
            candidate["evidence_image_license"] = "press"
    except Exception as exc:  # noqa: BLE001 - this gate reports failure to the model.
        candidate["evidence_error"] = re.sub(r"https?://\S+", "[地址已隐藏]", str(exc))[:200]


def _extract_article_og_image(page: str, *, base_url: str = "") -> dict[str, str]:
    """Pull the article's official share image and publisher name for attribution.

    Prefers ``og:image``, then ``twitter:image``, then the first absolute
    ``<img>`` as a fallback.  Returns ``{"image_url", "attribution"}`` with
    empty strings when no usable image is present.  Purely textual and
    side-effect free.
    """

    def find_meta(attr: str, value: str) -> str:
        for match in re.finditer(r"<meta\b[^>]*>", page, re.IGNORECASE):
            tag = match.group(0)
            if re.search(rf'\b{attr}\s*=\s*["\']{re.escape(value)}["\']', tag, re.IGNORECASE):
                content = re.search(r'\bcontent\s*=\s*["\']([^"\']+)["\']', tag, re.IGNORECASE)
                if content:
                    return html.unescape(content.group(1))
        return ""

    image = find_meta("property", "og:image") or find_meta("name", "twitter:image")
    attribution = find_meta("property", "og:site_name") or find_meta("name", "application-name")
    if not image:
        first_img = re.search(r'<img\b[^>]*\bsrc\s*=\s*["\'](https?://[^"\']+)["\']', page, re.IGNORECASE)
        if first_img:
            image = first_img.group(1)
    if image and base_url:
        image = urljoin(base_url, image)
    return {"image_url": (image or "").strip(), "attribution": _clean_line_text(attribution)}


def _evidence_priority(candidate: dict[str, Any]) -> int:
    """Rank title-level leads for the bounded publisher-evidence prefetch."""
    corpus = f"{candidate.get('title', '')} {candidate.get('summary', '')}"
    score = 0
    for pattern, weight in (
        (r"诈骗|泄露|漏洞|封号|违法|涨价|降价", 15),
        (r"实机|机器人|比赛|高难|进厂|芯片|火箭", 12),
        (r"首款|首次|电池|手机|游戏", 7),
        (r"免费", 3),
        (r"微信|华为|小米|荣耀|黑神话|DeepSeek", 5),
        (r"财报|营收|收入|研报|预测|招聘", -8),
    ):
        if re.search(pattern, corpus, re.IGNORECASE):
            score += weight
    if str(candidate.get("source_id") or "").startswith("google-news-cn-"):
        score += 2
    return score


def _prefetch_selection_evidence(candidates: list[dict[str, Any]], *, limit: int = 12) -> set[str]:
    """Enrich only a bounded set of high-potential leads before model selection."""
    ranked = sorted(candidates, key=_evidence_priority, reverse=True)
    attempted: set[str] = set()
    title_prefixes: set[str] = set()
    selected: list[dict[str, Any]] = []
    for candidate in ranked:
        title_prefix = re.sub(r"\W+", "", str(candidate.get("title") or "").lower())[:18]
        if title_prefix and title_prefix in title_prefixes:
            continue
        title_prefixes.add(title_prefix)
        selected.append(candidate)
        if len(selected) >= max(1, limit):
            break
    for candidate in selected:
        candidate_id = str(candidate.get("candidate_id") or "")
        if not candidate_id:
            continue
        if "evidence_status" not in candidate:
            _enrich_candidate_evidence(candidate)
        attempted.add(candidate_id)
    return attempted


def _topic_selection_payload_candidates(candidates: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Bound the model packet without weakening the full-corpus evidence gate.

    Daily Google News queries can legitimately return one hundred leads.  Sending
    every title, summary and evidence excerpt to a text model made the JSON
    selector return empty decisions under load.  Keep a deterministic,
    source-diverse short list for reasoning; ``validate_topic_selection`` still
    resolves selected IDs against the complete frozen research corpus.
    """
    def rank(candidate: dict[str, Any]) -> tuple[int, str]:
        hint = candidate.get("china_short_video_hint") if isinstance(candidate.get("china_short_video_hint"), dict) else {}
        china_score = {"high": 24, "medium": 10}.get(str(hint.get("likely_china_relevance") or ""), 0)
        evidence_score = 36 if candidate.get("evidence_status") == "ok" else 0
        authority_score = 12 if candidate.get("authority") == "official" else 0
        return (_evidence_priority(candidate) * 10 + china_score + evidence_score + authority_score, str(candidate.get("candidate_id") or ""))

    selected: list[dict[str, Any]] = []
    per_source: dict[str, int] = {}
    seen_titles: set[str] = set()
    for candidate in sorted(candidates, key=rank, reverse=True):
        source_id = str(candidate.get("source_id") or "unknown")
        title_key = re.sub(r"\W+", "", str(candidate.get("title") or "").lower())[:28]
        if (title_key and title_key in seen_titles) or per_source.get(source_id, 0) >= 5:
            continue
        selected.append(candidate)
        per_source[source_id] = per_source.get(source_id, 0) + 1
        if title_key:
            seen_titles.add(title_key)
        if len(selected) >= MAX_TOPIC_SELECTION_CANDIDATES:
            break
    return selected


def _parse_feed_root(content: bytes) -> ET.Element:
    """Parse real-world RSS while keeping dependencies and failure scope small.

    Some otherwise valid publisher feeds contain a bare ``&`` in a URL or an
    invalid XML control character.  A strict first pass preserves normal XML
    behaviour; the second pass repairs only those two well-known defects.
    """
    try:
        return ET.fromstring(content)
    except ET.ParseError as strict_error:
        text = content.decode("utf-8", errors="replace")
        text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
        text = re.sub(r"&(?!#\d+;|#x[0-9A-Fa-f]+;|[A-Za-z][A-Za-z0-9]+;)", "&amp;", text)
        try:
            return ET.fromstring(text.encode("utf-8"))
        except ET.ParseError:
            raise strict_error


def _collect_baidu_heat_signals(source: dict[str, Any], request_get: Callable[..., Any]) -> list[dict[str, Any]]:
    """Read Baidu's public board snapshot without treating it as article evidence."""
    response = request_get(
        str(source.get("url") or "https://top.baidu.com/board?tab=realtime"),
        timeout=20,
        headers={"User-Agent": "Mozilla/5.0 OpenMontage/1.0"},
    )
    response.raise_for_status()
    page = response.text
    match = re.search(r"<!--s-data:(\{.*?\})-->", page, re.DOTALL)
    if not match:
        raise DailyAutomationError("百度热榜页面没有可读取的公开榜单快照")
    payload = json.loads(match.group(1))
    cards = (((payload.get("data") or {}).get("cards")) or []) if isinstance(payload, dict) else []
    rows = next(
        (card.get("content") for card in cards if isinstance(card, dict) and isinstance(card.get("content"), list)),
        [],
    )
    maximum = max(1, min(int(source.get("max_items") or 50), 100))
    signals: list[dict[str, Any]] = []
    for position, item in enumerate(rows[:maximum], 1):
        if not isinstance(item, dict):
            continue
        title = _clean_line_text(item.get("word") or item.get("query"))
        if not title:
            continue
        try:
            score = int(item.get("hotScore") or 0)
        except (TypeError, ValueError):
            score = 0
        signals.append(
            {
                "signal_id": f"baidu-{position:02d}",
                "source_id": str(source.get("id") or "baidu-realtime"),
                "source_name": str(source.get("name") or "百度实时热榜"),
                "title": title[:160],
                "rank": position,
                "heat_value": score,
                "url": _canonical_url(str(item.get("rawUrl") or item.get("url") or "")),
                "captured_at": _now(),
                "scope": "domestic_public_heat_snapshot",
            }
        )
    return signals


def collect_heat_signals(
    *,
    sources: list[dict[str, Any]],
    request_get: Callable[..., Any] = requests.get,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    signals: list[dict[str, Any]] = []
    results: list[dict[str, Any]] = []
    for source in sources:
        source_id = str(source.get("id") or "heat-source")
        result = {"id": source_id, "name": str(source.get("name") or source_id), "status": "failed", "count": 0}
        try:
            kind = str(source.get("kind") or "")
            if kind == "baidu_board":
                rows = _collect_baidu_heat_signals(source, request_get)
            else:
                raise DailyAutomationError(f"不支持的热度来源类型：{kind or '[空]'}")
            signals.extend(rows)
            result.update({"status": "ok", "count": len(rows)})
        except Exception as exc:  # noqa: BLE001 - one heat source is never fatal to evidence collection.
            result["error"] = re.sub(r"https?://\S+", "[地址已隐藏]", str(exc))[:300]
        results.append(result)
    return signals, results


def _read_douyin_snapshot(source: dict[str, Any]) -> dict[str, Any] | None:
    relative = str(source.get("snapshot_path") or "").strip()
    if not relative:
        return None
    path = Path(relative)
    if not path.is_absolute():
        path = REPO_ROOT / relative
    if not path.is_file():
        return None
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None
    return data if isinstance(data, dict) else None


def _douyin_api_key() -> str:
    _, values = _read_env_file(_secrets_path())
    return _effective_value("DOUYIN_API_KEY", values)


def _safe_douyin_error(error: object) -> str:
    message = str(error or "")
    api_key = _douyin_api_key()
    if api_key:
        message = message.replace(api_key, "[密钥已隐藏]")
    message = re.sub(r"Bearer\s+[^\s,;]+", "Bearer [密钥已隐藏]", message, flags=re.IGNORECASE)
    return re.sub(r"https?://\S+", "[地址已隐藏]", message)[:300]


def _response_json(response: Any) -> Any:
    try:
        return response.json()
    except (AttributeError, TypeError, ValueError):
        return json.loads(str(getattr(response, "text", "") or "{}"))


def _normalize_douyin_api_snapshot(payload: Any, kind: str) -> dict[str, Any]:
    """Reduce vendor JSON to the small, non-evidentiary snapshot contract."""
    node = payload
    for _ in range(3):
        if not isinstance(node, dict):
            break
        nested = next(
            (
                node.get(key)
                for key in ("data", "result")
                if isinstance(node.get(key), (dict, list))
            ),
            None,
        )
        if nested is None:
            break
        node = nested
    if kind == "douyin_board":
        keys = ("items", "list", "word_list", "hot_list")
        output_key = "items"
        allowed = {"word", "query", "title", "hotScore", "heat_value", "url", "share_url"}
    else:
        keys = ("videos", "list", "items", "aweme_list")
        output_key = "videos"
        allowed = {"title", "desc", "caption", "account_name", "play_count", "url", "share_url"}
    rows: Any = node if isinstance(node, list) else None
    if isinstance(node, dict):
        rows = next((node.get(key) for key in keys if isinstance(node.get(key), list)), None)
    if not isinstance(rows, list):
        raise DailyAutomationError("抖音数据接口响应缺少可识别的榜单数组")
    normalized = [
        {key: item.get(key) for key in allowed if key in item}
        for item in rows
        if isinstance(item, dict)
    ]
    return {output_key: normalized, "captured_at": _now()}


def _request_douyin_api_snapshot(
    source: dict[str, Any],
    request_get: Callable[..., Any],
) -> dict[str, Any] | None:
    api_url = str(source.get("api_url") or "").strip()
    api_key = _douyin_api_key()
    if not api_url or not api_key:
        return None
    parsed = urlsplit(api_url)
    if parsed.scheme.lower() != "https" or not parsed.netloc:
        raise DailyAutomationError("抖音数据接口必须使用有效的 HTTPS 地址")
    if re.search(r"(?:key|token|secret)=", parsed.query, re.IGNORECASE):
        raise DailyAutomationError("抖音数据接口密钥不得写入 URL 查询参数")
    response = request_get(
        api_url,
        headers={
            "Accept": "application/json",
            "Authorization": f"Bearer {api_key}",
            "User-Agent": "OpenMontage/1.0 DailyTechBrief",
        },
        timeout=20,
    )
    response.raise_for_status()
    return _normalize_douyin_api_snapshot(_response_json(response), str(source.get("kind") or ""))


def _douyin_board_signals(
    source: dict[str, Any], snapshot: dict[str, Any] | None = None
) -> list[dict[str, Any]]:
    snapshot = snapshot or _read_douyin_snapshot(source)
    if snapshot is None:
        return []
    rows = snapshot.get("items")
    if not isinstance(rows, list):
        return []
    maximum = max(1, min(int(source.get("max_items") or 50), 100))
    signals: list[dict[str, Any]] = []
    for position, item in enumerate(rows[:maximum], 1):
        if not isinstance(item, dict):
            continue
        title = _clean_line_text(item.get("word") or item.get("query") or item.get("title"))
        if not title:
            continue
        try:
            heat_value = int(item.get("hotScore") or item.get("heat_value") or 0)
        except (TypeError, ValueError):
            heat_value = 0
        signals.append(
            {
                "signal_id": f"{str(source.get('id') or 'douyin-hotboard')}-{position:02d}",
                "source_id": str(source.get("id") or "douyin-hotboard"),
                "source_name": str(source.get("name") or "抖音热榜"),
                "title": title[:160],
                "rank": position,
                "heat_value": heat_value,
                "url": _canonical_url(str(item.get("url") or item.get("rawUrl") or "")),
                "captured_at": _now(),
                "scope": "douyin_public_heat_snapshot",
            }
        )
    return signals


def _douyin_creator_signals(
    source: dict[str, Any], snapshot: dict[str, Any] | None = None
) -> list[dict[str, Any]]:
    snapshot = snapshot or _read_douyin_snapshot(source)
    if snapshot is None:
        return []
    rows = snapshot.get("videos")
    if not isinstance(rows, list):
        return []
    maximum = max(1, min(int(source.get("max_items") or 30), 100))
    signals: list[dict[str, Any]] = []
    for position, item in enumerate(rows[:maximum], 1):
        if not isinstance(item, dict):
            continue
        title = _clean_line_text(item.get("title") or item.get("desc") or item.get("caption"))
        if not title:
            continue
        signals.append(
            {
                "signal_id": f"{str(source.get('id') or 'douyin-benchmark-accounts')}-{position:02d}",
                "source_id": str(source.get("id") or "douyin-benchmark-accounts"),
                "source_name": str(item.get("account_name") or source.get("name") or "对标账号"),
                "title": title[:160],
                "rank": position,
                "heat_value": int(item.get("play_count") or 0),
                "url": _canonical_url(str(item.get("url") or item.get("share_url") or "")),
                "captured_at": _now(),
                "scope": "douyin_benchmark_topic",
            }
        )
    return signals


def collect_douyin_sources(
    sources: list[dict[str, Any]],
    request_get: Callable[..., Any] = requests.get,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Collect Douyin hot-board / benchmark-account topic signals (opt-in, non-fatal).

    These are ranking hints only.  They never become article evidence: the
    existing RSS candidates plus the original-site evidence gate still own the
    frozen facts.  A missing snapshot or key yields ``skipped``, never an error.
    """
    signals: list[dict[str, Any]] = []
    results: list[dict[str, Any]] = []
    for source in sources:
        source_id = str(source.get("id") or "douyin-source")
        result = {"id": source_id, "name": str(source.get("name") or source_id), "status": "skipped", "count": 0}
        kind = str(source.get("kind") or "")
        api_url = str(source.get("api_url") or "").strip()
        api_attempted = bool(api_url and _douyin_api_key())
        api_error = ""
        try:
            snapshot: dict[str, Any] | None = None
            if api_attempted:
                try:
                    snapshot = _request_douyin_api_snapshot(source, request_get)
                except Exception as exc:  # noqa: BLE001 - an offline snapshot remains a safe fallback.
                    api_error = _safe_douyin_error(exc)
            mode = "api" if snapshot is not None else ""
            if snapshot is None:
                snapshot = _read_douyin_snapshot(source)
                if snapshot is not None:
                    mode = "snapshot"
            if kind == "douyin_board":
                rows = _douyin_board_signals(source, snapshot)
            elif kind == "douyin_creator":
                rows = _douyin_creator_signals(source, snapshot)
            else:
                raise DailyAutomationError(f"不支持的抖音信号类型：{kind or '[空]'}")
            if rows:
                signals.extend(rows)
                result.update({"status": "ok", "mode": mode, "count": len(rows), "captured_at": _now()})
                if api_error:
                    result["fallback_reason"] = api_error
            elif api_attempted and api_error:
                result.update({"status": "failed", "mode": "api", "error": api_error})
            else:
                result["error"] = "未配置可用接口或离线快照，本轮跳过"
        except Exception as exc:  # noqa: BLE001 - douyin is never fatal.
            result.update({
                "status": "failed" if api_attempted else "skipped",
                "error": _safe_douyin_error(exc),
            })
        results.append(result)
    return signals, results


def collect_news_candidates(
    target: date | str,
    *,
    sources: list[dict[str, Any]] | None = None,
    request_get: Callable[..., Any] = requests.get,
    copy_skill_root: str | Path | None = None,
    copy_skill_enabled: bool | None = None,
    copy_skill_snapshot_dir: str | Path | None = None,
) -> dict[str, Any]:
    target_value = date.fromisoformat(target) if isinstance(target, str) else target
    start, end = target_window(target_value)
    settings = read_config() if sources is None else {}
    configured_sources = sources or settings.get("news_sources") or list(DEFAULT_NEWS_SOURCES)
    candidates: list[dict[str, Any]] = []
    source_results: list[dict[str, Any]] = []
    seen_urls: set[str] = set()
    seen_candidate_ids: set[str] = set()
    for source in configured_sources:
        source_id = str(source.get("id") or "source")
        source_url = str(source.get("url") or "").strip()
        source_url = (
            source_url.replace("{previous_date}", (target_value - timedelta(days=1)).isoformat())
            .replace("{target_date}", target_value.isoformat())
            .replace("{next_date}", (target_value + timedelta(days=1)).isoformat())
        )
        max_candidates = max(1, min(int(source.get("max_candidates") or 100), 100))
        result = {"id": source_id, "name": str(source.get("name") or source_id), "url": source_url, "status": "failed", "count": 0}
        if not source_url:
            result["error"] = "缺少 RSS 地址"
            source_results.append(result)
            continue
        try:
            response = request_get(source_url, timeout=20, headers={"User-Agent": "OpenMontage/1.0 DailyTechBrief"})
            response.raise_for_status()
            root = _parse_feed_root(response.content)
            entries = [item for item in root.iter() if item.tag.rsplit("}", 1)[-1].lower() in {"item", "entry"}]
            for entry in entries:
                published = _parse_published(_first_text(entry, ("published", "updated", "pubDate", "date")))
                if published is None or not (start <= published < end):
                    continue
                title = _strip_markup(_first_text(entry, ("title",)))
                summary = _strip_markup(_first_text(entry, ("summary", "description", "content", "encoded")))
                url = _canonical_url(_entry_url(entry) or _first_text(entry, ("guid", "id")))
                if not title or not url or url in seen_urls:
                    continue
                seen_urls.add(url)
                fingerprint = hashlib.sha256(re.sub(r"\W+", "", title.lower()).encode("utf-8")).hexdigest()[:16]
                candidate_id = f"N-{fingerprint.upper()}"
                if candidate_id in seen_candidate_ids:
                    continue
                seen_candidate_ids.add(candidate_id)
                candidate = {
                    "candidate_id": candidate_id,
                    "title": title[:300],
                    "summary": summary[:1200],
                    "url": url,
                    "published_at": published.isoformat(timespec="seconds"),
                    "source_id": source_id,
                    "source_name": result["name"],
                    "authority": str(source.get("authority") or "media"),
                }
                candidate["china_short_video_hint"] = _china_short_video_hint(candidate)
                candidates.append(candidate)
                result["count"] += 1
                if result["count"] >= max_candidates:
                    break
            result["status"] = "ok"
        except Exception as exc:  # noqa: BLE001 - individual feeds are non-fatal.
            result["error"] = re.sub(r"https?://\S+", "[地址已隐藏]", str(exc))[:300]
        source_results.append(result)
    feed_config = settings.get("copy_skill_hotspot_feed") if isinstance(settings.get("copy_skill_hotspot_feed"), dict) else {}
    feed_enabled = (
        bool(copy_skill_enabled)
        if copy_skill_enabled is not None
        else bool(settings and feed_config.get("enabled") is not False)
    )
    feed_root = copy_skill_root or feed_config.get("root") or DEFAULT_COPY_SKILL_HOTSPOT_ROOT
    snapshot_dir = (
        Path(copy_skill_snapshot_dir)
        if copy_skill_snapshot_dir is not None
        else RUNS_ROOT / target_value.isoformat() / "inputs" / "copy_skill_hotspot"
    )
    if feed_enabled:
        copy_skill_feed = try_load_copy_skill_hotspot_feed(
            feed_root,
            target_value,
            snapshot_dir=snapshot_dir,
        )
    else:
        copy_skill_feed = {
            "schema": "openmontage-copy-skill-hotspot-feed-v1",
            "feed_status": "missing",
            "business_date": target_value.isoformat(),
            "run_id": None,
            "coverage_warning": "copy_skill 热点候选池未启用",
            "manifest_validation": {"valid": False, "error": "disabled"},
            "counts": {"candidates": 0, "raw_records": 0, "target_day_videos": 0, "images": 0},
            "partial_risks": [],
            "candidates": [],
        }
    for candidate in feed_to_discovery_candidates(copy_skill_feed):
        candidate_id = str(candidate.get("candidate_id") or "")
        if not candidate_id or candidate_id in seen_candidate_ids:
            continue
        seen_candidate_ids.add(candidate_id)
        candidates.append(candidate)
    heat_signals, heat_source_results = collect_heat_signals(
        sources=list(settings.get("heat_sources") or []), request_get=request_get,
    ) if settings else ([], [])
    # 抖音热榜/对标账号信号：非致命，只做热度排序补盲，不充当事实证据。
    douyin_sources = (settings.get("douyin_sources") if settings else None) or list(DEFAULT_DOUYIN_SOURCES)
    douyin_signals: list[dict[str, Any]] = []
    douyin_source_results: list[dict[str, Any]] = []
    douyin_signals, douyin_source_results = collect_douyin_sources(douyin_sources, request_get)
    heat_signals.extend(douyin_signals)
    copy_skill_signals = feed_to_heat_signals(copy_skill_feed)
    heat_signals.extend(copy_skill_signals)
    factual_candidates = [item for item in candidates if item.get("discovery_only") is not True]
    source_ids = {str(item.get("source_id") or "") for item in factual_candidates}
    official_count = sum(1 for item in factual_candidates if item.get("authority") == "official")
    quality_warnings: list[str] = []
    if candidates and len(source_ids) < 2:
        quality_warnings.append("候选新闻仅来自一个来源；脚本可以继续生成，但早间审核应重点核对来源集中风险")
    if candidates and official_count == 0:
        quality_warnings.append("本轮窗口内未收集到官方来源候选；脚本必须只复述媒体候选中的可核验事实")
    return {
        "version": "1.0",
        "target_date": target_value.isoformat(),
        "window": {"start": start.isoformat(), "end": end.isoformat()},
        "collected_at": _now(),
        "sources": source_results,
        "heat_sources": heat_source_results,
        "heat_signals": heat_signals,
        "douyin_sources": douyin_source_results,
        "copy_skill_feed": {
            key: copy_skill_feed.get(key)
            for key in (
                "schema",
                "feed_status",
                "package_status",
                "business_date",
                "run_id",
                "source_root",
                "source_day_dir",
                "source_pack_dir",
                "manifest_validation",
                "counts",
                "coverage_warning",
                "partial_risks",
                "snapshot_paths",
                "error",
            )
            if key in copy_skill_feed
        },
        "candidates": candidates,
        "quality": {
            "distinct_source_count": len(source_ids),
            "official_candidate_count": official_count,
            "warnings": quality_warnings,
        },
    }


def _china_short_video_hint(candidate: dict[str, Any]) -> dict[str, Any]:
    """Give the model a transparent relevance hint; do not invent heat data."""
    corpus = " ".join(str(candidate.get(key) or "") for key in ("title", "summary")).lower()
    matches = [brand for brand in CHINA_SHORT_VIDEO_PRIORITY_BRANDS if brand.lower() in corpus]
    domestic = [brand for brand in matches if brand in {"字节", "豆包", "火山", "阿里", "通义", "夸克", "腾讯", "混元", "元宝", "百度", "文心", "华为", "小米", "荣耀", "deepseek"}]
    return {
        "priority_brands": matches[:5],
        "likely_china_relevance": "high" if domestic else ("medium" if matches else "unknown"),
        "editorial_instruction": "优先解释普通中国用户、创作者或职场人会受到什么影响；没有明确影响时不要假称热搜或全民关注。",
    }


def _clean_line_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


TOPIC_SCORE_LIMITS = {
    "sab_combination": 30,
    "public_relevance": 25,
    "event_tension": 20,
    "interaction_potential": 15,
    "reliability_freshness": 10,
}
SCRIPT_SCORE_LIMITS = {
    "hook_strength": 25,
    "dual_host_structure": 20,
    "plain_density": 25,
    "pacing_duration": 15,
    "outro_interaction": 15,
}
TIER_LINE_COUNTS = {"S": 4, "A": 4, "B": 2}
TIER_INFORMATION_DIMENSIONS = {
    "S": ("scene_consequence", "mechanism_or_pattern", "user_impact", "action_tip"),
    "A": ("event", "audience_reaction_or_test", "distinct_detail", "why_it_matters"),
    "B": ("event", "user_value"),
}
COMMON_SPOKEN_ACRONYMS = {"AI", "PC"}
HOST_FUNCTIONS = {
    "yaya": {"hook", "event", "reaction", "user_question", "plain_translate", "quip", "closing", "fact", "impact", "context"},
    "mengmeng": {"fact", "mechanism", "context", "impact", "industry_take", "reference_tip", "summary", "translate", "quip"},
}


def _spoken_english_terms(value: Any) -> set[str]:
    return {
        token
        for token in re.findall(r"(?<![A-Za-z0-9])[A-Za-z][A-Za-z0-9.+-]{2,}(?![A-Za-z0-9])", str(value or ""))
        if token.upper() not in COMMON_SPOKEN_ACRONYMS
    }


def _numeric_score(value: Any, *, maximum: int) -> float:
    try:
        score = float(value)
    except (TypeError, ValueError):
        return -1.0
    return score if 0 <= score <= maximum else -1.0


def validate_topic_selection(selection: dict[str, Any], research: dict[str, Any] | None = None) -> dict[str, Any]:
    issues: list[str] = []
    stories = selection.get("selected_stories") if isinstance(selection.get("selected_stories"), list) else []
    scores = selection.get("topic_scores") if isinstance(selection.get("topic_scores"), dict) else {}
    if len(stories) != 4:
        issues.append("90秒公域快报必须选择 4 条新闻")
    valid_candidate_ids = {
        str(item.get("candidate_id") or "")
        for item in ((research or {}).get("candidates") or [])
        if isinstance(item, dict)
    }
    candidate_by_id = {
        str(item.get("candidate_id") or ""): item
        for item in ((research or {}).get("candidates") or [])
        if isinstance(item, dict)
    }
    tiers: list[str] = []
    categories: list[str] = []
    all_candidate_ids: list[str] = []
    for index, story in enumerate(stories, 1):
        if not isinstance(story, dict):
            issues.append(f"S{index:02d} 不是有效选题对象")
            continue
        tier = str(story.get("tier") or "").upper()
        category = _clean_line_text(story.get("category"))
        terms_to_explain = story.get("terms_to_explain") if isinstance(story.get("terms_to_explain"), list) else []
        candidate_ids = story.get("candidate_ids") if isinstance(story.get("candidate_ids"), list) else []
        tiers.append(tier)
        categories.append(category)
        all_candidate_ids.extend(map(str, candidate_ids))
        if tier not in TIER_LINE_COUNTS:
            issues.append(f"S{index:02d} 必须标注 S/A/B 等级")
        if tier == "B" and not re.search(r"省钱|省时间|更方便|能避坑|使用体验|消费决策", _clean_line_text(story.get("public_value"))):
            issues.append(f"S{index:02d} B级用户价值必须落到省钱、省时间、更方便、避坑、使用体验或消费决策")
        if len(terms_to_explain) > 1:
            issues.append(f"S{index:02d} 需要解释两个以上陌生名词，不适合公域播报")
        public_copy = f"{story.get('headline', '')} {story.get('three_second_summary', '')}"
        if len(_spoken_english_terms(public_copy)) > 1:
            issues.append(f"S{index:02d} 公域标题和摘要不得堆叠两个以上英文专名")
        if index == 1 and tier != "S":
            issues.append("第一条必须是 S 级大众爆点")
        if not candidate_ids:
            issues.append(f"S{index:02d} 未绑定候选来源")
        if valid_candidate_ids and not set(map(str, candidate_ids)) <= valid_candidate_ids:
            issues.append(f"S{index:02d} 引用了冻结候选之外的来源")
        source_corpus = " ".join(
            f"{candidate_by_id.get(str(candidate_id), {}).get('title', '')} "
            f"{candidate_by_id.get(str(candidate_id), {}).get('summary', '')} "
            f"{candidate_by_id.get(str(candidate_id), {}).get('evidence_excerpt', '')}"
            for candidate_id in candidate_ids
        )
        selected_sources = [candidate_by_id.get(str(candidate_id), {}) for candidate_id in candidate_ids]
        if tier in {"S", "A"} and any("evidence_status" in item for item in selected_sources) and not any(
            item.get("evidence_status") == "ok" for item in selected_sources
        ):
            issues.append(f"S{index:02d} {tier}级新闻缺少原站正文证据，不能扩写传播维度")
        if tier in {"S", "A"} and re.search(r"币安|Binance|加密货币|cryptocurrency|crypto trading|比特币", source_corpus, re.IGNORECASE):
            issues.append(f"S{index:02d} 加密交易平台功能不适合作为国内公域 S/A 级核心题")
        extreme_percentages = [int(value) for value in re.findall(r"(?<!\d)(\d{3,})\s*%", source_corpus)]
        has_official_source = any(str(item.get("authority") or "").lower() == "official" for item in selected_sources)
        independent_sources = {
            str(item.get("source_name") or item.get("url") or "") for item in selected_sources if item
        }
        if extreme_percentages and not has_official_source and len(independent_sources) < 2:
            issues.append(f"S{index:02d} 含超过 100% 的异常数字，必须有官方或两条独立来源交叉核验")
        for field, label, minimum in (
            ("headline", "标题", 4),
            ("three_second_summary", "三秒摘要", 8),
            ("why_public_cares", "大众利益点", 10),
            ("public_value", "用户价值锚点", 2),
            ("event_tension", "事件冲突", 8),
            ("comment_hook", "评论抓手", 6),
        ):
            if len(_clean_line_text(story.get(field))) < minimum:
                issues.append(f"S{index:02d} 缺少有效{label}")
    if tiers.count("S") < 1:
        issues.append("选题组合缺少 S 级大众爆点")
    if tiers.count("B") > 1:
        issues.append("B 级行业补充不得超过一条")
    if len(stories) == 4 and sorted(tiers) != ["A", "A", "B", "S"]:
        issues.append("90秒版本必须采用 1S+2A+1B 组合")
    if len(stories) >= 3 and len(set(categories)) < 2:
        issues.append("一期新闻至少覆盖两个内容类别，禁止同类行业事件占满全期")
    if len(all_candidate_ids) != len(set(all_candidate_ids)):
        issues.append("同一候选新闻不得重复入选")
    normalized_scores: dict[str, float] = {}
    for key, maximum in TOPIC_SCORE_LIMITS.items():
        value = _numeric_score(scores.get(key), maximum=maximum)
        normalized_scores[key] = value
        if value < 0:
            issues.append(f"选题评分 {key} 必须在 0—{maximum} 分之间")
    expected_total = sum(value for value in normalized_scores.values() if value >= 0)
    total = _numeric_score(scores.get("total"), maximum=100)
    if total < 0 or abs(total - expected_total) > 0.01:
        issues.append("选题总分必须等于五项分数之和")
    if normalized_scores.get("sab_combination", -1) < 21:
        issues.append("S/A/B 组合评分未达到 21 分")
    if normalized_scores.get("public_relevance", -1) < 15:
        issues.append("大众感知评分未达到 15 分")
    if normalized_scores.get("reliability_freshness") != 10:
        issues.append("事实可靠性与时新性必须满分 10 分")
    if total < 70:
        issues.append(f"选题总分仅 {max(total, 0):g}，低于脚本准入线 70 分")
    if issues:
        raise DailyTopicValidationError(list(dict.fromkeys(issues)))
    return {"valid": True, "story_count": len(stories), "tiers": tiers, "topic_score": round(total, 2)}


def _topic_selection_prompt() -> str:
    return """你是抖音公域科技快报的选题总编。只做选题，不写任何台词。只根据冻结候选输出 JSON，不得补充候选外事实。
输出结构：
{
  "selected_stories":[{"story_id":"S01","candidate_ids":["候选ID"],"headline":"短标题","tier":"S|A|B","category":"民生风险|国产硬核|消费数码|产业事件|国际前沿","three_second_summary":"三秒能懂的一句话","why_public_cares":"和普通人钱包、体验、风险或情绪的关系","public_value":"省钱、省时间、更方便、能避坑或更直观看懂进展中的一项","terms_to_explain":["最多一个必须解释的陌生名词"],"event_tension":"冲突、反差、风险、突破或价格变化","comment_hook":"观众能随手表达的观点"}],
  "combination_reason":"为什么这组题有节奏起伏",
  "topic_scores":{"sab_combination":0,"public_relevance":0,"event_tension":0,"interaction_potential":0,"reliability_freshness":10,"total":0}
}
硬规则：
1. 来源可靠且位于目标 24 小时窗口是准入项；传闻当事实、夸大国产突破或数据没有出处，reliability_freshness=0 并整组淘汰。RSS 的 title/summary 只算标题级线索；候选若有 evidence_excerpt，事实扩写必须以该原站正文证据为准。S/A级必须至少绑定一条 evidence_status=ok 的来源。
   - 涨跌超过100%、数量级突破或“全球第一”等异常强结论，必须绑定官方来源或至少两条独立媒体来源；只有一条聚合标题时不得入选，也不得给可靠性10分。
   - 不得把双来源要求扩大到所有普通新闻。官方来源或可信专业媒体可以单独支撑其标题和摘要中明确披露的常规事件；但只能使用来源已有结论，不得自行扩写受影响范围、解决方案或普遍风险。
2. S级是全民可感知、有情绪、有天然话题的事件；A级是科技圈重磅热点、国产硬核突破、重要数码或商业航天；B级是企业公告、版本迭代、API政策、财报或研报。
   - 高价值候选优先看：国家级技术或标准突破、千亿级战略动作、旗舰模型重大能力/价格变化、新能源与机器人真实交付、消费数码重磅变化。但“层级高”不自动等于“传播强”，仍要通过三秒理解、具体反差和用户利益三项测试。
   - 企业融资要同时具备明确规模、资金用途和可解释的行业影响才可判A；仅有募资公告仍是B。模型降价要讲清适用产品、幅度和期限；机器人从亮相走到开放购买可升A，但预订不等于普及，能力和交付均不得超出来源。
   - 厂商自称“重磅”“升级”“免费”“从找到到做到”，不构成A级证据；没有独立可核验的行业第一、性能跃迁、价格巨变、重大事故或广泛用户影响时，常规产品发布一律判B级。
   - 安全研究只有在来源给出已验证攻击、明确受影响产品和普通用户风险时才可判S级；抽象论文、概念漏洞或“可能有风险”最多判A级。
   - 机器人展会亮相本身不是A级；只有具体高难动作、量产交付、进厂作业、价格反差或事故名场面可升为A/S级。
3. 90秒版本严格采用四条 1S+2A+1B，顺序为S→A→A→B，第一条必须是S级。全B级一票否决。
4. 先做亲友测试：普通人能否三秒听懂？是否会感到厉害、离谱、踩坑或与钱包相关？是否能随手评论？至少两个答案为是才可入选。
   - 一期至少覆盖两个内容类别，不能三条都围绕AI安全、企业政策或厂商发布。A级优先补国产硬核、机器人具体名场面、芯片/航天突破或大众数码事件。
   - 面向中国大陆抖音公域，加密货币交易平台的产品功能、AI代币交易或投机工具不得作为S/A级核心题；除非是重大监管或安全事件，否则直接淘汰。
   - 理解成本也是选题门：一条新闻若必须解释两个以上陌生品牌、英文产品名或行业术语才能讲清，优先淘汰或换题。确无更好B级候选时，必须能完全改写成一个中文产品类别加一个具体用户价值，terms_to_explain 最多一个。
   - 每条新闻都必须填写 public_value；只有事实、没有省钱、省时间、更方便、能避坑或更直观看懂进展等价值锚点时，不得入选。
   - B级不能用“看懂行业进展”冒充用户价值，public_value 必须明确写出省钱、省时间、更方便、能避坑、使用体验或消费决策之一；财报和企业数字若不能落到这些场景，应换成更有体感的B级题。
5. 选题组合按五项评分：S/A/B组合30、大众感知25、事件冲突20、互动潜力15、事实与时新10。先按上述锚点重新判级，再给组合分；不得根据自己想要的组合反向抬级。一个题只有厂商措辞、没有独立事件证据时，事件冲突不得超过10分。普通人需要解释行业背景才能理解时，大众感知不得超过18分。严格据实评分，不得为了过线虚高。
   - 评分必须同时评价“素材潜力”和“可交付证据”：题材听起来很大但正文只能支撑公告标题时，不得用想象补足事件感。选题合规不等于传播优质；没有具体数字反差、生活场景或可争论抓手的组合，总分不得超过79。
6. 选题总分不足70不得写脚本；不足60应停止公域更新。宁可停止，也不拿企业PR凑数。
7. 优先形成“大众爆点→两个不同领域A级热点→B级实用补充”的节奏。S级来源必须足以支撑具体场景/后果、套路或原理、用户影响、行动提示四个不同维度；缺少来源证据时换题或停更，禁止靠常识补写。"""


def _chat_json_with_transient_retry(system: str, payload: dict[str, Any]) -> tuple[dict[str, Any], str]:
    try:
        return _chat_json(system, payload, timeout_seconds=180)
    except TextAIError as first_error:
        if not re.search(r"超时|HTTP 50[234]|连接(?:失败|中断|重置)", str(first_error), re.IGNORECASE):
            raise
        return _chat_json(system, payload, timeout_seconds=180)


def select_daily_topics(research: dict[str, Any], *, max_revision_rounds: int = 2) -> dict[str, Any]:
    candidates = research.get("candidates") if isinstance(research.get("candidates"), list) else []
    if len(candidates) < 3:
        raise DailyAutomationError("指定日期内可核验候选少于 3 条，无法进行公域选题")
    evidence_attempted = _prefetch_selection_evidence(candidates)
    selection_candidates = _topic_selection_payload_candidates(candidates)
    payload: dict[str, Any] = {
        "target_date": research.get("target_date"),
        "candidate_news": selection_candidates,
        "source_evidence_instruction": (
            "优先从 evidence_status=ok 的候选中选择S/A级，并只使用 evidence_excerpt 已明确支持的细节。"
            "evidence_status=failed 或缺少正文证据的候选不得靠标题扩写。"
        ),
    }
    last_issues: list[str] = []
    # A newly chosen story may fall outside the bounded evidence prefetch.
    # Fetching its source text is not a failed rewrite; it must leave one
    # model turn available to reassess that exact selection against the newly
    # frozen evidence.  Keep the cost bounded at one extra selection call.
    attempt_limit = max_revision_rounds + 1
    evidence_extension_available = True
    revision = 0
    while revision < attempt_limit:
        if last_issues:
            payload["validation_errors_to_fix"] = last_issues
            payload["revision_round"] = revision
        raw, model = _chat_json_with_transient_retry(_topic_selection_prompt(), payload)
        shortlisted_ids = {
            str(candidate_id)
            for story in raw.get("selected_stories") or []
            if isinstance(story, dict)
            for candidate_id in story.get("candidate_ids") or []
        }
        pending_evidence = shortlisted_ids - evidence_attempted
        if pending_evidence:
            for candidate in candidates:
                candidate_id = str(candidate.get("candidate_id") or "")
                if candidate_id in pending_evidence:
                    _enrich_candidate_evidence(candidate)
                    evidence_attempted.add(candidate_id)
            # The list keeps references to the frozen candidates, so newly
            # fetched evidence is visible without expanding the model packet.
            payload["candidate_news"] = selection_candidates
            payload["source_evidence_instruction"] = (
                "本轮入围项已补充 evidence_excerpt。必须重新选题和评分；只有 evidence_status=ok 的来源才能支撑"
                "S级四个信息维度及A级具体细节。不得把RSS标题重复当正文证据。"
            )
            last_issues = ["请使用刚补充的原站正文证据重新完成选题，证据不足的题必须替换或降级"]
            if revision == attempt_limit - 1 and evidence_extension_available:
                attempt_limit += 1
                evidence_extension_available = False
            revision += 1
            continue
        try:
            validation = validate_topic_selection(raw, research)
        except DailyTopicValidationError as exc:
            last_issues = exc.issues
            revision += 1
            continue
        raw.update({"version": "1.0", "target_date": str(research.get("target_date") or ""), "generated_at": _now(), "model": model, "validation": validation})
        return raw
    raise DailyTopicValidationError(last_issues or ["选题在两轮修订后仍未达到公域准入线"])


def validate_daily_script(script: dict[str, Any], selection: dict[str, Any] | None = None) -> dict[str, Any]:
    issues: list[str] = []
    stories = script.get("stories") if isinstance(script.get("stories"), list) else []
    lines = script.get("lines") if isinstance(script.get("lines"), list) else []
    quality = script.get("script_scores") if isinstance(script.get("script_scores"), dict) else {}
    editorial = script.get("editorial") if isinstance(script.get("editorial"), dict) else {}
    if len(stories) != 4:
        issues.append("90秒脚本必须包含 4 条新闻")
    tiers: list[str] = []
    story_ids: list[str] = []
    story_glosses: dict[str, dict[str, str]] = {}
    story_gloss_items: dict[str, list[tuple[str, str]]] = {}
    story_event_identities: dict[str, str] = {}
    story_tiers: dict[str, str] = {}
    expected_roles: list[str] = []
    expected_story_ids: list[str] = []
    expected_dimensions: list[str] = []
    for index, story in enumerate(stories, 1):
        if not isinstance(story, dict):
            issues.append(f"S{index:02d} 不是有效新闻对象")
            continue
        story_id = str(story.get("story_id") or "")
        tier = str(story.get("tier") or "").upper()
        count = story.get("line_count")
        source_ids = story.get("source_ids") if isinstance(story.get("source_ids"), list) else []
        gloss_rows = story.get("foreign_term_glosses") if isinstance(story.get("foreign_term_glosses"), list) else []
        glosses: dict[str, str] = {}
        for gloss in gloss_rows:
            if not isinstance(gloss, dict):
                continue
            term = _clean_line_text(gloss.get("term"))
            chinese_label = _clean_line_text(gloss.get("chinese_label"))
            if term and len(chinese_label) >= 2:
                glosses[term.lower()] = chinese_label
        story_glosses[story_id] = glosses
        story_gloss_items[story_id] = list(glosses.items())
        event_identity = _clean_line_text(story.get("event_identity"))
        story_event_identities[story_id] = event_identity
        story_tiers[story_id] = tier
        story_ids.append(story_id)
        tiers.append(tier)
        if story_id != f"S{index:02d}":
            issues.append(f"第 {index} 条新闻编号应为 S{index:02d}")
        if count != TIER_LINE_COUNTS.get(tier):
            issues.append(f"{story_id or f'S{index:02d}'} 必须按 {tier or '?'} 级分配 {TIER_LINE_COUNTS.get(tier, '?')} 句")
            count = 0
        if not source_ids:
            issues.append(f"{story_id or f'S{index:02d}'} 未绑定来源")
        if len(_clean_line_text(story.get("plain_summary"))) < 10:
            issues.append(f"{story_id or f'S{index:02d}'} 缺少大白话摘要")
        if len(_clean_line_text(story.get("why_viewers_care"))) < 10:
            issues.append(f"{story_id or f'S{index:02d}'} 缺少普通人利益关联")
        if len(event_identity) < 2:
            issues.append(f"{story_id or f'S{index:02d}'} 必须声明具体事件身份 event_identity")
        expected_roles.extend(("yaya", "mengmeng", "yaya", "mengmeng")[:count])
        expected_story_ids.extend([story_id] * count)
        expected_dimensions.extend(TIER_INFORMATION_DIMENSIONS.get(tier, ())[:count])
    expected_roles.append("yaya")
    expected_story_ids.append("")
    expected_dimensions.append("")
    expected_count = sum(TIER_LINE_COUNTS.get(tier, 0) for tier in tiers) + 1
    if len(lines) != expected_count:
        issues.append(f"弹性总句数应为 {expected_count}，实际为 {len(lines)}")
    selected = (selection or {}).get("selected_stories") if isinstance((selection or {}).get("selected_stories"), list) else []
    if selected:
        expected_tiers = [str(item.get("tier") or "").upper() for item in selected]
        if tiers != expected_tiers:
            issues.append("脚本新闻等级与已通过选题不一致")
        expected_ids = [str(item.get("story_id") or "") for item in selected]
        if story_ids != expected_ids:
            issues.append("脚本新闻顺序与已通过选题不一致")
    authority_chars = {"yaya": 0, "mengmeng": 0}
    host_functions: dict[str, set[str]] = {"yaya": set(), "mengmeng": set()}
    yaya_questions = 0
    information_keys: dict[str, set[str]] = {story_id: set() for story_id in story_ids}
    story_functions: dict[str, set[str]] = {story_id: set() for story_id in story_ids}
    story_line_texts: dict[str, list[str]] = {story_id: [] for story_id in story_ids}
    for index, line in enumerate(lines, 1):
        expected_id = f"T{index:03d}"
        if not isinstance(line, dict):
            issues.append(f"{expected_id} 不是有效台词对象")
            continue
        if str(line.get("turn_id") or "").upper() != expected_id:
            issues.append(f"第 {index} 句编号应为 {expected_id}")
        speaker = str(line.get("speaker_id") or "").lower()
        if index <= len(expected_roles) and speaker != expected_roles[index - 1]:
            issues.append(f"{expected_id} 说话人顺序错误，应为 {expected_roles[index - 1]}")
        story_id = str(line.get("story_id") or "")
        if index <= len(expected_story_ids) and story_id != expected_story_ids[index - 1]:
            issues.append(f"{expected_id} 未按选题顺序编排")
        text = _clean_line_text(line.get("text"))
        function = _clean_line_text(line.get("function"))
        information_key = _clean_line_text(line.get("information_key"))
        information_dimension = _clean_line_text(line.get("information_dimension"))
        reply_to = _clean_line_text(line.get("reply_to"))
        if index <= len(expected_dimensions) and information_dimension != expected_dimensions[index - 1]:
            issues.append(f"{expected_id} 信息维度应为 {expected_dimensions[index - 1] or '空'}")
        kind = str(line.get("kind") or "story")
        source_ids = line.get("source_ids") if isinstance(line.get("source_ids"), list) else []
        if not text:
            issues.append(f"{expected_id} 台词为空")
        if re.search(r"内容提示词|待补充|某公司|TODO|TBD", text, re.IGNORECASE):
            issues.append(f"{expected_id} 含占位内容")
        if speaker in authority_chars:
            authority_chars[speaker] += len(re.sub(r"\s+", "", text))
        if index == 1:
            if kind != "hook" or function != "hook" or not source_ids:
                issues.append("T001 必须是绑定 S 级来源的雅雅前三秒钩子")
            if not re.search(r"每日科技快讯|科技圈快讯|科技快报|科技新闻快报", text):
                issues.append("T001 必须在同一句带出栏目身份“每日科技快讯”或“科技快报”")
            if text.strip("。！？!? ") == INTRO_TEXT.strip("。！？!? "):
                issues.append("前三秒不能只有固定问候，必须同时抛出当期爆点")
            if re.search(r"^(?:API|Agent|代理系统|安全护栏|加密(?:恶意)?指令)", text, re.IGNORECASE):
                issues.append("前三秒不能用行业术语起句，必须先说普通人的具体后果")
            if not re.search(r"你|手机|文件|聊天|隐私|钱包|转账|价格|工作|家人|模型|调用", text):
                issues.append("前三秒必须出现普通人能立刻感知的对象或利益词")
            if not re.search(r"借|租|传|上传|点开|转账|聊天|购买|交给|录音|拉群|加好友|兼职|写代码|做方案|调用|付费|下单", text) or not re.search(
                r"被骗|损失|泄露|封号|担责|多花|少花|省下|降价|涨价|偷走|导出|违法|永久限制|法律责任|催.{0,4}转账", text
            ):
                issues.append("前三秒必须同时包含具体使用场景和明确后果或利益，不能只说泛化风险")
        elif index == len(lines):
            if kind != "outro" or speaker != "yaya":
                issues.append("最后一句必须由雅雅做互动收束")
            if not re.search(r"[？?]", text) or "评论区" not in text:
                issues.append("结尾必须包含低门槛具体问题和评论区邀请")
            outro_mentions = editorial.get("outro_mentions") if isinstance(editorial.get("outro_mentions"), list) else []
            mention_ids: list[str] = []
            for mention in outro_mentions:
                if not isinstance(mention, dict):
                    continue
                mention_id = str(mention.get("story_id") or "")
                phrase = _clean_line_text(mention.get("phrase"))
                body_text = " ".join(
                    _clean_line_text(item.get("text"))
                    for item in lines[:-1]
                    if isinstance(item, dict) and str(item.get("story_id") or "") == mention_id
                )
                if mention_id in story_ids and len(phrase) >= 2 and phrase in text and phrase in body_text:
                    mention_ids.append(mention_id)
            if len(set(mention_ids)) < 2:
                issues.append("结尾必须在台词中真实带出至少两条当期新闻的具体短语")
            if re.search(r"该先选谁|谁更重要|必须二选一|只能选一个", text):
                issues.append("结尾不得把无关或不可比较概念硬凑成二选一")
            if re.search(r"会不会先核实|要不要小心|是否应该保护|敢不敢借", text):
                issues.append("结尾不得提出只有标准答案的问题")
        else:
            if kind != "story" or not source_ids:
                issues.append(f"{expected_id} 必须是绑定来源的新闻台词")
        if story_id:
            story_line_texts.setdefault(story_id, []).append(text)
            if len(information_key) < 2:
                issues.append(f"{expected_id} 必须声明独立的信息增量键 information_key")
            elif information_key in information_keys.setdefault(story_id, set()):
                issues.append(f"{story_id} 内 information_key 重复，存在原地复述风险")
            else:
                information_keys.setdefault(story_id, set()).add(information_key)
            story_functions.setdefault(story_id, set()).add(function)
        if speaker in HOST_FUNCTIONS:
            if function not in HOST_FUNCTIONS[speaker]:
                issues.append(f"{expected_id} {('雅雅' if speaker == 'yaya' else '檬檬')}台词功能无效")
            else:
                host_functions[speaker].add(function)
        if re.search(r"值得关注|拭目以待|玩家能自己掂量|大家心里有数|你怎么看|未来可期|科技圈炸锅|肉眼可见", text):
            issues.append(f"{expected_id} 是空泛广播腔，必须改成具体事实、追问、解释或克制吐槽")
        if re.search(r"好家伙|撒胡椒面|牌桌加码|把筹码推上桌|掀桌", text):
            issues.append(f"{expected_id} 语气过痞，不符合明快但克制的栏目口吻")
        if re.search(r"做饭|收拾|陪护", text) and not re.search(r"未来|以后|有望|可能|如果|要是|真能|期待", text):
            issues.append(f"{expected_id} 未落地能力只能作为明确展望，必须带未来、可能、如果或真能等标记")
        if re.search(r"放弃其他赛道|破釜沉舟|成本腰斩|全世界.{0,6}(?:必须|都得)|全球已经.{0,6}(?:照办|执行)", text) and not re.search(
            r"不是|并非|不等于|不能说|别喊|别急着喊", text
        ):
            issues.append(f"{expected_id} 把传播性形容写成了无来源事实结论")
        if speaker == "yaya":
            if function == "user_question":
                yaya_questions += 1
                if not re.search(r"[？?]", text):
                    issues.append(f"{expected_id} user_question 必须是真实口语问句")
                if index >= len(lines) or str(lines[index].get("speaker_id") or "").lower() != "mengmeng":
                    issues.append(f"{expected_id} 雅雅的追问必须由下一句檬檬直接回答")
            if reply_to:
                issues.append(f"{expected_id} 雅雅不填写 reply_to")
        elif speaker == "mengmeng":
            if re.search(r"发布|公布|宣布|首次|数据显示|达到|突破|同比|环比|\d", text) and function not in {"fact", "mechanism", "context"}:
                issues.append(f"{expected_id} 檬檬播报核心事实时必须标注 fact、mechanism 或 context，不能伪装成接梗")
            previous_id = f"T{index - 1:03d}"
            if reply_to != previous_id:
                issues.append(f"{expected_id} 必须通过 reply_to 明确承接紧邻的 {previous_id}")
        char_count = len(re.sub(r"\s+", "", text))
        maximum = 42
        minimum = 22 if kind in {"hook", "outro"} else 16
        if text and not minimum <= char_count <= maximum:
            issues.append(f"{expected_id} 台词长度应为 {minimum}—{maximum} 个汉字")
    if len(host_functions["yaya"]) < 3 or len(host_functions["mengmeng"]) < 3:
        issues.append("两位主播整期都必须承担至少三种有效功能，不能退化成播报员和捧哏")
    if yaya_questions > 3:
        issues.append("雅雅正文追问不得超过三次，避免整期变成机械问答")
    outro_story_ids = editorial.get("outro_story_ids") if isinstance(editorial.get("outro_story_ids"), list) else []
    outro_mentions = editorial.get("outro_mentions") if isinstance(editorial.get("outro_mentions"), list) else []
    if len(outro_story_ids) != 2 or len(outro_mentions) != 2:
        issues.append("结尾必须只选择两条能自然比较的新闻，禁止串联三条以上硬凑互动")
    mention_story_ids = [str(item.get("story_id") or "") for item in outro_mentions if isinstance(item, dict)]
    if list(map(str, outro_story_ids)) != mention_story_ids:
        issues.append("outro_story_ids 必须与两条 outro_mentions 按顺序完全一致")
    for story_id in story_ids:
        story_texts = story_line_texts.get(story_id, [])
        identity_line_limit = 2 if story_tiers.get(story_id) == "S" else 1
        identity_window = " ".join(story_texts[:identity_line_limit])
        if story_event_identities.get(story_id) not in identity_window:
            position_label = "前两句" if identity_line_limit == 2 else "第一句"
            issues.append(f"{story_id} {position_label}必须明确说出具体事件身份“{story_event_identities.get(story_id)}”")
        joined_story_text = " ".join(story_texts)
        english_terms = _spoken_english_terms(joined_story_text)
        remaining_terms = {term.lower(): term for term in english_terms}
        grouped_identity_count = 0
        for grouped_term, chinese_label in story_gloss_items.get(story_id, []):
            grouped_tokens = _spoken_english_terms(grouped_term)
            if len(grouped_tokens) <= 1 or grouped_term.lower() not in joined_story_text.lower():
                continue
            grouped_identity_count += 1
            for grouped_token in grouped_tokens:
                remaining_terms.pop(grouped_token.lower(), None)
            first_group_line = next((text for text in story_texts if grouped_term.lower() in text.lower()), "")
            if chinese_label not in first_group_line:
                issues.append(f"{story_id} 组合专名 {grouped_term} 首次出现时必须附中文通俗名称")
        if grouped_identity_count + len(remaining_terms) > 1:
            issues.append(f"{story_id} 台词不得堆叠两个以上英文专名")
        for term_lower, term in remaining_terms.items():
            chinese_label = story_glosses.get(story_id, {}).get(term_lower)
            first_line = next((text for text in story_texts if re.search(rf"(?i)(?<![A-Za-z0-9]){re.escape(term)}(?![A-Za-z0-9])", text)), "")
            if not chinese_label or chinese_label not in first_line:
                issues.append(f"{story_id} 英文专名 {term} 首次出现时必须附中文通俗名称")
    total_chars = authority_chars["yaya"] + authority_chars["mengmeng"]
    ratio = authority_chars["yaya"] / total_chars if total_chars else 0.0
    if total_chars and not 0.30 <= ratio <= 0.70:
        issues.append(f"任一主播都不得低于总台词量的30%，当前雅雅为 {ratio:.1%}")
    normalized_scores: dict[str, float] = {}
    for key, maximum in SCRIPT_SCORE_LIMITS.items():
        value = _numeric_score(quality.get(key), maximum=maximum)
        normalized_scores[key] = value
        if value < 0:
            issues.append(f"脚本评分 {key} 必须在 0—{maximum} 分之间")
    script_total = _numeric_score(quality.get("total"), maximum=100)
    expected_total = sum(value for value in normalized_scores.values() if value >= 0)
    if script_total < 0 or abs(script_total - expected_total) > 0.01:
        issues.append("脚本总分必须等于五项分数之和")
    for key, minimum in {"hook_strength": 20, "dual_host_structure": 16, "plain_density": 18, "pacing_duration": 12, "outro_interaction": 10}.items():
        if normalized_scores.get(key, -1) < minimum:
            issues.append(f"脚本评分 {key} 未达强制线 {minimum}")
    if script_total < 80:
        issues.append("脚本总分低于交付线 80 分")
    try:
        duration = float(script.get("estimated_duration_seconds"))
    except (TypeError, ValueError):
        duration = -1
    duration_range = (85, 95)
    if not duration_range[0] <= duration <= duration_range[1]:
        issues.append(f"预计时长必须在 {duration_range[0]}—{duration_range[1]} 秒")
    if not 365 <= total_chars <= 425:
        issues.append(f"90秒版本总口播字数必须在 365—425 字，当前为 {total_chars}")
    selection_score = float(((selection or {}).get("validation") or {}).get("topic_score") or 0)
    combined_score = selection_score * 0.6 + max(script_total, 0) * 0.4
    if selection and combined_score < 78:
        issues.append(f"综合得分 {combined_score:.1f} 低于发布线 78 分")
    if issues:
        raise DailyScriptValidationError(list(dict.fromkeys(issues)))
    return {
        "valid": True,
        "story_count": len(stories),
        "tiers": tiers,
        "line_count": len(lines),
        "expected_line_count": expected_count,
        "speaker_character_counts": authority_chars,
        "yaya_ratio": round(ratio, 4),
        "script_score": round(script_total, 2),
        "selection_score": round(selection_score, 2) if selection else None,
        "combined_score": round(combined_score, 2) if selection else None,
        "estimated_duration_seconds": duration,
    }


def _script_generation_prompt() -> str:
    return """你是抖音双主持科技快报的脚本主编。选题已通过独立评分，禁止换题，只根据 selected_stories 和 selected_candidates 写稿。只输出 JSON。
输出结构：
{
  "title":"YYYY-MM-DD 科技快报",
  "editorial":{"daily_theme":"主题","audience_promise":"普通人看完得到什么","outro_story_ids":["S02","S03"],"outro_mentions":[{"story_id":"S02","phrase":"必须逐字出现在结尾的短语"},{"story_id":"S03","phrase":"必须逐字出现在结尾的短语"}]},
  "stories":[{"story_id":"S01","tier":"S|A|B","headline":"标题","event_identity":"S级前两句、A/B级第一句必须逐字说出的公司+产品/模型或具体事件名","plain_summary":"大白话事实","why_viewers_care":"普通人利益点","foreign_term_glosses":[{"term":"正文保留的单个英文专名，或公司+型号组成的一个组合专名","chinese_label":"首次出现时逐字带出的中文通俗名称"}],"line_count":4,"source_ids":["候选ID"]}],
  "estimated_duration_seconds":90,
  "script_scores":{"hook_strength":0,"dual_host_structure":0,"plain_density":0,"pacing_duration":0,"outro_interaction":0,"total":0},
  "lines":[{"turn_id":"T001","speaker_id":"yaya|mengmeng","speaker_name":"雅雅|檬檬","kind":"hook|story|outro","function":"hook|event|reaction|user_question|plain_translate|quip|fact|mechanism|context|impact|industry_take|reference_tip|summary|closing","information_dimension":"scene_consequence|mechanism_or_pattern|user_impact|action_tip|event|audience_reaction_or_test|distinct_detail|why_it_matters|user_value或空","information_key":"本句独有的新增信息标签，结尾可为空","reply_to":"檬檬填写紧邻上一句turn_id，雅雅留空","story_id":"S01或空","source_ids":["候选ID"],"text":"台词"}]
}
硬规则：
1. 直接用中文构思并输出中文台词，禁止先生成英文稿再翻译。T001既是栏目开场、前三秒钩子，也是S01四句中的第一句，不额外增加开头。第一句优先以“每日科技快讯来了”带出栏目身份，先说大众能理解的利益、风险或反差；S级的公司、产品和型号允许放在紧邻的第二句补齐，禁止第一句为了报型号挤掉钩子。
2. 90秒版本固定四条并严格全程交替：S级4句1212、第一条A级4句1212、第二条A级4句1212、B级2句12、雅雅结尾1句，共15句。任何相邻两句不得同一角色，跨新闻也不例外；顺序固定S→A→A→B。
3. 技术角色名保持不变，但人设锁定：雅雅是生活化用户视角，负责钩子、把新闻抛到具体生活场景、自然追问、接梗和结尾；檬檬是理性技术党，负责硬事实、底层逻辑、行业影响和克制吐槽。两人都可以说事实，但不得互相复述；雅雅正文追问最多三次，每次必须由下一句檬檬直接回答。檬檬每句 reply_to 填紧邻上一句，并接住上一句的具体对象或动作。任一主播字数不得低于30%，不再追求机械7:3。禁止“没错—我懂了”式流水问答，也禁止为了句句有梗而塞网络热词。
4. information_dimension 必须严格按等级逐句填写且不可重复，每句还要填写独立 information_key，同一 story_id 下禁止重复：S级依次是 scene_consequence（具体场景和后果）、mechanism_or_pattern（套路或原理）、user_impact（影响）、action_tip（行动提示）；每条A级依次是 event（事件）、audience_reaction_or_test（观众反应或检验点）、distinct_detail（独有细节）、why_it_matters（为什么值得看）；B级依次是 event（事件）、user_value（用户价值）。换句式不算新维度，来源不足以支撑某维度时必须停稿，严禁补写。
   - 每条 stories 必须填写 event_identity。S级必须在前两句内逐字说出它，允许第一句先吸引人、第二句再落具体公司和型号；A/B级必须在第一句说出。event_identity 应是“公司+产品/模型”或具体事件名，禁止只写“顶级AI、某机器人、国产突破、大厂动作”等泛称。公司与型号构成的一个完整组合专名按一个事件身份计算，不按多个无关英文名处罚，但同句仍需带“旗舰模型、云游戏服务”等中文类别说明。
5. 所有事实、数字、产品名、行动建议和收益描述必须来自 selected_candidates 并绑定 source_ids。不得把“支持部分产品”扩大成“全部可用”，不得补写来源没有的解决方案、适用范围或效果；数字必须给出参照或说明它为何重要。没有来源或AB实验数据时，禁止编写百分比涨幅和量化效果承诺。
   - 事实层严谨，表达层可以有力度：降价超过两成可概括为“大幅降低”；净募资全部投AI可说“大力支持AI、重押AI”；牵头国际标准可说“提高国际话语权”。但不得写成“成本腰斩、放弃其他赛道、全世界已经必须照办”等超出来源的既成结论。
   - 展望层允许合理想象：机器人未来做饭、收拾、陪护等可作为愿景、吐槽或问题，但必须带“未来、可能、如果、真能、期待”等明确展望词，不能写成已有能力或承诺具体实现时间。
   - 强结论必须保留证据姿态：公司公告写“公司称/官方公告显示”，标准立项写清只是进入制定阶段，未来判断使用“可能、取决于、还要看”，不得把判断写成既成事实。
   - 每条新闻口播最多保留一个英文专名，通用缩写AI、PC不计。能译成中文时优先使用中文，如“英伟达云游戏”“火狐浏览器”；若必须保留Grok等英文名，首次出现的同一句必须带中文通俗名称，并在 foreign_term_glosses 中逐字登记。
   - 每条新闻必须通过规定的 information_dimension 带出用户价值；S级的 user_impact/action_tip、A级的 audience_reaction_or_test/why_it_matters、B级的 user_value 都不能写成空泛评价。檬檬可用 reference_tip 给提示，但不得为凑功能编造来源外建议。
6. 语言适合90秒数字人口播：普通台词16—42字，钩子和结尾22—42字，总口播365—425字，预计85—95秒；不用行业黑话和书面翻译腔。整体模仿明快的双主持对话节奏，优先使用“没错、确实、精准、这也太狠了吧、太提气了、这下有意思了”等带主持感的情绪回应；每条新闻建议一至两处，且同一句必须继续提供事实、解释或用户影响，不能只喊口号。禁止“好家伙、撒胡椒面、牌桌加码、把筹码推上桌”等偏痞表达，也禁止连续使用“没错—我懂了”。每条新闻至少出现一次自然的反差、生活类比或具体追问，但不是每句强行玩梗。输出JSON前逐句计数，越界先改写。
7. 结尾必须只选择两条有共同比较维度的当期新闻，outro_story_ids 和 outro_mentions 都必须恰好两项且顺序一致；优先比较两个产品、两种进展或两种使用习惯，禁止把反诈提醒与游戏、机器人、办公工具三件事串成一道题。问题要有行为分享或真实站队空间，包含问号和“欢迎在评论区讨论”。禁止“会不会先核实、要不要小心”等只有标准答案的问题，也禁止无关假二选一。outro_mentions 每个短语必须先在正文出现，再逐字出现在结尾；结尾 kind=outro、story_id、information_dimension和reply_to均为空。
8. 按五项据实自评：前三秒钩子25、双人结构20、通俗与密度25、节奏时长15、结尾互动15；总分低于80不得交付，不得虚高。结构合规只代表没有犯错，不能因此给高分；若台词像公告摘要、没有具体细节、两人只是“提问—复述”，或删掉任一主播后几乎不影响理解和情绪，脚本总分不得超过79。"""


def generate_daily_script(
    research: dict[str, Any],
    *,
    selection: dict[str, Any] | None = None,
    max_revision_rounds: int = 2,
    editorial_feedback: list[str] | None = None,
    previous_script: dict[str, Any] | None = None,
) -> dict[str, Any]:
    candidates = research.get("candidates") if isinstance(research.get("candidates"), list) else []
    if len(candidates) < 3:
        raise DailyAutomationError("指定日期内可核验候选少于 3 条，已停止在付费阶段之前")
    selection = selection or select_daily_topics(research, max_revision_rounds=max_revision_rounds)
    selected_ids = {
        str(candidate_id)
        for story in selection.get("selected_stories") or []
        if isinstance(story, dict)
        for candidate_id in story.get("candidate_ids") or []
    }
    selected_candidates: list[dict[str, Any]] = []
    for item in candidates:
        if str(item.get("candidate_id") or "") not in selected_ids:
            continue
        compact = dict(item)
        evidence = str(compact.get("evidence_excerpt") or "")
        if len(evidence) > 3600:
            compact["evidence_excerpt"] = f"{evidence[:2500]}\n[正文中段已压缩]\n{evidence[-1100:]}"
        selected_candidates.append(compact)
    payload: dict[str, Any] = {
        "target_date": research.get("target_date"),
        "selected_stories": selection.get("selected_stories"),
        "topic_scores": selection.get("topic_scores"),
        "selected_candidates": selected_candidates,
    }
    if editorial_feedback:
        payload["codex_editorial_feedback"] = [str(item) for item in editorial_feedback]
    if previous_script:
        payload["previous_rejected_script"] = previous_script
    last_issues: list[str] = []
    for revision in range(max_revision_rounds + 1):
        if last_issues:
            payload["validation_errors_to_fix"] = last_issues
            payload["revision_round"] = revision
        raw, model = _chat_json_with_transient_retry(_script_generation_prompt(), payload)
        try:
            validation = validate_daily_script(raw, selection)
        except DailyScriptValidationError as exc:
            last_issues = exc.issues
            continue
        raw.update({
            "version": "2.0",
            "target_date": str(research.get("target_date") or ""),
            "generated_at": _now(),
            "model": model,
            "topic_selection": selection,
            "validation": validation,
        })
        return raw
    raise DailyScriptValidationError(
        last_issues or ["脚本在两轮修订后仍未达到抖音交付线"],
        candidate=raw if isinstance(raw, dict) else None,
    )


def revise_daily_script_dialogue(
    script: dict[str, Any],
    selection: dict[str, Any],
    *,
    editorial_feedback: list[str],
    max_revision_rounds: int = 2,
) -> dict[str, Any]:
    """Let the project model polish dialogue while freezing selected facts."""
    prompt = """你是抖音双主持科技快报的对白编辑。输入稿的选题、事实、数字、来源、顺序、分级和信息维度已经冻结；只允许改写台词口语感、接话感和结尾问法，不得新增、删除或改变任何事实。
只输出与输入稿完全同结构的JSON。保留15句、turn_id、speaker_id、story_id、source_ids、information_dimension和新闻顺序；檬檬reply_to必须指向紧邻上一句。普通台词16—42字，钩子和结尾22—42字，总口播365—425字，预计85—95秒，任一主播不得低于30%。
目标是熟人搭档自然接话：雅雅负责生活视角、具体追问和接梗，檬檬负责硬事实、底层逻辑和克制吐槽。优先用“没错、确实、精准、这也太狠了吧、太提气了”等明快主持语气，每个情绪短语后必须继续给信息；禁用“好家伙、撒胡椒面、牌桌加码、把筹码推上桌”等偏痞表达。允许把两成以上降价概括为“大幅降低”、把明确AI投入说成“大力支持AI”、把牵头标准说成“提高国际话语权”；机器人做饭、收拾、陪护只能写成带“未来、可能、如果、真能”的展望。第一句把栏目身份和爆点合并，结尾只保留两条可比较新闻。
模型自评分不得因结构合规自动给高分；没有明显口语改善时总分不得超过84。"""
    payload: dict[str, Any] = {
        "frozen_selection": selection.get("selected_stories"),
        "rejected_script": script,
        "codex_editorial_feedback": [str(item) for item in editorial_feedback],
    }
    last_issues: list[str] = []
    raw: dict[str, Any] = {}
    model = ""
    for revision in range(max_revision_rounds + 1):
        if last_issues:
            payload["validation_errors_to_fix"] = last_issues
            payload["revision_round"] = revision
        raw, model = _chat_json_with_transient_retry(prompt, payload)
        try:
            validation = validate_daily_script(raw, selection)
        except DailyScriptValidationError as exc:
            last_issues = exc.issues
            payload["rejected_script"] = raw
            continue
        raw.update({
            "version": "2.1-dialogue-revision",
            "target_date": str(script.get("target_date") or selection.get("target_date") or ""),
            "generated_at": _now(),
            "model": model,
            "validation": validation,
        })
        return raw
    raise DailyScriptValidationError(
        last_issues or ["对白返工后仍未达到抖音交付线"],
        candidate=raw if isinstance(raw, dict) else None,
    )


def script_to_import_text(script: dict[str, Any]) -> str:
    lines = [str(script.get("title") or "每日科技快报")]
    for row in script.get("lines") or []:
        lines.append(f"{row['turn_id']} {row['speaker_name']}：{_clean_line_text(row['text'])}")
    return "\n".join(lines).strip() + "\n"


def speaker_pure_text(script: dict[str, Any], speaker_id: str) -> str:
    return "\n".join(
        _clean_line_text(row.get("text"))
        for row in script.get("lines") or []
        if isinstance(row, dict) and row.get("speaker_id") == speaker_id
    ).strip()


def export_script_docx(script: dict[str, Any], target: Path) -> Path:
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(str(script.get("title") or "每日科技快报"), level=0)
    document.add_heading("顺序编号版", level=1)
    for row in script.get("lines") or []:
        document.add_paragraph(f"{row['turn_id']} {row['speaker_name']}：{_clean_line_text(row['text'])}")
    document.add_page_break()
    document.add_heading("纯净版脚本", level=1)
    for speaker_id, speaker_name in (("yaya", "雅雅"), ("mengmeng", "檬檬")):
        document.add_heading(f"{speaker_name}台词", level=2)
        for row in script.get("lines") or []:
            if isinstance(row, dict) and row.get("speaker_id") == speaker_id:
                document.add_paragraph(_clean_line_text(row.get("text")))
    temporary = target.with_suffix(target.suffix + ".tmp")
    document.save(temporary)
    os.replace(temporary, target)
    return target


def create_daily_project(target: date | str, script: dict[str, Any]) -> Path:
    from backlot.script_imports import stage_text_preview
    from backlot.workbench import apply_daily_story_contract, import_avatar_user_script

    target_value = date.fromisoformat(target) if isinstance(target, str) else target
    project_id = f"daily-tech-{target_value.strftime('%Y%m%d')}"
    project_dir = PROJECTS_DIR / project_id
    if not project_dir.exists():
        project_dir = init_project(
            project_id,
            title=f"{target_value.isoformat()} 科技快报",
            pipeline_type="avatar-spokesperson",
            pipeline_dir=PROJECTS_DIR,
            style_playbook="clean-professional",
        )
        marker_path = project_dir / "project.json"
        marker = _read_json(marker_path) or {}
        marker["render_profile"] = {
            "aspect_ratio": "portrait", "width": 1080, "height": 1920, "fps": 30, "audio_sample_rate": 48000,
        }
        marker["intake"] = {
            "brief": "每日自动科技快报；早间人工审核后发布。",
            "duration_seconds": 120,
            "aspect": "portrait",
            "aspect_label": "竖版 9:16",
            "created_from": "daily_automation",
            "avatar": {
                "source_status": "planned", "import_mode": "longform", "default_treatment": "pip_top_left", "background_mode": "opaque",
            },
        }
        _atomic_json(marker_path, marker)
    if not (project_dir / "artifacts" / "script.json").is_file():
        preview = stage_text_preview(project_dir, script_to_import_text(script), title=str(script.get("title") or ""))
        import_avatar_user_script(project_dir, {
            "import_token": preview["import_token"],
            "generation_mode": "manual_import",
            "import_mode": "longform",
            "background_mode": "opaque",
            "default_treatment": "pip_top_left",
            "replace_confirmed": False,
            "adopt_source_title": True,
        })
    apply_daily_story_contract(project_dir, script)
    _atomic_json(project_dir / "artifacts" / "daily_script.json", script)
    _ensure_daily_decision_log(project_dir)
    export_script_docx(script, project_dir / "docs" / f"{target_value.isoformat()}-科技快报脚本.docx")
    return project_dir


def _ensure_daily_decision_log(project_dir: Path, run: dict[str, Any] | None = None) -> dict[str, Any]:
    """Persist the user's bounded unattended-production approval and policy."""
    path = project_dir / "artifacts" / "decision_log.json"
    value = _read_json(path) or {"version": "1.0", "project_id": project_dir.name, "decisions": []}
    decisions = value.setdefault("decisions", [])
    provider_policy = run.get("provider_policy") if isinstance((run or {}).get("provider_policy"), dict) else {}
    standard_authorized = str(provider_policy.get("authorized_instance") or "").lower() in {"default", "standard", "standard_24gb"}
    global_standard = str(provider_policy.get("authorization_scope") or "") == "global_default"
    provider_decision = {
        "decision_id": "daily-provider-policy-v1",
        "stage": "production",
        "category": "provider_selection",
        "subject": "每日数字人 RunningHub 机型策略",
        "options_considered": [
            {"option_id": "enterprise-lite", "label": "企业 Lite", "score": 0.4 if standard_authorized else 1.0, "reason": "全局默认低成本运行"},
            {
                "option_id": "standard-24gb",
                "label": "Standard 24GB",
                "score": 1.0 if standard_authorized else 0.4,
                "reason": (
                    "用户已设为全局默认机型" if global_standard
                    else ("当前运行已获明确单次授权" if standard_authorized else "仅在明确显存不足时作为恢复选项")
                ),
                **({} if standard_authorized else {"rejected_because": "不得因排队、慢速、超时、限流或网络错误升级"}),
            },
        ],
        "selected": "standard-24gb" if standard_authorized else "enterprise-lite",
        "reason": (
            ("用户已将 Standard 24GB 设为每日默认机型；Plus 48GB 禁用且每期仍受5元预算约束。" if global_standard
             else f"用户针对 {run.get('target_date')} 明确授权 Standard 24GB；授权仅限该运行且受5元总预算约束。")
            if standard_authorized
            else "用户要求企业 Lite 始终优先，Standard 24GB 仅限明确 OOM/显存不足证据。"
        ),
        "user_visible": True,
        "user_approved": True,
        "confidence": 1.0,
    }
    wanted = [
        provider_decision,
        {
            "decision_id": "daily-budget-policy-v1",
            "stage": "production",
            "category": "budget_tradeoff",
            "subject": "每日无人值守 RunningHub 预算",
            "options_considered": [
                {"option_id": "cny-5-cap", "label": "人民币 5 元硬上限", "score": 1.0, "reason": "用户明确授权且限制成本"},
            ],
            "selected": "cny-5-cap",
            "reason": "总承诺费用达到 5 元后不再提交新的付费任务；已完成资产继续保留。",
            "user_visible": True,
            "user_approved": True,
            "confidence": 1.0,
        },
        {
            "decision_id": "daily-human-publish-gate-v1",
            "stage": "publish",
            "category": "composition_mode",
            "subject": "无人值守交付边界",
            "options_considered": [
                {"option_id": "review-candidate", "label": "生成待审全片预览", "score": 1.0, "reason": "允许早间人工抽检和局部热插拔"},
                {"option_id": "auto-publish", "label": "自动发布正式视频", "score": 0.0, "reason": "超出本次授权", "rejected_because": "正式发布必须人工确认"},
            ],
            "selected": "review-candidate",
            "reason": "自动化只生成待审预览，不替代人工发布门。",
            "user_visible": True,
            "user_approved": True,
            "confidence": 1.0,
        },
    ]
    replacements = {item["decision_id"]: item for item in wanted}
    retained = [item for item in decisions if not isinstance(item, dict) or str(item.get("decision_id")) not in replacements]
    decisions[:] = retained + wanted
    _atomic_json(path, value)
    return value


def run_research_and_script(target: date | str, *, trigger: str = "manual") -> dict[str, Any]:
    """Run every free stage and stop before Voicebox/RunningHub.

    This entry point is safe for the frontend's first feedback loop.  Paid
    media orchestration uses the same run manifest and starts only after this
    function has produced a validated script and project.
    """
    target_value = date.fromisoformat(target) if isinstance(target, str) else target
    if target_value >= datetime.now(LOCAL_TIMEZONE).date():
        raise DailyAutomationError("只能在目标自然日结束后生成科技快报，避免遗漏当天晚间新闻")
    run = create_or_resume_run(target_value, trigger=trigger)
    try:
        research_stage = run["stages"]["research"]
        if research_stage.get("status") != "succeeded":
            update_stage(run, "research", "running", message="正在检索并核验上一自然日科技新闻")
            research = collect_news_candidates(target_value)
            if len(research["candidates"]) < 3:
                raise DailyAutomationError(f"只找到 {len(research['candidates'])} 条窗口内候选新闻，至少需要 3 条")
            research_path = _run_path(target_value).parent / "news_research.json"
            _atomic_json(research_path, research)
            run["copy_skill_feed"] = research.get("copy_skill_feed") or {}
            update_stage(
                run,
                "research",
                "succeeded",
                message=f"已冻结 {len(research['candidates'])} 条候选新闻",
                output={
                    "candidate_count": len(research["candidates"]),
                    "artifact": str(research_path),
                    "copy_skill_feed": research.get("copy_skill_feed") or {},
                },
            )
        else:
            research = _read_json(_run_path(target_value).parent / "news_research.json") or {}
            if isinstance(research.get("copy_skill_feed"), dict):
                run["copy_skill_feed"] = research["copy_skill_feed"]

        if run["stages"]["script"].get("status") != "succeeded":
            update_stage(run, "script", "running", message="正在进行V2公域选题、冻结事实并生成中文双主持脚本")
            # Imports stay local because both V2 modules use the durable
            # primitives in this module.  The scheduled production path is now
            # the only authority; V2 is no longer a detached test artifact.
            from backlot.daily_text_resilience import generate_resilient_script_v2
            from backlot.news_selection_v2 import select_daily_news_v2

            run_dir = _run_path(target_value).parent
            selection_path = run_dir / "topic_selection.json"
            selection_v2_path = run_dir / "topic_selection_v2.json"
            audited_research_path = run_dir / "news_research_v2.json"
            script_path = run_dir / "daily_script.json"
            selection = _read_json(selection_v2_path)
            if not selection:
                selection = select_daily_news_v2(research)
                # Selection is independently valid. Persist it before dialogue
                # generation so local script repairs do not repeat a long model
                # assessment over the same candidates.
                _atomic_json(selection_v2_path, selection)
                _atomic_json(audited_research_path, research)
            def text_progress(state: str, message: str, details: dict[str, Any]) -> None:
                nonlocal run
                run = heartbeat_stage(
                    run,
                    "script",
                    message=message,
                    output={
                        "text_resilience": {
                            "state": state,
                            **details,
                        }
                    },
                )

            text_result = generate_resilient_script_v2(
                selection,
                research,
                run_dir=run_dir,
                policy=read_config().get("text_resilience") or {},
                progress=text_progress,
            )
            script = text_result["script"]
            selection = text_result["selection"]
            script["topic_selection"] = selection
            _atomic_json(selection_path, selection)
            _atomic_json(selection_v2_path, selection)
            _atomic_json(script_path, script)
            terminal_release = evaluate_media_release(script)
            fallback_candidate = (
                text_result["status"] == "awaiting_human"
                and terminal_release["decision"] == "fallback_review_candidate"
            )
            if text_result["status"] == "awaiting_human" and not fallback_candidate:
                stage = run["stages"]["script"]
                stage["status"] = "awaiting_human"
                stage["finished_at"] = _now()
                stage["message"] = "有界文本恢复已完成；已保留最佳稿，付费媒体未启动"
                stage["error"] = None
                stage["output"] = _merge_dict(
                    stage.get("output") or {},
                    {
                        "artifact": str(script_path),
                        "topic_selection_artifact": str(selection_path),
                        "topic_selection_v2_artifact": str(selection_v2_path),
                        "text_attempts_artifact": text_result["ledger_path"],
                        "text_resilience": {
                            "state": "awaiting_human",
                            "attempt_count": len(text_result["ledger"].get("attempts") or []),
                            "editorial_reviews_used": text_result["ledger"].get("editorial_reviews_used", 0),
                            "best_score": ((text_result["ledger"].get("best_candidate") or {}).get("editorial_score") or 0),
                        },
                    },
                )
                run["status"] = "awaiting_human"
                run["current_stage"] = "script"
                run.setdefault("approval_policy", {})["editorial_recovery_reason"] = (
                    text_result["ledger"].get("terminal_reason")
                    or "文本恢复额度已用完，最佳稿等待人工处理。"
                )
                _save_run(run)
                return read_run(target_value) or run
            validation = script["validation"]
            if fallback_candidate:
                text_result["ledger"]["recovery_state"] = "fallback_review_candidate"
                text_result["ledger"]["safe_resume_point"] = "review_only_media"
                text_result["ledger"]["next_action"] = "generate_review_candidate"
                text_result["ledger"]["terminal_reason"] = (
                    "最佳稿达到78分可靠线；继续生成待审视频，正式发布仍需人工确认。"
                )
                _atomic_json(run_dir / "daily_text_attempts.json", text_result["ledger"])
            update_stage(
                run,
                "script",
                "succeeded",
                message=(
                    f"V2已选 {len(selection.get('selected_stories') or [])} 条新闻；"
                    f"{validation['line_count']}句、约{validation['estimated_duration_seconds']:g}秒通过冻结事实与结构校验"
                ),
                output={
                    "artifact": str(script_path),
                    "topic_selection_artifact": str(selection_path),
                    "topic_selection_v2_artifact": str(selection_v2_path),
                    "audited_research_artifact": str(audited_research_path),
                    "text_attempts_artifact": text_result["ledger_path"],
                    "text_resilience": {
                        "state": "fallback_review_candidate" if fallback_candidate else "passed",
                        "attempt_count": len(text_result["ledger"].get("attempts") or []),
                        "editorial_reviews_used": text_result["ledger"].get("editorial_reviews_used", 0),
                        "best_score": ((text_result["ledger"].get("best_candidate") or {}).get("editorial_score") or 0),
                    },
                    "validation": validation,
                    "media_release_decision": terminal_release,
                },
            )
            if fallback_candidate:
                run.setdefault("approval_policy", {})["review_candidate_notice"] = (
                    "文本恢复已用完，但最佳稿达到78分可靠线；继续生成待审视频，不自动发布。"
                )
        else:
            script = _read_json(_run_path(target_value).parent / "daily_script.json") or {}

        if run["stages"]["project"].get("status") != "succeeded":
            update_stage(run, "project", "running", message="正在创建 9:16 数字人口播项目")
            project_dir = create_daily_project(target_value, script)
            run["project_id"] = project_dir.name
            _save_run(run)
            _ensure_daily_decision_log(project_dir, run)
            update_stage(run, "project", "succeeded", message="项目与双版本脚本已初始化", output={"project_id": project_dir.name})
        run = read_run(target_value) or run
        release = evaluate_media_release(script)
        run["media_release_decision"] = release
        if release["decision"] == "blocked":
            run["status"] = "blocked"
            run["current_stage"] = "voice"
        elif release["decision"] == "fallback_review_candidate":
            run["status"] = "queued"
            run["current_stage"] = "voice"
            run.setdefault("approval_policy", {})["review_candidate_notice"] = (
                "脚本达到可靠可用线，将自动生成待审视频；正式发布仍需人工确认。"
            )
        else:
            run["status"] = "queued"
            run["current_stage"] = "voice"
        _save_run(run)
        return read_run(target_value) or run
    except Exception as exc:
        current = str(run.get("current_stage") or "research")
        if current in STAGE_ORDER:
            update_stage(run, current, "failed", message="自动化阶段失败，可从本阶段重试", error=str(exc))
        raise


_RUN_LOCK = threading.Lock()
_ACTIVE_RUN_LOCK: dict[str, str] | None = None


def _process_is_alive(pid: int) -> bool:
    if pid <= 0:
        return False
    if os.name == "nt":
        return _windows_process_is_alive(pid)
    try:
        os.kill(pid, 0)
    except OSError:
        return False
    return True


def _windows_process_is_alive(pid: int) -> bool:
    """Query a Windows process handle without sending a console signal."""
    import ctypes
    from ctypes import wintypes

    process_query_limited_information = 0x1000
    still_active = 259
    kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
    kernel32.OpenProcess.argtypes = [wintypes.DWORD, wintypes.BOOL, wintypes.DWORD]
    kernel32.OpenProcess.restype = wintypes.HANDLE
    kernel32.GetExitCodeProcess.argtypes = [wintypes.HANDLE, ctypes.POINTER(wintypes.DWORD)]
    kernel32.GetExitCodeProcess.restype = wintypes.BOOL
    kernel32.CloseHandle.argtypes = [wintypes.HANDLE]
    kernel32.CloseHandle.restype = wintypes.BOOL
    handle = kernel32.OpenProcess(process_query_limited_information, False, pid)
    if not handle:
        # Access denied still proves that a protected process owns the PID.
        return ctypes.get_last_error() == 5
    try:
        exit_code = wintypes.DWORD()
        if not kernel32.GetExitCodeProcess(handle, ctypes.byref(exit_code)):
            return False
        return exit_code.value == still_active
    finally:
        kernel32.CloseHandle(handle)


def _remove_stale_run_lock() -> bool:
    lock = _read_json(RUN_LOCK_PATH)
    if lock and _process_is_alive(int(lock.get("pid") or 0)):
        return False
    try:
        RUN_LOCK_PATH.unlink(missing_ok=True)
    except OSError:
        return False
    return True


def try_acquire_run_lock(target: date | str | None = None, *, trigger: str = "unknown") -> bool:
    """Acquire one machine-wide daily production lock.

    The previous in-process mutex did not protect a server process and a
    Windows scheduled CLI process from submitting the same paid task.  This
    exclusive lock file survives browser/server boundaries and is reclaimed
    only after its owner PID is no longer alive.
    """
    global _ACTIVE_RUN_LOCK
    if not _RUN_LOCK.acquire(blocking=False):
        return False
    token = uuid.uuid4().hex
    payload = {
        "version": "1.0",
        "token": token,
        "pid": os.getpid(),
        "target_date": str(target or "unknown"),
        "trigger": str(trigger or "unknown"),
        "acquired_at": _now(),
    }
    RUN_LOCK_PATH.parent.mkdir(parents=True, exist_ok=True)
    for attempt in range(2):
        try:
            descriptor = os.open(RUN_LOCK_PATH, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
        except FileExistsError:
            if attempt == 0 and _remove_stale_run_lock():
                continue
            _RUN_LOCK.release()
            return False
        try:
            os.write(descriptor, (json.dumps(payload, ensure_ascii=False, indent=2) + "\n").encode("utf-8"))
        finally:
            os.close(descriptor)
        _ACTIVE_RUN_LOCK = {"token": token, "path": str(RUN_LOCK_PATH)}
        return True
    _RUN_LOCK.release()
    return False


def release_run_lock() -> None:
    global _ACTIVE_RUN_LOCK
    active = _ACTIVE_RUN_LOCK
    if active:
        persisted = _read_json(RUN_LOCK_PATH)
        if persisted and persisted.get("token") == active.get("token"):
            RUN_LOCK_PATH.unlink(missing_ok=True)
        _ACTIVE_RUN_LOCK = None
    if _RUN_LOCK.locked():
        _RUN_LOCK.release()


def scheduler_command() -> list[str]:
    python = REPO_ROOT / ".venv" / "Scripts" / "python.exe"
    if not python.is_file():
        raise DailyAutomationError("项目虚拟环境不存在，无法创建稳定的每日计划任务")
    wrapper = REPO_ROOT / "scripts" / "run_daily_automation.py"
    if not wrapper.is_file():
        raise DailyAutomationError("每日自动化启动脚本不存在")
    return [str(python), str(wrapper), "run", "--previous-day", "--trigger", "schedule"]


def scheduler_task_spec(config: dict[str, Any] | None = None) -> dict[str, Any]:
    config = config or read_config()
    return {
        "task_name": "OpenMontage-Daily-Tech-Brief",
        "enabled": bool(config["enabled"]),
        "schedule_time": config["schedule_time"],
        "working_directory": str(REPO_ROOT),
        "command": scheduler_command(),
        "start_when_available": True,
        "multiple_instances": "ignore_new",
        "execution_time_limit_hours": 12,
    }


def _windows_command_line(parts: list[str]) -> str:
    return subprocess.list2cmdline(parts)


def _scheduler_task_xml(spec: dict[str, Any], *, username: str | None = None) -> bytes:
    """Build a Task Scheduler contract that survives sleep/reboot and avoids duplicates."""
    namespace = "http://schemas.microsoft.com/windows/2004/02/mit/task"
    ET.register_namespace("", namespace)
    task = ET.Element(f"{{{namespace}}}Task", {"version": "1.4"})
    registration = ET.SubElement(task, f"{{{namespace}}}RegistrationInfo")
    ET.SubElement(registration, f"{{{namespace}}}Author").text = "OpenMontage"
    triggers = ET.SubElement(task, f"{{{namespace}}}Triggers")
    calendar = ET.SubElement(triggers, f"{{{namespace}}}CalendarTrigger")
    today = datetime.now(LOCAL_TIMEZONE).date().isoformat()
    ET.SubElement(calendar, f"{{{namespace}}}StartBoundary").text = f"{today}T{spec['schedule_time']}:00+08:00"
    ET.SubElement(calendar, f"{{{namespace}}}Enabled").text = "true"
    by_day = ET.SubElement(calendar, f"{{{namespace}}}ScheduleByDay")
    ET.SubElement(by_day, f"{{{namespace}}}DaysInterval").text = "1"
    principals = ET.SubElement(task, f"{{{namespace}}}Principals")
    principal = ET.SubElement(principals, f"{{{namespace}}}Principal", {"id": "Author"})
    user = username or "\\".join(filter(None, (os.environ.get("USERDOMAIN"), os.environ.get("USERNAME"))))
    if user:
        ET.SubElement(principal, f"{{{namespace}}}UserId").text = user
    ET.SubElement(principal, f"{{{namespace}}}LogonType").text = "InteractiveToken"
    ET.SubElement(principal, f"{{{namespace}}}RunLevel").text = "LeastPrivilege"
    settings = ET.SubElement(task, f"{{{namespace}}}Settings")
    for name, value in (
        ("MultipleInstancesPolicy", "IgnoreNew"),
        ("DisallowStartIfOnBatteries", "false"),
        ("StopIfGoingOnBatteries", "false"),
        ("AllowHardTerminate", "true"),
        ("StartWhenAvailable", "true"),
        ("WakeToRun", "true"),
        ("RunOnlyIfNetworkAvailable", "true"),
        ("Enabled", "true"),
        ("ExecutionTimeLimit", f"PT{int(spec['execution_time_limit_hours'])}H"),
        ("Priority", "7"),
    ):
        ET.SubElement(settings, f"{{{namespace}}}{name}").text = value
    actions = ET.SubElement(task, f"{{{namespace}}}Actions", {"Context": "Author"})
    execute = ET.SubElement(actions, f"{{{namespace}}}Exec")
    command = list(spec["command"])
    ET.SubElement(execute, f"{{{namespace}}}Command").text = command[0]
    ET.SubElement(execute, f"{{{namespace}}}Arguments").text = _windows_command_line(command[1:])
    ET.SubElement(execute, f"{{{namespace}}}WorkingDirectory").text = str(spec["working_directory"])
    return ET.tostring(task, encoding="utf-16", xml_declaration=True)


def sync_windows_scheduler(config: dict[str, Any] | None = None) -> dict[str, Any]:
    """Create or disable the per-user Windows scheduled task.

    ``schtasks`` is intentionally called with a fixed task name and an
    absolute project-venv command.  No user-supplied command text reaches the
    shell.
    """
    settings = config or read_config()
    spec = scheduler_task_spec(settings)
    if os.name != "nt":
        return {**spec, "installed": False, "platform_supported": False, "detail": "当前系统不是 Windows，未创建计划任务"}
    task_name = str(spec["task_name"])
    if settings.get("enabled") is True:
        handle, temporary = tempfile.mkstemp(prefix="openmontage-daily-", suffix=".xml")
        try:
            with os.fdopen(handle, "wb") as stream:
                stream.write(_scheduler_task_xml(spec))
            result = subprocess.run(
                ["schtasks", "/Create", "/TN", task_name, "/XML", temporary, "/F"],
                capture_output=True, text=True, errors="replace", timeout=30, check=False,
            )
        finally:
            Path(temporary).unlink(missing_ok=True)
        if result.returncode != 0:
            detail = (result.stderr or result.stdout or "无法创建 Windows 计划任务").strip()
            raise DailyAutomationError(f"每日自动运行开关保存失败：{detail[:400]}")
        return {**spec, "installed": True, "platform_supported": True, "detail": "Windows 计划任务已启用"}
    query = subprocess.run(
        ["schtasks", "/Query", "/TN", task_name], capture_output=True, text=True, errors="replace", timeout=15, check=False,
    )
    if query.returncode == 0:
        result = subprocess.run(
            ["schtasks", "/Change", "/TN", task_name, "/DISABLE"],
            capture_output=True, text=True, errors="replace", timeout=15, check=False,
        )
        if result.returncode != 0:
            detail = (result.stderr or result.stdout or "无法禁用 Windows 计划任务").strip()
            raise DailyAutomationError(f"每日自动运行开关保存失败：{detail[:400]}")
        return {**spec, "installed": True, "platform_supported": True, "detail": "Windows 计划任务已停用"}
    return {**spec, "installed": False, "platform_supported": True, "detail": "每日计划任务尚未创建"}


def scheduler_runtime_status() -> dict[str, Any]:
    spec = scheduler_task_spec()
    if os.name != "nt":
        return {**spec, "installed": False, "platform_supported": False}
    result = subprocess.run(
        ["schtasks", "/Query", "/TN", str(spec["task_name"]), "/FO", "LIST", "/V"],
        capture_output=True, text=True, errors="replace", timeout=15, check=False,
    )
    raw_status = (result.stdout or "")[:3000] if result.returncode == 0 else ""
    return _parse_scheduler_runtime_output(raw_status, installed=result.returncode == 0, spec=spec)


def _parse_scheduler_runtime_output(raw_status: str, *, installed: bool, spec: dict[str, Any]) -> dict[str, Any]:
    """Normalize English or Chinese ``schtasks /FO LIST /V`` output."""
    fields: dict[str, str] = {}
    for line in raw_status.splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        fields[key.strip().lower()] = value.strip()
    state_text = (
        fields.get("scheduled task state") or fields.get("计划任务状态")
        or fields.get("status") or fields.get("模式") or ""
    )
    runtime_enabled = installed and state_text.lower() not in {"disabled", "已禁用", "禁用"}
    task_command = fields.get("task to run") or fields.get("要运行的任务") or ""
    command_matches = not task_command or all(str(part).lower() in task_command.lower() for part in spec["command"][:2])
    last_result_text = fields.get("last result") or fields.get("上次结果") or ""
    try:
        last_result: int | str | None = int(last_result_text, 0)
    except (TypeError, ValueError):
        last_result = last_result_text or None
    return {
        **spec,
        "installed": installed,
        "platform_supported": True,
        "runtime_enabled": runtime_enabled,
        "runtime_state": state_text or ("missing" if not installed else "unknown"),
        "command_matches": command_matches,
        "next_run_time": fields.get("next run time") or fields.get("下次运行时间") or None,
        "last_run_time": fields.get("last run time") or fields.get("上次运行时间") or None,
        "last_result": last_result,
        "raw_status": raw_status,
    }


def scheduler_effective_state(config: dict[str, Any], scheduler: dict[str, Any]) -> dict[str, Any]:
    desired = bool(config.get("enabled"))
    installed = bool(scheduler.get("installed"))
    runtime_enabled = bool(scheduler.get("runtime_enabled"))
    command_matches = bool(scheduler.get("command_matches", True))
    configured_healthy = (installed and runtime_enabled and command_matches) if desired else (not installed or not runtime_enabled)
    last_result = scheduler.get("last_result")
    last_run_succeeded = last_result in {None, 0, "0", "0x0"}
    healthy = configured_healthy and (last_run_succeeded or not desired)
    effective_enabled = desired and configured_healthy
    conflict = desired != effective_enabled or (desired and not command_matches)
    if effective_enabled and not last_run_succeeded:
        message = f"凌晨调度已生效，但上次任务退出码 {last_result}；请检查失败阶段或等待同日期恢复。"
    elif effective_enabled:
        message = "项目配置与 Windows 计划任务一致，将按设定时间自动运行。"
    elif desired:
        message = "项目配置已开启，但 Windows 计划任务未生效；凌晨任务不会自动运行。"
    elif installed and runtime_enabled:
        message = "项目配置已关闭，但 Windows 计划任务仍在运行；请重新关闭一次自动化。"
    else:
        message = "每日自动生产当前已关闭。"
    return {
        "desired_enabled": desired,
        "effective_enabled": effective_enabled,
        "healthy": healthy,
        "configured_healthy": configured_healthy,
        "last_run_succeeded": last_run_succeeded,
        "last_result": last_result,
        "conflict": conflict,
        "message": message,
    }
